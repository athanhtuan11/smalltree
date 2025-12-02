from werkzeug.security import generate_password_hash
from PIL import Image
from flask import Blueprint, render_template, request, redirect, url_for, flash, send_file, session, jsonify, current_app
from app.models import db, Activity, Curriculum, Child, AttendanceRecord, Staff, BmiRecord, ActivityImage, Supplier, Product, StudentAlbum, StudentPhoto, StudentProgress, Dish, Menu, Class, MonthlyService, UserActivity
from app.forms import EditProfileForm, ActivityCreateForm, ActivityEditForm, SupplierForm, ProductForm
from calendar import monthrange
from datetime import datetime, date, timedelta
import io, zipfile, os, json, re, secrets, tempfile

# Cloudflare R2 Storage
try:
    from r2_storage import get_r2_storage
    R2_ENABLED = True
except ImportError:
    R2_ENABLED = False
    print("‚ö†Ô∏è  R2 Storage kh√¥ng kh·∫£ d·ª•ng. ·∫¢nh s·∫Ω l∆∞u local.")

# Check for optional dependencies
try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    from docx.shared import RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

main = Blueprint('main', __name__)

def assess_child_bmi(bmi, age_months, gender='unknown'):
    """
    ƒê√°nh gi√° BMI cho tr·∫ª em theo percentile ƒë∆°n gi·∫£n h√≥a
    Tr·∫£ v·ªÅ: "Thi·∫øu c√¢n: <5%", "B√¨nh th∆∞·ªùng: 5-85%", "Th·ª´a c√¢n: 85-95%", "B√©o ph√¨: >95%"
    """
    if age_months < 0 or bmi is None:
        return 'Ch∆∞a c√≥ ƒë·ªß th√¥ng tin'
    
    # Ng∆∞·ª°ng BMI ƒë∆°n gi·∫£n theo tu·ªïi (x·∫•p x·ªâ percentile 5%, 85%, 95%)
    if age_months < 24:  # 0-2 tu·ªïi
        p5, p85, p95 = 13.5, 17.5, 18.5
    elif age_months < 36:  # 2-3 tu·ªïi  
        p5, p85, p95 = 13.5, 17.0, 18.0
    elif age_months < 48:  # 3-4 tu·ªïi
        p5, p85, p95 = 13.5, 16.5, 17.5
    elif age_months < 60:  # 4-5 tu·ªïi
        p5, p85, p95 = 14.0, 16.5, 17.5
    else:  # > 5 tu·ªïi
        p5, p85, p95 = 14.0, 17.0, 18.5
    
    if bmi < p5:
        return 'Thi·∫øu c√¢n: <5%'
    elif bmi < p85:
        return 'B√¨nh th∆∞·ªùng: 5-85%'
    elif bmi < p95:
        return 'Th·ª´a c√¢n: 85-95%'
    else:
        return 'B√©o ph√¨: >95%'

def log_activity(action, resource_type=None, resource_id=None, description=None):
    """Helper function ƒë·ªÉ ghi nh·∫≠n ho·∫°t ƒë·ªông ng∆∞·ªùi d√πng"""
    try:
        from datetime import datetime, timezone, timedelta
        user_type = session.get('role', 'guest')
        user_id = session.get('user_id')
        user_name = session.get('name', 'Kh√°ch v√£ng lai')
        
        # L·∫•y th·ªùi gian Vi·ªát Nam (UTC+7)
        vietnam_tz = timezone(timedelta(hours=7))
        vietnam_time = datetime.now(vietnam_tz)
        
        activity = UserActivity(
            user_id=user_id,
            user_type=user_type,
            user_name=user_name,
            action=action,
            resource_type=resource_type,
            resource_id=resource_id,
            description=description,
            ip_address=request.remote_addr,
            user_agent=request.headers.get('User-Agent', '')[:500],
            timestamp=vietnam_time
        )
        db.session.add(activity)
        db.session.commit()
    except Exception as e:
        db.session.rollback()  # Rollback transaction ƒë·ªÉ kh√¥ng ·∫£nh h∆∞·ªüng DB
        print(f"[ERROR] Failed to log activity: {str(e)}")
        # Kh√¥ng raise exception ƒë·ªÉ kh√¥ng ·∫£nh h∆∞·ªüng lu·ªìng ch√≠nh
        # N·∫øu b·∫£ng UserActivity ch∆∞a t·ªìn t·∫°i, app v·∫´n ch·∫°y b√¨nh th∆∞·ªùng

def redirect_no_permission():
    try:
        log_activity('access_denied', description=f'Attempted to access: {request.path}')
    except:
        pass  # B·ªè qua n·∫øu b·∫£ng ch∆∞a t·ªìn t·∫°i
    flash('B·∫°n kh√¥ng c√≥ quy·ªÅn truy c·∫≠p ch·ª©c nƒÉng n√†y!', 'danger')
    return redirect(url_for('main.login'))

def get_class_order(class_name):
    """ƒê·ªãnh nghƒ©a th·ª© t·ª± s·∫Øp x·∫øp l·ªõp h·ªçc"""
    class_order = {
        'L·ªõp M·∫ßm': 1,
        'L·ªõp Ch·ªìi': 2, 
        'L·ªõp L√°': 3,
        'Kay 01': 4,
        'Kay01': 4,
        'Kay 02': 5,
        'Kay02': 5,
        'Kay 03': 6,
        'Kay03': 6,
    }
    return class_order.get(class_name, 999)  # 999 cho l·ªõp kh√¥ng x√°c ƒë·ªãnh

def optimize_image(file_stream, max_size=(1200, 900), quality=85):
    """
    T·ªëi ∆∞u h√≥a ·∫£nh: resize v√† compress - LU√îN TH√ÄNH C√îNG
    Args:
        file_stream: File stream c·ªßa ·∫£nh
        max_size: K√≠ch th∆∞·ªõc t·ªëi ƒëa (width, height)
        quality: Ch·∫•t l∆∞·ª£ng JPEG (1-100)
    Returns:
        Tuple (optimized_image_data, format)
    """
    try:
        file_stream.seek(0)
        img = Image.open(file_stream)
        
        # Convert b·∫•t c·ª© format n√†o v·ªÅ RGB ƒë·ªÉ ƒë·∫£m b·∫£o t∆∞∆°ng th√≠ch
        if img.mode in ('RGBA', 'LA', 'P', 'CMYK', '1', 'L'):
            if img.mode == 'RGBA':
                # T·∫°o background tr·∫Øng cho ·∫£nh trong su·ªët
                background = Image.new('RGB', img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[-1])
                img = background
            elif img.mode == 'P':
                img = img.convert('RGB')
            elif img.mode in ('CMYK', 'LAB'):
                img = img.convert('RGB')
            elif img.mode in ('1', 'L', 'LA'):
                img = img.convert('RGB')
        
        # Resize if qu√° l·ªõn - lu√¥n resize v·ªÅ k√≠ch th∆∞·ªõc h·ª£p l√Ω
        original_size = img.size
        if img.size[0] > max_size[0] or img.size[1] > max_size[1]:
            img.thumbnail(max_size, Image.Resampling.LANCZOS)
            print(f"[INFO] Resize ·∫£nh t·ª´ {original_size} xu·ªëng {img.size}")
        
        # Gi·∫£m ch·∫•t l∆∞·ª£ng d·∫ßn n·∫øu file v·∫´n qu√° l·ªõn
        import io
        output = io.BytesIO()
        img_format = 'JPEG'  # Always save as JPEG for consistency
        
        # Th·ª≠ c√°c m·ª©c ch·∫•t l∆∞·ª£ng kh√°c nhau
        for test_quality in [quality, 70, 50, 30]:
            output.seek(0)
            output.truncate()
            img.save(output, format=img_format, quality=test_quality, optimize=True)
            
            # N·∫øu ·∫£nh nh·ªè h∆°n 2MB th√¨ OK
            if output.tell() <= 2 * 1024 * 1024:  # 2MB
                break
            print(f"[INFO] Gi·∫£m ch·∫•t l∆∞·ª£ng xu·ªëng {test_quality}% ƒë·ªÉ t·ªëi ∆∞u k√≠ch th∆∞·ªõc")
        
        output.seek(0)
        return output, img_format
        
    except Exception as e:
        print(f"[ERROR] L·ªói t·ªëi ∆∞u ·∫£nh: {str(e)}")
        # Fallback: t·∫°o ·∫£nh placeholder nh·ªè
        import io
        placeholder_img = Image.new('RGB', (400, 300), color=(200, 200, 200))
        output = io.BytesIO()
        placeholder_img.save(output, format='JPEG', quality=80)
        output.seek(0)
        return output, 'JPEG'

def verify_and_repair_image(file_stream):
    """
    Ki·ªÉm tra v√† s·ª≠a ·∫£nh b·ªã l·ªói
    Returns: (is_readable, repaired_stream)
    """
    try:
        file_stream.seek(0)
        img = Image.open(file_stream)
        img.verify()  # Ki·ªÉm tra integrity
        file_stream.seek(0)  # Reset l·∫°i ƒë·ªÉ ƒë·ªçc l·∫°i
        img = Image.open(file_stream)  # Open l·∫°i sau verify
        
        # Th·ª≠ load to√†n b·ªô ·∫£nh ƒë·ªÉ ƒë·∫£m b·∫£o kh√¥ng corrupt
        img.load()
        return True, file_stream
    except Exception as e:
        print(f"[WARNING] ·∫¢nh b·ªã l·ªói, th·ª≠ s·ª≠a ch·ªØa: {e}")
        try:
            # Th·ª≠ ƒë·ªçc l·∫°i v·ªõi mode kh√°c nhau
            file_stream.seek(0)
            img = Image.open(file_stream)
            
            # Convert v·ªÅ RGB ƒë·ªÉ fix m·ªôt s·ªë l·ªói
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # T·∫°o stream m·ªõi v·ªõi ·∫£nh ƒë√£ s·ª≠a
            import io
            repaired_stream = io.BytesIO()
            img.save(repaired_stream, format='JPEG', quality=90)
            repaired_stream.seek(0)
            return True, repaired_stream
        except Exception as e2:
            print(f"[ERROR] Kh√¥ng th·ªÉ s·ª≠a ·∫£nh: {e2}")
            return False, None

def validate_image_file(file, max_size_mb=50):  # TƒÉng l√™n 50MB ƒë·ªÉ ch·∫•p nh·∫≠n h·∫ßu h·∫øt file
    """
    Validate file ·∫£nh - B√ÇY GI·ªú CH·ªà KI·ªÇM TRA C∆† B·∫¢N, KH√îNG T·ª™ CH·ªêI
    Args:
        file: FileStorage object
        max_size_mb: K√≠ch th∆∞·ªõc t·ªëi ƒëa (MB) - ch·ªâ ƒë·ªÉ warning
    Returns:
        Tuple (is_valid, warning_message) - Lu√¥n tr·∫£ True ƒë·ªÉ ch·∫•p nh·∫≠n
    """
    if not file or not file.filename:
        return False, "Kh√¥ng c√≥ file ƒë∆∞·ª£c ch·ªçn"
    
    # Check extension - ch·∫•p nh·∫≠n t·∫•t c·∫£ ·∫£nh ph·ªï bi·∫øn
    allowed_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.jfif', '.webp', '.bmp', '.tiff', '.svg'}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_extensions:
        return False, f"File {file.filename} kh√¥ng ph·∫£i ·∫£nh, s·∫Ω b·ªè qua"
    
    # Check size - CH·ªà WARNING, KH√îNG T·ª™ CH·ªêI
    try:
        file.seek(0, 2)
        size = file.tell()
        file.seek(0)
        size_mb = size / (1024 * 1024)
        if size_mb > max_size_mb:
            return True, f"File {file.filename} l·ªõn ({size_mb:.1f}MB), s·∫Ω ƒë∆∞·ª£c n√©n t·ª± ƒë·ªông"
        else:
            return True, None  # File OK
    except Exception:
        return True, f"Kh√¥ng ƒë·ªçc ƒë∆∞·ª£c k√≠ch th∆∞·ªõc {file.filename}, s·∫Ω th·ª≠ x·ª≠ l√Ω"
    
    # Check file size
    file.seek(0, 2)  # Seek to end
    size = file.tell()
    file.seek(0)  # Reset to beginning
    
    if size > max_size_mb * 1024 * 1024:
        return False, f"File qu√° l·ªõn: {size // (1024*1024)}MB > {max_size_mb}MB"
    
    # Try to open as image
    try:
        file.seek(0)
        img = Image.open(file.stream)
        img.verify()  # Verify it's a valid image
        file.seek(0)  # Reset stream
        return True, ""
    except Exception as e:
        return False, f"File kh√¥ng ph·∫£i l√† ·∫£nh h·ª£p l·ªá: {str(e)}"

# CRUD Class

@main.route('/classes/new', methods=['GET', 'POST'])
def new_class():
    if session.get('role') != 'admin' and session.get('role') != 'teacher':
        flash('B·∫°n kh√¥ng c√≥ quy·ªÅn t·∫°o l·ªõp m·ªõi!', 'danger')
        return redirect(url_for('main.attendance'))
    # Th√™m l·ªõp m·ªõi
    if request.method == 'POST':
        class_name = request.form.get('class_name')
        description = request.form.get('description')
        if not class_name or len(class_name) < 3:
            flash('T√™n l·ªõp ph·∫£i c√≥ √≠t nh·∫•t 3 k√Ω t·ª±!', 'danger')
            return redirect(url_for('main.new_class'))
        existing = Class.query.filter_by(name=class_name).first()
        if existing:
            flash('L·ªõp n√†y ƒë√£ t·ªìn t·∫°i!', 'warning')
            return redirect(url_for('main.new_class'))
        new_class = Class(name=class_name, description=description)
        db.session.add(new_class)
        db.session.commit()
        log_activity('create', 'class', new_class.id, f'T·∫°o l·ªõp: {class_name}')
        flash(f'ƒê√£ t·∫°o l·ªõp m·ªõi: {class_name}', 'success')
        return redirect(url_for('main.new_class'))
    # Hi·ªÉn th·ªã danh s√°ch l·ªõp
    classes = Class.query.order_by(Class.name).all()
    mobile = is_mobile()
    return render_template('new_class.html', title='T·∫°o L·ªõp m·ªõi', mobile=mobile, classes=classes)

@main.route('/classes/<int:class_id>/edit', methods=['GET', 'POST'])
def edit_class(class_id):
    if session.get('role') != 'admin' and session.get('role') != 'teacher':
        flash('B·∫°n kh√¥ng c√≥ quy·ªÅn s·ª≠a l·ªõp!', 'danger')
        return redirect(url_for('main.new_class'))
    class_obj = Class.query.get_or_404(class_id)
    if request.method == 'POST':
        class_obj.name = request.form.get('class_name')
        class_obj.description = request.form.get('description')
        db.session.commit()
        log_activity('edit', 'class', class_id, f'S·ª≠a l·ªõp: {class_obj.name}')
        flash('ƒê√£ c·∫≠p nh·∫≠t l·ªõp!', 'success')
        return redirect(url_for('main.new_class'))
    return render_template('edit_class.html', class_obj=class_obj)

@main.route('/classes/<int:class_id>/delete', methods=['POST'])
def delete_class(class_id):
    if session.get('role') != 'admin' and session.get('role') != 'teacher':
        flash('B·∫°n kh√¥ng c√≥ quy·ªÅn x√≥a l·ªõp!', 'danger')
        return redirect(url_for('main.new_class'))
    class_obj = Class.query.get_or_404(class_id)
    class_name = class_obj.name
    db.session.delete(class_obj)
    db.session.commit()
    log_activity('delete', 'class', class_id, f'X√≥a l·ªõp: {class_name}')
    flash('ƒê√£ x√≥a l·ªõp!', 'success')
    return redirect(url_for('main.new_class'))




@main.route('/attendance/save', methods=['POST'])
def save_attendance():
    if not session.get('role'):
        flash('B·∫°n ph·∫£i ƒëƒÉng nh·∫≠p m·ªõi truy c·∫≠p ƒë∆∞·ª£c trang n√†y!', 'danger')
        return redirect(url_for('main.about'))
    from datetime import date
    attendance_date = request.form.get('attendance_date') or date.today().strftime('%Y-%m-%d')
    selected_class = request.form.get('class_name')
    # L∆∞u h√†ng lo·∫°t (kh√¥ng c√≥ student_id ri√™ng l·∫ª)
    if selected_class and selected_class != 'None':
        students = Child.query.filter_by(class_name=selected_class, is_active=True).all()
    else:
        students = Child.query.filter_by(is_active=True).all()
    for student in students:
        present_value = request.form.get(f'present_{student.id}')
        if present_value == 'yes':
            status = 'C√≥ m·∫∑t'
        elif present_value == 'absent_excused':
            status = 'V·∫Øng m·∫∑t c√≥ ph√©p'
        elif present_value == 'absent_unexcused':
            status = 'V·∫Øng m·∫∑t kh√¥ng ph√©p'
        else:
            status = 'V·∫Øng'
        breakfast = request.form.get(f'breakfast_{student.id}')
        lunch = request.form.get(f'lunch_{student.id}')
        snack = request.form.get(f'snack_{student.id}')
        toilet = request.form.get(f'toilet_{student.id}')
        toilet_times = request.form.get(f'toilet_times_{student.id}') or None
        note = request.form.get(f'note_{student.id}')
        record = AttendanceRecord.query.filter_by(child_id=student.id, date=attendance_date).first()
        if record:
            record.status = status
            record.breakfast = breakfast
            record.lunch = lunch
            record.snack = snack
            record.toilet = toilet
            record.toilet_times = toilet_times
            record.note = note
        else:
            record = AttendanceRecord(child_id=student.id, date=attendance_date, status=status,
                                     breakfast=breakfast, lunch=lunch, snack=snack, toilet=toilet, toilet_times=toilet_times, note=note)
            db.session.add(record)
    db.session.commit()
    #flash('ƒê√£ l∆∞u ƒëi·ªÉm danh!', 'success')
    if not selected_class or selected_class == 'None':
        return redirect(url_for('main.attendance', attendance_date=attendance_date))
    return redirect(url_for('main.attendance', attendance_date=attendance_date, class_name=selected_class))

def is_mobile():
    ua = request.user_agent.string.lower()
    return 'mobile' in ua or 'android' in ua or 'iphone' in ua

def calculate_age(birth_date):
    today = datetime.today()
    try:
        birthday = datetime.strptime(birth_date, '%Y-%m-%d')
    except Exception:
        return 0
    age = today.year - birthday.year - ((today.month, today.day) < (birthday.month, birthday.day))
    return age

MAX_LOGIN_ATTEMPTS = 5
LOCKOUT_TIME_MINUTES = 10
LOGIN_COOLDOWN_SECONDS = 30
login_attempts = {}
lockout_until = {}
last_login_time = {}

# --- GLOBAL ERROR HANDLER FOR API JSON RESPONSE ---
from werkzeug.exceptions import HTTPException

@main.errorhandler(Exception)
def handle_api_exception(e):
    from flask import request
    import traceback
    # Ch·ªâ tr·∫£ v·ªÅ JSON n·∫øu l√† API (application/json ho·∫∑c /ai/ route)
    if request.path.startswith('/ai/') or request.is_json or request.headers.get('Accept', '').startswith('application/json'):
        code = 500
        if isinstance(e, HTTPException):
            code = e.code
        return jsonify({
            'success': False,
            'error': str(e),
            'trace': traceback.format_exc()[:1000]  # Gi·ªõi h·∫°n ƒë·ªô d√†i trace cho debug
        }), code
    # N·∫øu kh√¥ng ph·∫£i API th√¨ tr·∫£ v·ªÅ m·∫∑c ƒë·ªãnh
    raise e
# Password hash/check imports (with fallback)
try:
    from werkzeug.security import generate_password_hash, check_password_hash
    WERKZEUG_AVAILABLE = True
except ImportError:
    print("Warning: werkzeug.security not available, using fallback")
    import hashlib
    WERKZEUG_AVAILABLE = False
    def generate_password_hash(password):
        return hashlib.sha256(password.encode()).hexdigest()
    def check_password_hash(hash_password, password):
        return hash_password == hashlib.sha256(password.encode()).hexdigest()


# ================== DANH S√ÅCH M√ìN ƒÇN ==================

@main.route('/dish-list')
def dish_list():
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    try:
        dishes = Dish.query.all()
        mobile = is_mobile()
        return render_template('dish_list.html', dishes=dishes, mobile=mobile)
    except Exception as e:
        flash(f'L·ªói khi t·∫£i danh s√°ch m√≥n ƒÉn: {str(e)}', 'danger')
        return redirect(url_for('main.menu'))

# Route ƒë·ªÉ b·∫≠t/t·∫Øt tr·∫°ng th√°i m√≥n ƒÉn
@main.route('/dish/<int:dish_id>/toggle-active', methods=['POST'])
def toggle_dish_active(dish_id):
    from app.models import Dish, db
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    dish = Dish.query.get_or_404(dish_id)
    dish.is_active = not dish.is_active
    db.session.commit()
    flash(f"ƒê√£ {'b·∫≠t' if dish.is_active else '·∫©n'} m√≥n ƒÉn!", 'success')
    return redirect(url_for('main.dish_list'))

# ================== S·ª¨A/X√ìA M√ìN ƒÇN ==================
@main.route('/dish/<int:dish_id>/edit', methods=['GET', 'POST'])
def edit_dish(dish_id):
    from app.models import Dish, DishIngredient, Product, db
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    dish = Dish.query.get_or_404(dish_id)
    products = Product.query.all()
    product_units = sorted(list(set([p.unit for p in products if p.unit])))
    if request.method == 'POST':
        dish.name = request.form.get('name')
        dish.description = request.form.get('description')
        meal_times = request.form.getlist('meal_times')
        dish.meal_times = meal_times
        # X√≥a nguy√™n li·ªáu c≈©
        DishIngredient.query.filter_by(dish_id=dish.id).delete()
        ingredient_ids = request.form.getlist('ingredient_id')
        units = request.form.getlist('unit')
        quantities = request.form.getlist('quantity')
        for idx, pid in enumerate(ingredient_ids):
            if not pid or not units[idx] or not quantities[idx]:
                continue
            di = DishIngredient(
                dish_id=dish.id,
                product_id=int(pid),
                quantity=float(quantities[idx]),
                unit=units[idx],
                created_date=datetime.now(),
                is_active=True
            )
            db.session.add(di)
        db.session.commit()
        log_activity('edit', 'dish', dish.id, f'C·∫≠p nh·∫≠t m√≥n ƒÉn: {dish.name}')
        flash('ƒê√£ c·∫≠p nh·∫≠t m√≥n ƒÉn!', 'success')
        return redirect(url_for('main.dish_list'))
    # Chu·∫©n b·ªã d·ªØ li·ªáu nguy√™n li·ªáu cho form
    ingredients = dish.ingredients
    return render_template('edit_dish.html', dish=dish, products=products, ingredients=ingredients, product_units=product_units)

@main.route('/dish/<int:dish_id>/delete', methods=['POST'])
def delete_dish(dish_id):
    from app.models import Dish, db
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    dish = Dish.query.get_or_404(dish_id)
    dish_name = dish.name
    
    try:
        # X√≥a m√≥n ƒÉn (cascade s·∫Ω t·ª± ƒë·ªông x√≥a DishIngredient)
        db.session.delete(dish)
        db.session.commit()
        log_activity('delete', 'dish', dish_id, f'X√≥a m√≥n ƒÉn: {dish_name}')
        flash(f'ƒê√£ x√≥a m√≥n ƒÉn "{dish_name}"! L∆∞u √Ω: M√≥n n√†y c√≥ th·ªÉ v·∫´n c√≤n trong th·ª±c ƒë∆°n c≈©, vui l√≤ng ki·ªÉm tra v√† c·∫≠p nh·∫≠t l·∫°i th·ª±c ƒë∆°n.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Kh√¥ng th·ªÉ x√≥a m√≥n ƒÉn: {str(e)}', 'danger')
        
    return redirect(url_for('main.dish_list'))
# ================== T·∫†O M√ìN ƒÇN ==================
@main.route('/dish/new', methods=['GET', 'POST'])
def create_dish():
    from app.models import Dish, DishIngredient, Product, db
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    products = Product.query.filter_by(is_active=True).join(Supplier).order_by(Product.category, Product.name).all()
    product_units = sorted(list(set([p.unit for p in products if p.unit])))
    if request.method == 'POST':
        name = request.form.get('name')
        description = request.form.get('description')
        ingredient_ids = request.form.getlist('ingredient_id')
        units = request.form.getlist('unit')
        quantities = request.form.getlist('quantity')
        from sqlalchemy.exc import IntegrityError
        meal_times = request.form.getlist('meal_times')
        dish = Dish(name=name, description=description, meal_times=meal_times)
        db.session.add(dish)
        try:
            db.session.flush()  # ƒê·ªÉ l·∫•y dish.id
        except IntegrityError:
            db.session.rollback()
            flash('T√™n m√≥n ƒÉn ƒë√£ t·ªìn t·∫°i, vui l√≤ng ch·ªçn t√™n kh√°c!', 'danger')
            return render_template('create_dish.html', products=products, product_units=product_units)
        # Th√™m nguy√™n li·ªáu
        for idx, pid in enumerate(ingredient_ids):
            if not pid or not units[idx] or not quantities[idx]:
                continue
            di = DishIngredient(
                dish_id=dish.id,
                product_id=int(pid),
                quantity=float(quantities[idx]),
                unit=units[idx],
                created_date=datetime.now(),
                is_active=True
            )
            db.session.add(di)
        db.session.commit()
        log_activity('create', 'dish', dish.id, f'T·∫°o m√≥n ƒÉn: {name}')
        flash('ƒê√£ t·∫°o m√≥n ƒÉn th√†nh c√¥ng!', 'success')
        return redirect(url_for('main.dish_list'))
    return render_template('create_dish.html', products=products, product_units=product_units)



@main.route('/')
def index():
    mobile = is_mobile()
    return render_template('about.html', title='Home', mobile=mobile)

@main.route('/about')
def about():
    # Kh√¥ng log activity ·ªü trang ch·ªß ƒë·ªÉ tr√°nh l·ªói khi ch∆∞a migrate
    mobile = is_mobile()
    return render_template('about.html', title='About Us', mobile=mobile)

@main.route('/gallery')
def gallery():
    mobile = is_mobile()
    from app.models import ActivityImage, Activity, Child, Class
    role = session.get('role')
    user_id = session.get('user_id')
    images = []
    if role in ['admin', 'teacher']:
        images = ActivityImage.query.order_by(ActivityImage.upload_date.desc()).all()
    elif role == 'parent':
        # L·∫•y class_id c·ªßa con
        child = Child.query.filter_by(id=user_id).first()
        class_id = None
        if child and child.class_name:
            class_obj = Class.query.filter_by(name=child.class_name).first()
            if class_obj:
                class_id = class_obj.id
        # Ch·ªâ l·∫•y ·∫£nh c·ªßa ho·∫°t ƒë·ªông thu·ªôc l·ªõp con ho·∫∑c cho kh√°ch v√£ng lai
        if class_id:
            images = ActivityImage.query.join(Activity).filter(
                (Activity.class_id == class_id) | (Activity.class_id == None)
            ).order_by(ActivityImage.upload_date.desc()).all()
        else:
            # Kh√¥ng x√°c ƒë·ªãnh ƒë∆∞·ª£c l·ªõp, ch·ªâ l·∫•y ·∫£nh cho kh√°ch v√£ng lai
            images = ActivityImage.query.join(Activity).filter(Activity.class_id == None).order_by(ActivityImage.upload_date.desc()).all()
    else:
        # Kh√°ch v√£ng lai ch·ªâ xem ·∫£nh ho·∫°t ƒë·ªông cho kh√°ch v√£ng lai
        images = ActivityImage.query.join(Activity).filter(Activity.class_id == None).order_by(ActivityImage.upload_date.desc()).all()
    return render_template('gallery.html', title='Gallery', mobile=mobile, images=images)

@main.route('/contact')
def contact():
    mobile = is_mobile()
    return render_template('contact.html', title='Contact Us', mobile=mobile)

@main.route('/activities/new', methods=['GET', 'POST'])
def new_activity():
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    classes = Class.query.order_by(Class.name).all()
    class_choices = [(0, 'T·∫•t c·∫£ kh√°ch v√£ng lai')] + [(c.id, c.name) for c in classes]
    form = ActivityCreateForm()
    form.class_id.choices = class_choices

    if request.method == 'POST':

        # Debug file uploads - BOTH from form and request
        files = request.files.getlist('images')

        # Also try direct access to form field
        form_files = form.images.data if hasattr(form.images, 'data') else []

        for key in request.files.keys():
            files_for_key = request.files.getlist(key)
            
            for i, file in enumerate(files_for_key):
                if file and file.filename:
                    print(f"[DEBUG]   File {i}: {file.filename}, size: {getattr(file, 'content_length', 'unknown')}")
                else:
                    print(f"[DEBUG]   File {i}: Empty or no filename")
        
        for i, file in enumerate(files):
            if file and file.filename:
                print(f"[DEBUG] File {i}: {file.filename}, size: {getattr(file, 'content_length', 'unknown')}")
            else:
                print(f"[DEBUG] File {i}: Empty or no filename")
    
    if form.validate_on_submit():
        print(f"[DEBUG] Form validation SUCCESS - proceeding to process")
        title = form.title.data
        content = form.description.data
        date_val = datetime.strptime(form.date.data, '%Y-%m-%d') if form.date.data else date.today()
        background_file = form.background.data
        image_url = ''
        if background_file and background_file.filename:
            allowed_ext = {'.jpg', '.jpeg', '.png', '.gif', '.jfif'}
            ext = os.path.splitext(background_file.filename)[1].lower()
            safe_filename = re.sub(r'[^a-zA-Z0-9_.-]', '', background_file.filename)
            if ext not in allowed_ext:
                flash('Ch·ªâ cho ph√©p t·∫£i l√™n c√°c file ·∫£nh c√≥ ƒëu√¥i: .jpg, .jpeg, .png, .gif, .jfif!', 'danger')
                return render_template('new_activity.html', form=form, title='ƒêƒÉng b√†i vi·∫øt m·ªõi', mobile=is_mobile(), classes=classes)
            filename = 'bg_' + datetime.now().strftime('%Y%m%d%H%M%S') + '_' + safe_filename
            save_path = os.path.join('app', 'static', 'images', filename)
            # Resize background
            img = Image.open(background_file)
            img.thumbnail((1200, 800))
            img.save(save_path)
            image_url = url_for('static', filename=f'images/{filename}')
        class_id = form.class_id.data if form.class_id.data != 0 else None
        new_post = Activity(title=title, description=content, date=date_val, image=image_url, class_id=class_id)
        db.session.add(new_post)
        db.session.commit()
        
        log_activity('create', 'activity', new_post.id, f'T·∫°o ho·∫°t ƒë·ªông: {title}')
        
        # L∆∞u activity_id v√†o session ƒë·ªÉ upload batch sau
        session['temp_activity_id'] = new_post.id
        
        # T·∫°o th∆∞ m·ª•c l∆∞u ·∫£nh ho·∫°t ƒë·ªông
        activity_dir = os.path.join('app', 'static', 'images', 'activities', str(new_post.id))
        os.makedirs(activity_dir, exist_ok=True)
        
        # X·ª≠ l√Ω ·∫£nh upload v·ªõi t·ªëi ∆∞u h√≥a - LU√îN CH·∫§P NH·∫¨N V√Ä T·ª∞ ƒê·ªòNG S·ª¨A
        files = request.files.getlist('images')
        print(f"[DEBUG] Processing {len(files)} files from request.files")
        
        if files and files[0].filename:  # C√≥ ·∫£nh upload
            print(f"[INFO] X·ª≠ l√Ω upload v·ªõi auto-fix: {len(files)} ·∫£nh")
            total_files = len(files)
            success_count = 0
            warning_messages = []
            
            # L·ªçc file th·ª±c s·ª± (b·ªè qua file tr·ªëng)
            valid_files = []
            for file in files:
                if not file or not file.filename:
                    continue
                    
                # Validate - b√¢y gi·ªù lu√¥n ch·∫•p nh·∫≠n
                is_valid, warning_msg = validate_image_file(file, max_size_mb=50)
                
                if is_valid:
                    valid_files.append(file)
                    if warning_msg:
                        warning_messages.append(warning_msg)
                else:
                    # Ch·ªâ skip file kh√¥ng ph·∫£i ·∫£nh, kh√¥ng b√°o l·ªói
                    print(f"[INFO] B·ªè qua file kh√¥ng ph·∫£i ·∫£nh: {file.filename}")
            
            # X·ª≠ l√Ω t·ª´ng ·∫£nh h·ª£p l·ªá - LU√îN TH√ÄNH C√îNG
            for i, file in enumerate(valid_files):
                try:
                    print(f"[DEBUG] Auto-processing file {i+1}/{len(valid_files)}: {file.filename}")
                    
                    # Ki·ªÉm tra v√† s·ª≠a ·∫£nh b·ªã l·ªói tr∆∞·ªõc
                    is_readable, processed_stream = verify_and_repair_image(file.stream)
                    if not is_readable:
                        print(f"[WARNING] ·∫¢nh {file.filename} kh√¥ng ƒë·ªçc ƒë∆∞·ª£c, b·ªè qua")
                        continue
                    
                    # T·ªëi ∆∞u ·∫£nh - lu√¥n th√†nh c√¥ng v·ªõi fallback
                    optimized_data, img_format = optimize_image(processed_stream, max_size=(1200, 900), quality=80)
                    
                    # T·∫°o t√™n file an to√†n
                    safe_filename = re.sub(r'[^a-zA-Z0-9_.-]', '', file.filename)
                    base_name = os.path.splitext(safe_filename)[0] if safe_filename else 'image'
                    img_filename = f"{datetime.now().strftime('%Y%m%d%H%M%S%f')}_{base_name}.jpg"
                    img_path = os.path.join(activity_dir, img_filename)
                    
                    # ƒê·ªçc data t·ª´ optimized_data tr∆∞·ªõc (lu√¥n c·∫ßn cho fallback)
                    optimized_data.seek(0)
                    image_data = optimized_data.read()
                    
                    # Upload l√™n R2 (n·∫øu c√≥)
                    r2_url = None
                    if R2_ENABLED:
                        try:
                            r2 = get_r2_storage()
                            if r2.enabled:
                                # T·∫°o BytesIO m·ªõi t·ª´ image_data cho R2
                                from io import BytesIO
                                r2_stream = BytesIO(image_data)
                                r2_url = r2.upload_file(r2_stream, img_filename, folder='activities')
                                if r2_url:
                                    print(f"‚úÖ ƒê√£ upload l√™n R2: {img_filename}")
                        except Exception as e:
                            print(f"‚ö†Ô∏è  L·ªói upload R2: {e}")
                    
                    # Fallback: L∆∞u local n·∫øu R2 kh√¥ng th√†nh c√¥ng
                    if not r2_url:
                        with open(img_path, 'wb') as f:
                            f.write(image_data)
                        rel_path = f'images/activities/{new_post.id}/{img_filename}'
                        print(f"üíæ ƒê√£ l∆∞u local: {rel_path}")
                    else:
                        rel_path = r2_url  # D√πng R2 URL
                    
                    # L∆∞u v√†o database
                    db.session.add(ActivityImage(
                        filename=img_filename, 
                        filepath=rel_path, 
                        upload_date=datetime.now(), 
                        activity_id=new_post.id
                    ))
                    
                    success_count += 1
                    print(f"[DEBUG] Successfully auto-processed file {i+1}: {img_filename}")
                    
                except Exception as e:
                    print(f"[WARNING] File {file.filename} kh√¥ng x·ª≠ l√Ω ƒë∆∞·ª£c, b·ªè qua: {e}")
                    # Kh√¥ng b√°o l·ªói, ch·ªâ skip
                    continue
            
            # Commit database
            try:
                db.session.commit()
                print(f"[DEBUG] Database commit successful: {success_count} images saved")
            except Exception as e:
                print(f"[ERROR] L·ªói commit database: {e}")
                db.session.rollback()
                flash('L·ªói l∆∞u v√†o database!', 'danger')
                return render_template('new_activity.html', form=form, title='ƒêƒÉng b√†i vi·∫øt m·ªõi', mobile=is_mobile(), classes=classes)
            
            # Th√¥ng b√°o k·∫øt qu·∫£
            if success_count > 0:
                base_msg = f'ƒê√£ ƒëƒÉng b√†i vi·∫øt m·ªõi v·ªõi {success_count}/{total_files} ·∫£nh th√†nh c√¥ng!'
                if warning_messages:
                    base_msg += f' (ƒê√£ t·ª± ƒë·ªông t·ªëi ∆∞u {len(warning_messages)} ·∫£nh l·ªõn)'
                flash(base_msg, 'success')
            else:
                flash('ƒê√£ ƒëƒÉng b√†i vi·∫øt m·ªõi!', 'success')
        else:
            print(f"[DEBUG] No files to process")
            flash('ƒê√£ t·∫°o b√†i vi·∫øt! H·ªá th·ªëng s·∫Ω x·ª≠ l√Ω ·∫£nh trong gi√¢y l√°t...', 'success')
        return redirect(url_for('main.activities'))
    else:
        print(f"[DEBUG] Form validation FAILED")
        print(f"[DEBUG] Validation errors: {form.errors}")
        # Hi·ªÉn th·ªã form validation errors to user
        for field, errors in form.errors.items():
            for error in errors:
                flash(f'{field}: {error}', 'danger')
    mobile = is_mobile()
    from datetime import date
    current_date_iso = date.today().isoformat()
    return render_template('new_activity.html', form=form, title='ƒêƒÉng b√†i vi·∫øt m·ªõi', mobile=mobile, current_date_iso=current_date_iso, classes=classes)

# Route ri√™ng ƒë·ªÉ upload batch ·∫£nh (t·ª´ client-side compression)
@main.route('/activities/upload-batch', methods=['POST'])
def upload_batch_images():
    if session.get('role') not in ['admin', 'teacher']:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 403
    
    activity_id = request.form.get('activity_id') or session.get('temp_activity_id')
    if not activity_id:
        return jsonify({'success': False, 'error': 'No activity ID'}), 400
    
    files = request.files.getlist('batch_images')
    activity_dir = os.path.join('app', 'static', 'images', 'activities', str(activity_id))
    os.makedirs(activity_dir, exist_ok=True)
    
    success_count = 0
    for file in files:
        if file and file.filename:
            try:
                # File ƒë√£ ƒë∆∞·ª£c n√©n client-side, ch·ªâ c·∫ßn l∆∞u
                safe_filename = re.sub(r'[^a-zA-Z0-9_.-]', '', file.filename)
                img_filename = datetime.now().strftime('%Y%m%d%H%M%S%f') + '_' + safe_filename
                img_path = os.path.join(activity_dir, img_filename)
                file.save(img_path)
                
                rel_path = f'images/activities/{activity_id}/{img_filename}'
                db.session.add(ActivityImage(
                    filename=img_filename,
                    filepath=rel_path,
                    upload_date=datetime.now(),
                    activity_id=activity_id
                ))
                success_count += 1
            except Exception as e:
                print(f"[ERROR] Upload batch error: {e}")
                continue
    
    db.session.commit()
    print(f"[INFO] Uploaded {success_count} images to activity {activity_id}")
    return jsonify({'success': True, 'uploaded': success_count})

# Test route ƒë·ªÉ ki·ªÉm tra upload ·∫£nh
@main.route('/test-activity-images/<int:activity_id>')
def test_activity_images(activity_id):
    activity = Activity.query.get_or_404(activity_id)
    images = ActivityImage.query.filter_by(activity_id=activity_id).all()
    return jsonify({
        'activity_id': activity_id,
        'activity_title': activity.title,
        'image_count': len(images),
        'images': [{'filename': img.filename, 'filepath': img.filepath} for img in images]
    })

@main.route('/activities/<int:id>/delete', methods=['POST'])
def delete_activity(id):
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    post = Activity.query.get_or_404(id)
    if post:
        activity_title = post.title
        
        # Ph√¢n lo·∫°i ·∫£nh R2 v√† local
        r2_images = []
        local_images = []
        
        for img in post.images:
            if img.filepath.startswith('http'):
                r2_images.append(img.filepath)
            else:
                local_images.append(img)
        
        # X√≥a batch tr√™n R2 (nhanh h∆°n)
        if r2_images and R2_ENABLED:
            try:
                r2 = get_r2_storage()
                if r2.enabled:
                    result = r2.delete_files_batch(r2_images)
                    print(f"‚úÖ X√≥a R2: {result['success']} th√†nh c√¥ng, {result['failed']} l·ªói")
            except Exception as e:
                print(f"‚ö†Ô∏è  L·ªói x√≥a batch R2: {e}")
        
        # X√≥a files local
        for img in local_images:
            img_path = os.path.join('app', 'static', img.filepath)
            if os.path.exists(img_path):
                try:
                    os.remove(img_path)
                except:
                    pass
        
        # X√≥a t·∫•t c·∫£ records trong database
        for img in post.images:
            db.session.delete(img)
        
        # X√≥a ho·∫°t ƒë·ªông
        db.session.delete(post)
        db.session.commit()
        log_activity('delete', 'activity', id, f'X√≥a ho·∫°t ƒë·ªông: {activity_title}')
        flash('ƒê√£ xo√° b√†i vi·∫øt v√† t·∫•t c·∫£ ·∫£nh!', 'success')
    else:
        flash('Kh√¥ng t√¨m th·∫•y b√†i vi·∫øt ƒë·ªÉ xo√°!', 'danger')
    mobile = is_mobile()
    return redirect(url_for('main.activities', mobile=mobile))

@main.route('/activities/<int:id>')
def activity_detail(id):
    post = Activity.query.get_or_404(id)
    if not post:
        flash('Kh√¥ng t√¨m th·∫•y b√†i vi·∫øt!', 'danger')
        return redirect(url_for('main.activities'))
    user_role = session.get('role')
    user_id = session.get('user_id')
    if user_role == 'parent':
        child = Child.query.filter_by(id=user_id).first()
        class_name = child.class_name if child else None
        class_obj = Class.query.filter_by(name=class_name).first() if class_name else None
        class_id = class_obj.id if class_obj else None
        # N·∫øu b√†i vi·∫øt kh√¥ng ph·∫£i c·ªßa l·ªõp con m√¨nh v√† kh√¥ng ph·∫£i kh√°ch v√£ng lai th√¨ kh√¥ng cho xem
        if post.class_id is not None and post.class_id != class_id:
            flash('B·∫°n kh√¥ng c√≥ quy·ªÅn xem b√†i vi·∫øt n√†y!', 'danger')
            return redirect(url_for('main.activities'))
    activity = {
        'id': post.id,
        'title': post.title,
        'content': post.description,
        'image': post.image,
        'date_posted': post.date.strftime('%Y-%m-%d'),
        'gallery': post.images
    }
    mobile = is_mobile()
    from app.forms import DeleteActivityForm
    form = DeleteActivityForm()
    return render_template('activity_detail.html', activity=activity, title=post.title, mobile=mobile, form=form)

@main.route('/activities/<int:id>/download')
def download_activity_images(id):
    post = Activity.query.get_or_404(id)
    if not post:
        flash('Kh√¥ng t√¨m th·∫•y b√†i vi·∫øt!', 'danger')
        return redirect(url_for('main.activities'))
    
    # Ki·ªÉm tra quy·ªÅn truy c·∫≠p
    user_role = session.get('role')
    user_id = session.get('user_id')
    if user_role == 'parent':
        child = Child.query.filter_by(id=user_id).first()
        class_name = child.class_name if child else None
        class_obj = Class.query.filter_by(name=class_name).first() if class_name else None
        class_id = class_obj.id if class_obj else None
        if post.class_id is not None and post.class_id != class_id:
            flash('B·∫°n kh√¥ng c√≥ quy·ªÅn t·∫£i b√†i vi·∫øt n√†y!', 'danger')
            return redirect(url_for('main.activities'))
    
    # Ki·ªÉm tra c√≥ h√¨nh ·∫£nh kh√¥ng
    if not post.images or len(post.images) == 0:
        flash('B√†i vi·∫øt n√†y kh√¥ng c√≥ h√¨nh ·∫£nh ƒë·ªÉ t·∫£i!', 'warning')
        return redirect(url_for('main.activity_detail', id=id))
    
    # T·∫°o file ZIP trong b·ªô nh·ªõ
    memory_file = io.BytesIO()
    with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
        for idx, img in enumerate(post.images, 1):
            try:
                # X·ª≠ l√Ω ƒë∆∞·ªùng d·∫´n h√¨nh ·∫£nh
                filepath = img.filepath
                
                # N·∫øu l√† R2 URL, t·∫£i t·ª´ R2
                if filepath.startswith('http'):
                    if R2_ENABLED:
                        try:
                            import requests
                            response = requests.get(filepath, timeout=10)
                            if response.status_code == 200:
                                # L·∫•y extension t·ª´ URL ho·∫∑c d√πng jpg m·∫∑c ƒë·ªãnh
                                ext = os.path.splitext(filepath)[1] or '.jpg'
                                filename = f"{idx:03d}{ext}"
                                zf.writestr(filename, response.content)
                        except Exception as e:
                            print(f"L·ªói t·∫£i ·∫£nh t·ª´ R2: {e}")
                            continue
                    else:
                        continue
                else:
                    # H√¨nh ·∫£nh local
                    # Lo·∫°i b·ªè ti·ªÅn t·ªë 'images/' n·∫øu c√≥
                    if filepath.startswith('images/'):
                        filepath = filepath.replace('images/', '', 1)
                    
                    local_path = os.path.join(current_app.root_path, 'static', 'images', filepath)
                    
                    if os.path.exists(local_path):
                        # L·∫•y t√™n file g·ªëc ho·∫∑c t·∫°o t√™n m·ªõi
                        original_filename = os.path.basename(local_path)
                        ext = os.path.splitext(original_filename)[1]
                        filename = f"{idx:03d}{ext}"
                        
                        with open(local_path, 'rb') as f:
                            zf.writestr(filename, f.read())
            except Exception as e:
                print(f"L·ªói x·ª≠ l√Ω ·∫£nh {idx}: {e}")
                continue
    
    memory_file.seek(0)
    
    # T·∫°o t√™n file ZIP t·ª´ ti√™u ƒë·ªÅ b√†i vi·∫øt
    safe_title = re.sub(r'[^\w\s-]', '', post.title)
    safe_title = re.sub(r'[-\s]+', '_', safe_title)
    zip_filename = f"{safe_title}_{post.date.strftime('%Y%m%d')}.zip"
    
    log_activity('download', 'activity', id, f'T·∫£i h√¨nh ·∫£nh ho·∫°t ƒë·ªông: {post.title}')
    
    return send_file(
        memory_file,
        mimetype='application/zip',
        as_attachment=True,
        download_name=zip_filename
    )

@main.route('/curriculum/new', methods=['GET', 'POST'])
def new_curriculum():
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    classes = Class.query.order_by(Class.name).all()
    if request.method == 'POST':
        week_number = request.form.get('week_number')
        class_id = request.form.get('class_id')
        
        # Check for duplicate curriculum (same week + same class)
        existing = Curriculum.query.filter_by(week_number=week_number, class_id=class_id).first()
        if existing:
            class_name = Class.query.get(class_id).name if class_id else "Ch∆∞a ch·ªçn l·ªõp"
            flash(f'Ch∆∞∆°ng tr√¨nh h·ªçc tu·∫ßn {week_number} cho l·ªõp {class_name} ƒë√£ t·ªìn t·∫°i!', 'danger')
            classes = Class.query.order_by(Class.name).all()
            days = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat']
            morning_slots = ['morning_0', 'morning_1', 'morning_2', 'morning_3', 'morning_4', 'morning_5', 'morning_6']
            afternoon_slots = ['afternoon_1', 'afternoon_2', 'afternoon_3', 'afternoon_4']
            default_data = {}
            for day in days:
                default_data[day] = {}
                for slot in morning_slots + afternoon_slots:
                    default_data[day][slot] = ""
            mobile = is_mobile()
            return render_template('new_curriculum.html', title='T·∫°o ch∆∞∆°ng tr√¨nh m·ªõi', mobile=mobile, classes=classes, data=default_data, error_week=week_number, error_class=class_id)
        
        days = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat']
        morning_slots = ['morning_0', 'morning_1', 'morning_2', 'morning_3', 'morning_4', 'morning_5', 'morning_6']
        afternoon_slots = ['afternoon_1', 'afternoon_2', 'afternoon_3', 'afternoon_4']
        curriculum_data = {}
        for day in days:
            curriculum_data[day] = {}
            for slot in morning_slots:
                curriculum_data[day][slot] = request.form.get(f'{day}_{slot}')
            for slot in afternoon_slots:
                curriculum_data[day][slot] = request.form.get(f'{day}_{slot}')
        content = json.dumps(curriculum_data, ensure_ascii=False)
        new_week = Curriculum(week_number=week_number, class_id=class_id, content=content, material=None)
        db.session.add(new_week)
        db.session.commit()
        log_activity('create', 'curriculum', new_week.id, f'T·∫°o ch∆∞∆°ng tr√¨nh tu·∫ßn {week_number}')
        flash('ƒê√£ th√™m ch∆∞∆°ng tr√¨nh h·ªçc m·ªõi!', 'success')
        return redirect(url_for('main.curriculum'))
    # Set default data for new curriculum form
    days = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat']
    morning_slots = ['morning_0', 'morning_1', 'morning_2', 'morning_3', 'morning_4', 'morning_5', 'morning_6']
    afternoon_slots = ['afternoon_1', 'afternoon_2', 'afternoon_3', 'afternoon_4']
    default_data = {}
    for i, day in enumerate(days):
        default_data[day] = {}
        for slot in morning_slots:
            default_data[day][slot] = ""
        for j, slot in enumerate(afternoon_slots):
            if slot == 'afternoon_2':
                # 15h-15h30
                if i == 1 or i == 3:  # Th·ª© 3, Th·ª© 5
                    default_data[day][slot] = "Ho·∫°t ƒë·ªông v·ªõi gi√°o c·ª•"
                elif i == 2 or i == 4:  # Th·ª© 4, Th·ª© 6
                    default_data[day][slot] = "Lego time"
                else:
                    default_data[day][slot] = ""
            else:
                default_data[day][slot] = ""
    mobile = is_mobile()
    # Compute current week and week start/end (Monday-Sunday) for display
    from datetime import datetime, timedelta, date
    today = datetime.now().date()
    week_start = today - timedelta(days=today.weekday())
    current_week = week_start.isocalendar()[1]
    current_year = week_start.year
    try:
        first_iso_week_ref = date(current_year, 1, 4)
        week_start_iso = first_iso_week_ref + timedelta(days=(current_week - 1) * 7)
        week_start_iso = week_start_iso - timedelta(days=week_start_iso.isoweekday() - 1)
        week_end_iso = week_start_iso + timedelta(days=6)
        current_week_start = week_start_iso.strftime('%d/%m/%Y')
        current_week_end = week_end_iso.strftime('%d/%m/%Y')
    except Exception:
        current_week_start = week_start.strftime('%d/%m/%Y')
        current_week_end = (week_start + timedelta(days=6)).strftime('%d/%m/%Y')

    return render_template('new_curriculum.html', title='T·∫°o ch∆∞∆°ng tr√¨nh m·ªõi', mobile=mobile, classes=classes, data=default_data,
                           current_week=current_week, current_year=current_year,
                           current_week_start=current_week_start, current_week_end=current_week_end)


@main.route('/curriculum')
def curriculum():
    import secrets
    if 'csrf_token' not in session or not session['csrf_token']:
        session['csrf_token'] = secrets.token_hex(16)
    class_id = None
    classes = Class.query.order_by(Class.name).all()
    # N·∫øu l√† ph·ª• huynh, ch·ªâ cho xem curriculum c·ªßa l·ªõp con m√¨nh, kh√¥ng cho override qua URL
    if session.get('role') == 'parent':
        user_id = session.get('user_id')
        child = Child.query.filter_by(id=user_id).first()
        if child and child.class_name:
            class_obj = Class.query.filter_by(name=child.class_name).first()
            if class_obj:
                class_id = class_obj.id
    else:
        # Ch·ªâ admin/teacher m·ªõi ƒë∆∞·ª£c ch·ªçn class_id qua URL
        class_id = request.args.get('class_id', type=int)
    if class_id:
        weeks = Curriculum.query.filter_by(class_id=class_id).order_by(Curriculum.week_number.desc()).all()
    else:
        weeks = Curriculum.query.order_by(Curriculum.week_number.desc()).all()
    curriculum = []
    for week in weeks:
        try:
            data = json.loads(week.content)
        except Exception:
            data = {}
        curriculum.append({
            'week_number': week.week_number,
            'data': data,
            'class_id': week.class_id,
            'class_name': week.class_obj.name if week.class_obj else ''
        })
    mobile = is_mobile()
    return render_template('curriculum.html', curriculum=curriculum, title='Ch∆∞∆°ng tr√¨nh h·ªçc', mobile=mobile, classes=classes, selected_class_id=class_id)

@main.route('/attendance/new', methods=['GET', 'POST'])
def new_student():
    if session.get('role') != 'admin' and session.get('role') != 'teacher':
        return redirect_no_permission()
    classes = Class.query.order_by(Class.name).all()
    # Sinh m√£ s·ªë h·ªçc sinh t·ª± ƒë·ªông: l·∫•y max student_code d·∫°ng s·ªë, +1
    from sqlalchemy import func, cast, Integer
    last_code = db.session.query(func.max(cast(Child.student_code, Integer))).scalar()
    if last_code is None:
        next_code = '001'
    else:
        next_code = str(int(last_code) + 1).zfill(3)
    if request.method == 'POST':
        name = request.form.get('name')
        # student_code l·∫•y t·ª´ hidden input, ƒë√£ sinh s·∫µn
        student_code = request.form.get('student_code') or next_code
        class_name = request.form.get('class_name')
        birth_date = request.form.get('birth_date')
        parent_contact = request.form.get('parent_contact')
        
        # L·∫•y th√¥ng tin ph·ª• huynh chi ti·∫øt
        father_name = request.form.get('father_name')
        father_phone = request.form.get('father_phone')
        mother_name = request.form.get('mother_name')
        mother_phone = request.form.get('mother_phone')
        
        # Validate class first
        if not any(c.name == class_name for c in classes):
            flash('L·ªõp kh√¥ng h·ª£p l·ªá!', 'danger')
            return redirect(url_for('main.new_student'))
            
        # Process avatar separately - Save locally on VPS
        avatar_path = None
        avatar_file = request.files.get('avatar')
        
        if avatar_file and avatar_file.filename:
            import os
            from werkzeug.utils import secure_filename
            ext = os.path.splitext(avatar_file.filename)[1].lower()
            if ext not in ['.jpg', '.jpeg', '.png', '.gif']:
                flash('Ch·ªâ cho ph√©p upload ·∫£nh jpg, jpeg, png, gif! H·ªçc sinh s·∫Ω ƒë∆∞·ª£c t·∫°o kh√¥ng c√≥ ·∫£nh ƒë·∫°i di·ªán.', 'warning')
            else:
                try:
                    filename = f"student_{student_code}_{secure_filename(avatar_file.filename)}"
                    save_dir = os.path.join('app', 'static', 'images', 'students')
                    os.makedirs(save_dir, exist_ok=True)
                    local_path = os.path.join(save_dir, filename)
                    avatar_file.save(local_path)
                    avatar_path = local_path.replace('app/static/', '').replace('\\', '/')
                except Exception as e:
                    flash(f'L·ªói khi l∆∞u ·∫£nh ƒë·∫°i di·ªán: {str(e)}. H·ªçc sinh s·∫Ω ƒë∆∞·ª£c t·∫°o kh√¥ng c√≥ ·∫£nh.', 'warning')
                    avatar_path = None
        
        # Create student
        try:
            new_child = Child(
                name=name, 
                age=0, 
                parent_contact=parent_contact,
                father_name=father_name,
                father_phone=father_phone,
                mother_name=mother_name,
                mother_phone=mother_phone,
                class_name=class_name, 
                birth_date=birth_date, 
                student_code=student_code, 
                avatar=avatar_path
            )
            db.session.add(new_child)
            db.session.commit()
            log_activity('create', 'student', new_child.id, f'T·∫°o h·ªçc sinh: {name}')
            if avatar_path:
                flash('ƒê√£ th√™m h·ªçc sinh m·ªõi v·ªõi ·∫£nh ƒë·∫°i di·ªán!', 'success')
            else:
                flash('ƒê√£ th√™m h·ªçc sinh m·ªõi!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'L·ªói khi th√™m h·ªçc sinh: {str(e)}', 'danger')
            
        return redirect(url_for('main.attendance'))
    mobile = is_mobile()
    return render_template('new_attendance.html', title='T·∫°o h·ªçc sinh m·ªõi', mobile=mobile, classes=classes, next_code=next_code)

@main.route('/attendance', methods=['GET', 'POST'])
def attendance():
    if not session.get('role'):
        flash('B·∫°n ph·∫£i ƒëƒÉng nh·∫≠p m·ªõi truy c·∫≠p ƒë∆∞·ª£c trang n√†y!', 'danger')
        return redirect(url_for('main.about'))
    if session.get('role') == 'parent':
        return redirect(url_for('main.attendance_history'))
    from datetime import date
    attendance_date = request.args.get('attendance_date') or date.today().strftime('%Y-%m-%d')
    selected_class = request.args.get('class_name')
    # L·∫•y danh s√°ch l·ªõp t·ª´ b·∫£ng Class
    class_names = [c.name for c in Class.query.order_by(Class.name).all()]
    # L·ªçc h·ªçc sinh theo l·ªõp
    if selected_class:
        students = Child.query.filter_by(class_name=selected_class, is_active=True).order_by(Child.student_code).all()
    else:
        students = Child.query.filter_by(is_active=True).order_by(Child.student_code).all()
    
    # L·∫•y T·∫§T C·∫¢ attendance records cho ng√†y n√†y m·ªôt l·∫ßn (thay v√¨ query t·ª´ng student)
    student_ids = [s.id for s in students]
    records_dict = {}
    if student_ids:
        records = AttendanceRecord.query.filter(
            AttendanceRecord.child_id.in_(student_ids),
            AttendanceRecord.date == attendance_date
        ).all()
        records_dict = {r.child_id: r for r in records}
    
    # G√°n tr·∫°ng th√°i cho t·ª´ng h·ªçc sinh
    for student in students:
        record = records_dict.get(student.id)
        if record:
            student.status = record.status
            student.breakfast = record.breakfast
            student.lunch = record.lunch
            student.snack = record.snack
            student.toilet = record.toilet
            student.toilet_times = record.toilet_times
            student.note = record.note
        else:
            student.status = 'V·∫Øng'
            student.breakfast = ''
            student.lunch = ''
            student.snack = ''
            student.toilet = ''
            student.toilet_times = ''
            student.note = ''
    if request.method == 'POST':
        for student in students:
            present_value = request.form.get(f'present_{student.id}')
            if present_value == 'yes':
                status = 'C√≥ m·∫∑t'
            elif present_value == 'absent_excused':
                status = 'V·∫Øng m·∫∑t c√≥ ph√©p'
            elif present_value == 'absent_unexcused':
                status = 'V·∫Øng m·∫∑t kh√¥ng ph√©p'
            else:
                status = 'V·∫Øng'
            breakfast = request.form.get(f'breakfast_{student.id}')
            lunch = request.form.get(f'lunch_{student.id}')
            snack = request.form.get(f'snack_{student.id}')
            toilet = request.form.get(f'toilet_{student.id}')
            toilet_times = request.form.get(f'toilet_times_{student.id}') or None
            note = request.form.get(f'note_{student.id}')
            record = AttendanceRecord.query.filter_by(child_id=student.id, date=attendance_date).first()
            if record:
                record.status = status
                record.breakfast = breakfast
                record.lunch = lunch
                record.snack = snack
                record.toilet = toilet
                record.toilet_times = toilet_times
                record.note = note
            else:
                record = AttendanceRecord(child_id=student.id, date=attendance_date, status=status,
                                         breakfast=breakfast, lunch=lunch, snack=snack, toilet=toilet, toilet_times=toilet_times, note=note)
                db.session.add(record)
            student.status = status
            student.breakfast = breakfast
            student.lunch = lunch
            student.snack = snack
            student.toilet = toilet
            student.toilet_times = toilet_times
            student.note = note
        db.session.commit()
        flash('ƒê√£ l∆∞u ƒëi·ªÉm danh!', 'success')
        if not selected_class or selected_class == 'None':
            return redirect(url_for('main.attendance', attendance_date=attendance_date))
        return redirect(url_for('main.attendance', attendance_date=attendance_date, class_name=selected_class))
    mobile = is_mobile()
    return render_template('attendance.html', students=students, title='ƒêi·ªÉm danh', current_date=attendance_date, mobile=mobile, class_names=class_names, selected_class=selected_class)

@main.route('/attendance/mark', methods=['GET', 'POST'])
def mark_attendance():
    students = Child.query.filter_by(is_active=True).all()
    if request.method == 'POST':
        for student in students:
            present = request.form.get(f'present_{student.id}') == 'on'
            # TODO: L∆∞u tr·∫°ng th√°i ƒëi·ªÉm danh v√†o database (c·∫ßn th√™m tr∆∞·ªùng status v√†o model Child)
            student.status = 'C√≥ m·∫∑t' if present else 'V·∫Øng'
        db.session.commit()
        flash('ƒê√£ ƒëi·ªÉm danh cho t·∫•t c·∫£ h·ªçc sinh!', 'success')
        return redirect(url_for('main.attendance'))
    mobile = is_mobile()
    return render_template('mark_attendance.html', students=students, title='ƒêi·ªÉm danh h·ªçc sinh', mobile=mobile)

@main.route('/attendance/history')
def attendance_history():
    if session.get('role') == 'parent':
        # Ch·ªâ cho ph·ª• huynh xem l·ªãch s·ª≠ ƒëi·ªÉm danh c·ªßa con m√¨nh
        user_id = session.get('user_id')
        child = Child.query.filter_by(id=user_id).first()
        students = [child] if child else []
    else:
        students = Child.query.filter_by(is_active=True).all()
    month = request.args.get('month')
    if month:
        year, m = map(int, month.split('-'))
    else:
        today = datetime.today()
        year, m = today.year, today.month
        month = f"{year:04d}-{m:02d}"
    num_days = monthrange(year, m)[1]
    days_in_month = [f"{year:04d}-{m:02d}-{day:02d}" for day in range(1, num_days+1)]
    records_raw = AttendanceRecord.query.filter(AttendanceRecord.date.like(f"{year:04d}-{m:02d}-%")).all()
    records = {}
    for r in records_raw:
        records[(r.child_id, r.date)] = r
    mobile = is_mobile()
    return render_template('attendance_history.html', records=records, students=students, days_in_month=days_in_month, selected_month=month, title='L·ªãch s·ª≠ ƒëi·ªÉm danh', mobile=mobile)

@main.route('/api/save_monthly_service', methods=['POST'])
def save_monthly_service():
    """API ƒë·ªÉ l∆∞u th√¥ng tin d·ªãch v·ª• h√†ng th√°ng khi checkbox thay ƒë·ªïi"""
    try:
        data = request.get_json()
        child_id = data.get('child_id')
        month = data.get('month')
        has_english = data.get('has_english', True)
        has_steamax = data.get('has_steamax', True)
        
        print(f"[DEBUG] Nh·∫≠n request save service: child_id={child_id}, month={month}, english={has_english}, steamax={has_steamax}")
        
        if not child_id or not month:
            return jsonify({'error': 'Missing child_id or month'}), 400
        
        # T√¨m ho·∫∑c t·∫°o record
        service = MonthlyService.query.filter_by(child_id=child_id, month=month).first()
        if not service:
            service = MonthlyService(child_id=child_id, month=month)
            db.session.add(service)
            print(f"[DEBUG] T·∫°o m·ªõi MonthlyService cho child_id={child_id}, month={month}")
        else:
            print(f"[DEBUG] C·∫≠p nh·∫≠t MonthlyService existing cho child_id={child_id}, month={month}")
        
        # C·∫≠p nh·∫≠t th√¥ng tin
        service.has_english = has_english
        service.has_steamax = has_steamax
        
        db.session.commit()
        print(f"[DEBUG] ƒê√£ commit th√†nh c√¥ng!")
        return jsonify({'success': True, 'message': 'ƒê√£ l∆∞u th√¥ng tin d·ªãch v·ª•'})
        
    except Exception as e:
        db.session.rollback()
        print(f"[ERROR] L·ªói save service: {e}")
        return jsonify({'error': str(e)}), 500

@main.route('/invoice', methods=['GET', 'POST'])
def invoice():
    month = request.args.get('month')
    if month:
        year, m = map(int, month.split('-'))
    else:
        today = datetime.today()
        year, m = today.year, today.month
        month = f"{year:04d}-{m:02d}"
    
    # T√≠nh th√°ng h·ªçc ph√≠ (th√°ng ti·∫øp theo)
    next_m = m + 1
    next_year = year
    if next_m > 12:
        next_m = 1
        next_year += 1
    next_month = f"{next_year:04d}-{next_m:02d}"
    next_month_num = f"{next_m:02d}"
    
    num_days = monthrange(year, m)[1]
    days_in_month = [f"{year:04d}-{m:02d}-{day:02d}" for day in range(1, num_days+1)]
    students = Child.query.filter_by(is_active=True).all()
    records_raw = AttendanceRecord.query.filter(AttendanceRecord.date.like(f"{year:04d}-{m:02d}-%")).all()
    # T√≠nh s·ªë ng√†y c√≥ m·∫∑t, s·ªë ng√†y v·∫Øng m·∫∑t kh√¥ng ph√©p v√† c√≥ ph√©p cho t·ª´ng h·ªçc sinh
    attendance_days = {student.id: 0 for student in students}
    absent_unexcused_days = {student.id: 0 for student in students}
    absent_excused_days = {student.id: 0 for student in students}  # Th√™m s·ªë ng√†y v·∫Øng c√≥ ph√©p
    valid_student_ids = set(attendance_days.keys())
    for r in records_raw:
        if r.child_id not in valid_student_ids:
            continue
        if r.status == 'C√≥ m·∫∑t':
            attendance_days[r.child_id] += 1
        elif r.status == 'V·∫Øng m·∫∑t kh√¥ng ph√©p':
            absent_unexcused_days[r.child_id] += 1
        elif r.status == 'V·∫Øng m·∫∑t c√≥ ph√©p':  # Th√™m logic n√†y
            absent_excused_days[r.child_id] += 1
    
    # Load th√¥ng tin d·ªãch v·ª• t·ª´ database
    monthly_services = MonthlyService.query.filter_by(month=month).all()
    services_dict = {service.child_id: service for service in monthly_services}
    
    # T·∫°o default service cho h·ªçc sinh ch∆∞a c√≥ record
    for student in students:
        if student.id not in services_dict:
            # T·∫°o record m·∫∑c ƒë·ªãnh (m·∫∑c ƒë·ªãnh tick c·∫£ 2 d·ªãch v·ª•)
            new_service = MonthlyService(child_id=student.id, month=month, has_english=True, has_steamax=True)
            db.session.add(new_service)
            services_dict[student.id] = new_service
    
    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"[ERROR] L·ªói t·∫°o default services: {e}")
    
    invoices = []
    if request.method == 'POST':
        selected_ids = request.form.getlist('student_ids')
        
        # X·ª≠ l√Ω checkbox data t·ª´ form khi submit v√† c·∫≠p nh·∫≠t database ngay
        for student in students:
            english_checked = request.form.get(f'english_{student.id}') == '1'
            steamax_checked = request.form.get(f'steamax_{student.id}') == '1'
            
            # C·∫≠p nh·∫≠t database ngay l·∫≠p t·ª©c v·ªõi gi√° tr·ªã t·ª´ form
            service = services_dict.get(student.id)
            if service:
                service.has_english = english_checked
                service.has_steamax = steamax_checked
            else:
                # T·∫°o m·ªõi n·∫øu ch∆∞a c√≥
                new_service = MonthlyService(child_id=student.id, month=month, has_english=english_checked, has_steamax=steamax_checked)
                db.session.add(new_service)
                services_dict[student.id] = new_service
        
        # Commit changes to database
        try:
            db.session.commit()
            print(f"[DEBUG] ƒê√£ c·∫≠p nh·∫≠t t·∫•t c·∫£ d·ªãch v·ª• t·ª´ form cho th√°ng {month}")
        except Exception as e:
            db.session.rollback()
            print(f"[ERROR] L·ªói c·∫≠p nh·∫≠t d·ªãch v·ª• t·ª´ form: {e}")
        
        if request.form.get('export_word'):
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w') as zipf:
                for student in students:
                    if str(student.id) in selected_ids:
                        doc = Document()
                        
                        # C√†i ƒë·∫∑t page size A5 n·∫±m ngang
                        if DOCX_AVAILABLE:
                            try:
                                from docx.shared import Inches
                                from docx.enum.section import WD_SECTION_START
                                
                                section = doc.sections[0]
                                # A5 size: 148mm x 210mm, nh∆∞ng n·∫±m ngang n√™n ƒë·∫£o ng∆∞·ª£c
                                section.page_width = Inches(8.27)  # 210mm = 8.27 inches
                                section.page_height = Inches(5.83)  # 148mm = 5.83 inches
                                section.left_margin = Inches(0.3)
                                section.right_margin = Inches(0.3)
                                section.top_margin = Inches(0.2)
                                section.bottom_margin = Inches(0.2)
                            except ImportError:
                                pass
                        
                        # B·∫£ng header: logo b√™n tr√°i, th√¥ng tin tr∆∞·ªùng ·ªü gi·ªØa
                        header_table = doc.add_table(rows=1, cols=3)  # Thay ƒë·ªïi t·ª´ 2 c·ªôt th√†nh 3 c·ªôt
                        header_table.style = None  # Remove borders for a cleaner look
                        left_cell = header_table.cell(0,0)    # Logo
                        center_cell = header_table.cell(0,1)  # Th√¥ng tin tr∆∞·ªùng
                        right_cell = header_table.cell(0,2)   # Tr·ªëng
                        
                        left_cell.vertical_alignment = 1  # Top
                        center_cell.vertical_alignment = 1  # Top
                        right_cell.vertical_alignment = 1  # Top
                        # Logo on the left - to h∆°n
                        logo_path = os.path.join(os.path.dirname(__file__), 'static', 'images', 'logo.jpg')
                        if os.path.exists(logo_path):
                            run_logo = left_cell.paragraphs[0].add_run()
                            if DOCX_AVAILABLE:
                                try:
                                    from docx.shared import Inches
                                    run_logo.add_picture(logo_path, width=Inches(1.0))  # TƒÉng t·ª´ 0.6 l√™n 1.0
                                except ImportError:
                                    pass
                            left_cell.paragraphs[0].alignment = 0  # Left
                        # School info ·ªü gi·ªØa
                        center_paragraph = center_cell.paragraphs[0]
                        center_paragraph.alignment = 1  # Center
                        
                        school_run1 = center_paragraph.add_run('SMALL TREE\n')
                        school_run1.bold = True
                        school_run1.font.size = Pt(10)  # TƒÉng size v√¨ ·ªü gi·ªØa
                        
                        school_run2 = center_paragraph.add_run('M·∫¶M NON C√ÇY NH·ªé\n')
                        school_run2.bold = True
                        school_run2.font.size = Pt(10)
    
                        school_run3 = center_paragraph.add_run('S·ªë 1, Rchai 2, ƒê·ª©c Tr·ªçng, L√¢m ƒê·ªìng\n')
                        school_run3.font.size = Pt(8)
                        
                        school_run4 = center_paragraph.add_run('SDT: 0917618868 / STK: Nguy·ªÖn Th·ªã V√¢n 108875858567 NH VietinBank')
                        school_run4.font.size = Pt(7)
                        
                        # ƒê·∫£m b·∫£o m·ªçi paragraph trong center cell ƒë·ªÅu cƒÉn gi·ªØa
                        for para in center_cell.paragraphs:
                            para.alignment = 1
                        # Lo·∫°i b·ªè paragraph tr·ªëng ƒë·ªÉ ti·∫øt ki·ªám kh√¥ng gian
                        # Format title with proper month and year display
                        # T√≠nh th√°ng h·ªçc ph√≠ (th√°ng ti·∫øp theo)
                        current_year, current_month = map(int, month.split('-'))
                        fee_month = current_month + 1
                        fee_year = current_year
                        if fee_month > 12:
                            fee_month = 1
                            fee_year += 1
                        title = doc.add_heading(f'TH√îNG B√ÅO H·ªåC PH√ç TH√ÅNG {fee_month:02d} NƒÇM {fee_year}', 0)
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = title.runs[0]
                        run.font.size = Pt(12)  # Gi·∫£m t·ª´ 14 xu·ªëng 12 cho A5
                        run.font.color.rgb = RGBColor(76, 175, 80)
                        run.font.name = 'Comic Sans MS'
                        
                        # Thi·∫øt l·∫≠p line spacing compact
                        from docx.shared import Pt as PtUnit
                        title.paragraph_format.space_before = PtUnit(0)
                        title.paragraph_format.space_after = PtUnit(6)
                        
                        # B·∫£ng th√¥ng tin h·ªçc sinh - Layout ngang cho A5
                        info_table = doc.add_table(rows=1, cols=4)  # ƒê·ªïi t·ª´ 2x2 th√†nh 1x4
                        info_table.style = 'Table Grid'
                        for row in info_table.rows:
                            for cell in row.cells:
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:fill'), 'e8f5e9')
                                tcPr.append(shd)
                        info_table.cell(0,0).text = 'H·ªç v√† t√™n:'
                        info_table.cell(0,1).text = student.name
                        info_table.cell(0,2).text = 'Ng√†y sinh:'
                        info_table.cell(0,3).text = student.birth_date or "-"
                        # Lo·∫°i b·ªè paragraph tr·ªëng ƒë·ªÉ ti·∫øt ki·ªám kh√¥ng gian
                        # B·∫£ng t·ªïng k·∫øt
                        days = attendance_days.get(student.id, 0)
                        absents = absent_unexcused_days.get(student.id, 0)
                        age = calculate_age(student.birth_date) if student.birth_date else 0
                        if age == 1:
                            tuition = 1850000
                        elif age == 2:
                            tuition = 1750000
                        elif age == 3:
                            tuition = 1650000
                        elif age == 4:
                            tuition = 1550000
                        else:
                            tuition = 1500000
                        excused_absents = sum(1 for r in records_raw if r.child_id == student.id and r.status == 'V·∫Øng m·∫∑t c√≥ ph√©p')
                        
                        # T√≠nh meal_cost theo c√¥ng th·ª©c m·ªõi: (26 - s·ªë ng√†y v·∫Øng c√≥ ph√©p) * 38000
                        meal_cost = (26 - excused_absents) * 38000
                        
                        # L·∫•y th√¥ng tin d·ªãch v·ª• t·ª´ database sau khi ƒë√£ c·∫≠p nh·∫≠t
                        service = services_dict.get(student.id)
                        has_english = service.has_english if service else True
                        has_steamax = service.has_steamax if service else True
                        
                        # B·∫£ng t√≥m t·∫Øt compact cho A5 - chia l√†m 2 c·ªôt
                        summary_table = doc.add_table(rows=4, cols=4)  # 4x4 grid cho compact
                        summary_table.style = 'Table Grid'
                        for row in summary_table.rows:
                            for cell in row.cells:
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:fill'), 'e8f5e9')
                                tcPr.append(shd)
                                # Set font size cho t·∫•t c·∫£ text trong cell
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(8)  # Gi·∫£m t·ª´ 9 xu·ªëng 8
                        
                        # ƒêi·ªÅn th√¥ng tin c∆° b·∫£n - c·ªôt tr√°i
                        cell = summary_table.cell(0,0)
                        cell.text = 'S·ªë ng√†y ƒëi h·ªçc:'
                        cell.paragraphs[0].runs[0].font.size = Pt(8)
                        
                        cell = summary_table.cell(0,1)
                        cell.text = str(days)
                        cell.paragraphs[0].runs[0].font.size = Pt(8)
                        
                        cell = summary_table.cell(1,0)
                        cell.text = 'S·ªë ng√†y v·∫Øng kh√¥ng ph√©p:'
                        cell.paragraphs[0].runs[0].font.size = Pt(8)
                        
                        cell = summary_table.cell(1,1)
                        cell.text = str(absents)
                        cell.paragraphs[0].runs[0].font.size = Pt(8)
                        
                        cell = summary_table.cell(2,0)
                        cell.text = 'S·ªë ng√†y v·∫Øng c√≥ ph√©p:'
                        cell.paragraphs[0].runs[0].font.size = Pt(8)
                        
                        cell = summary_table.cell(2,1)
                        cell.text = str(excused_absents)
                        cell.paragraphs[0].runs[0].font.size = Pt(8)
                        
                        cell = summary_table.cell(3,0)
                        cell.text = 'Ti·ªÅn ƒÉn:'
                        cell.paragraphs[0].runs[0].font.size = Pt(8)
                        
                        cell = summary_table.cell(3,1)
                        cell.text = f'{meal_cost:,} ƒë'
                        cell.paragraphs[0].runs[0].font.size = Pt(8)
                        
                        # ƒêi·ªÅn th√¥ng tin h·ªçc ph√≠ v√† d·ªãch v·ª• - c·ªôt ph·∫£i
                        cell = summary_table.cell(0,2)
                        cell.text = 'Ti·ªÅn h·ªçc ph√≠:'
                        cell.paragraphs[0].runs[0].font.size = Pt(8)
                        
                        cell = summary_table.cell(0,3)
                        cell.text = f'{tuition:,} ƒë'
                        cell.paragraphs[0].runs[0].font.size = Pt(8)
                        
                        row_index = 1
                        if has_english:
                            cell = summary_table.cell(row_index,2)
                            cell.text = 'Ti·ªÅn h·ªçc anh vƒÉn:'
                            cell.paragraphs[0].runs[0].font.size = Pt(8)
                            
                            cell = summary_table.cell(row_index,3)
                            cell.text = '250,000 ƒë'
                            cell.paragraphs[0].runs[0].font.size = Pt(8)
                            row_index += 1
                            
                        if has_steamax:
                            cell = summary_table.cell(row_index,2)
                            cell.text = 'Ti·ªÅn h·ªçc STEAMAX:'
                            cell.paragraphs[0].runs[0].font.size = Pt(8)
                            
                            cell = summary_table.cell(row_index,3)
                            cell.text = '200,000 ƒë'
                            cell.paragraphs[0].runs[0].font.size = Pt(8)
                            row_index += 1
                        
                        
                        # T√≠nh t·ªïng
                        english_cost = 250000 if has_english else 0
                        steamax_cost = 200000 if has_steamax else 0
                        total = tuition + meal_cost + english_cost + steamax_cost
                        
                        total_paragraph = doc.add_paragraph(f'T·ªïng ti·ªÅn c·∫ßn thanh to√°n: {total:,} ƒë')
                        total_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        total_run = total_paragraph.runs[0]
                        total_run.font.color.rgb = RGBColor(76, 175, 80)
                        total_run.font.bold = True
                        total_run.font.name = 'Comic Sans MS'
                        total_run.font.size = Pt(10)  # Gi·∫£m font size cho A5
                        
                        # Thi·∫øt l·∫≠p compact spacing cho total
                        total_paragraph.paragraph_format.space_before = PtUnit(3)
                        total_paragraph.paragraph_format.space_after = PtUnit(6)

                        # Add payment info table - Compact cho A5
                        # Lo·∫°i b·ªè kho·∫£ng c√°ch ƒë·ªÉ ti·∫øt ki·ªám kh√¥ng gian
                        payment_table = doc.add_table(rows=1, cols=2)
                        payment_table.style = None  # No border for clean look
                        left_payment_cell = payment_table.cell(0,0)
                        right_payment_cell = payment_table.cell(0,1)
                        left_payment_cell.vertical_alignment = 1  # Top
                        right_payment_cell.vertical_alignment = 1  # Top
                        
                        # Left cell v·ªõi font size nh·ªè - cƒÉn gi·ªØa
                        left_para = left_payment_cell.paragraphs[0]
                        left_para.alignment = 1  # Center
                        left_run1 = left_para.add_run('Ng∆∞·ªùi n·ªôp ti·ªÅn:')
                        left_run1.font.size = Pt(8)
                        left_run1.bold = True
                        left_para2 = left_payment_cell.add_paragraph('(K√≠ v√† ghi r√µ h·ªç t√™n)')
                        left_para2.alignment = 1  # Center
                        left_para2.runs[0].font.size = Pt(7)
                        
                        # Right cell v·ªõi font size nh·ªè                      
                        now = datetime.now()
                        # Extract month number and year from month string (format: "YYYY-MM")
                        month_year, month_num = month.split('-')
                        right_para1 = right_payment_cell.paragraphs[0]
                        right_para1.alignment = 1
                        right_run1 = right_para1.add_run(f'Ng√†y 1 th√°ng {month_num} nƒÉm {month_year}')
                        right_run1.font.size = Pt(7)

                        right_para2 = right_payment_cell.add_paragraph('Ch·ªß Tr∆∞·ªùng')
                        right_para2.alignment = 1
                        right_para2.runs[0].font.size = Pt(8)
                        right_para2.runs[0].bold = True
                        
                        right_para3 = right_payment_cell.add_paragraph('(K√≠ v√† ghi r√µ h·ªç t√™n)')
                        right_para3.alignment = 1
                        right_para3.runs[0].font.size = Pt(7)
                        
                        # paragraph tr·ªëng ƒë·ªÉ ti·∫øt ki·ªám kh√¥ng gian
                        right_payment_cell.add_paragraph().alignment = 1
                        
                        right_para_name = right_payment_cell.add_paragraph('Nguy·ªÖn Th·ªã V√¢n')
                        right_para_name.alignment = 1
                        right_para_name.runs[0].font.size = Pt(8)
                        right_para_name.runs[0].bold = True
                        
                        file_stream = io.BytesIO()
                        doc.save(file_stream)
                        file_stream.seek(0)
                        filename = f"invoice_{student.name}_{month}.docx"
                        zipf.writestr(filename, file_stream.read())
            zip_buffer.seek(0)
            return send_file(zip_buffer, download_name=f"invoices_{month}.zip", as_attachment=True)
        else:
            for student in students:
                if str(student.id) in selected_ids:
                    days_present = attendance_days.get(student.id, 0)
                    days_absent_unexcused = absent_unexcused_days.get(student.id, 0)
                    days_absent_excused = absent_excused_days.get(student.id, 0)
                    
                    # L·∫•y th√¥ng tin d·ªãch v·ª• t·ª´ database
                    service = services_dict.get(student.id)
                    has_english = service.has_english if service else True
                    has_steamax = service.has_steamax if service else True
                    
                    # H·ªçc ph√≠ theo ƒë·ªô tu·ªïi - s·ª≠ d·ª•ng student_ages ƒë√£ t√≠nh
                    age = student_ages[student.id]
                    if age == 1:
                        tuition = 1850000
                    elif age == 2:
                        tuition = 1750000
                    elif age == 3:
                        tuition = 1650000
                    elif age == 4:
                        tuition = 1550000
                    else:
                        tuition = 1500000
                    
                    # T√≠nh c√°c kho·∫£n ph√≠ theo c√¥ng th·ª©c m·ªõi
                    meal_cost = (26 - days_absent_excused) * 38000  # 26 ng√†y m·∫∑c ƒë·ªãnh tr·ª´ ng√†y v·∫Øng c√≥ ph√©p
                    english_cost = 250000 if has_english else 0
                    steamax_cost = 200000 if has_steamax else 0
                    total = meal_cost + tuition + english_cost + steamax_cost
                    
                    # T·∫°o chu·ªói m√¥ t·∫£ chi ti·∫øt
                    extras = []
                    if has_english: extras.append(f"Anh vƒÉn: {english_cost:,}ƒë")
                    if has_steamax: extras.append(f"STEAMAX: {steamax_cost:,}ƒë")
                    extra_text = " + " + " + ".join(extras) if extras else ""
                    
                    invoices.append(f"H·ªçc sinh {student.name}: C√≥ m·∫∑t {days_present} ng√†y, v·∫Øng kh√¥ng ph√©p {days_absent_unexcused} ng√†y, v·∫Øng c√≥ ph√©p {days_absent_excused} ng√†y. Ti·ªÅn ƒÉn: {meal_cost:,}ƒë + H·ªçc ph√≠: {tuition:,}ƒë{extra_text} = T·ªïng: {total:,}ƒë")
    mobile = is_mobile()
    student_ages = {student.id: calculate_age(student.birth_date) if student.birth_date else 0 for student in students}
    return render_template('invoice.html', students=students, attendance_days=attendance_days, absent_unexcused_days=absent_unexcused_days, absent_excused_days=absent_excused_days, services_dict=services_dict, selected_month=month, next_month=next_month, next_month_num=next_month_num, invoices=invoices, days_in_month=days_in_month, records={ (r.child_id, r.date): r for r in records_raw }, student_ages=student_ages, title='Xu·∫•t h√≥a ƒë∆°n', mobile=mobile)


@main.route('/login', methods=['GET', 'POST'])
def login():
    user_ip = request.remote_addr
    now = datetime.now()
    # Ki·ªÉm tra lockout do nh·∫≠p sai
    if user_ip in lockout_until and now < lockout_until[user_ip]:
        flash(f'T√†i kho·∫£n ho·∫∑c IP n√†y b·ªã kh√≥a ƒëƒÉng nh·∫≠p t·∫°m th·ªùi. Vui l√≤ng th·ª≠ l·∫°i sau!', 'danger')
        return render_template('login.html', title='ƒêƒÉng nh·∫≠p')
    # Ki·ªÉm tra cooldown sau ƒëƒÉng nh·∫≠p th√†nh c√¥ng
    if user_ip in last_login_time and (now - last_login_time[user_ip]).total_seconds() < LOGIN_COOLDOWN_SECONDS:
        wait_time = LOGIN_COOLDOWN_SECONDS - int((now - last_login_time[user_ip]).total_seconds())
        flash(f'B·∫°n v·ª´a ƒëƒÉng nh·∫≠p th√†nh c√¥ng. Vui l√≤ng ch·ªù {wait_time} gi√¢y tr∆∞·ªõc khi ƒëƒÉng nh·∫≠p l·∫°i!', 'warning')
        return render_template('login.html', title='ƒêƒÉng nh·∫≠p')
    if request.method == 'POST':
        email_or_phone = request.form.get('email')
        password = request.form.get('password')
        # Ki·ªÉm tra admin
        admin = Staff.query.filter_by(position='admin').first()
        if admin and (email_or_phone == admin.email or email_or_phone == admin.phone) and check_password_hash(admin.password, password):
            session['user_id'] = admin.id
            session['role'] = 'admin'
            session['name'] = admin.name
            log_activity('login', description=f'Admin {admin.name} ƒëƒÉng nh·∫≠p')
            flash('ƒêƒÉng nh·∫≠p admin th√†nh c√¥ng!', 'success')
            login_attempts[user_ip] = 0
            last_login_time[user_ip] = now
            return redirect(url_for('main.about'))
        user = Child.query.filter(((Child.email==email_or_phone)|(Child.phone==email_or_phone))).first()
        staff = Staff.query.filter(((Staff.email==email_or_phone)|(Staff.phone==email_or_phone))).first()
        if user and check_password_hash(user.password, password):
            session['user_id'] = user.id
            session['role'] = 'parent'
            session['name'] = user.name
            log_activity('login', description=f'Ph·ª• huynh {user.name} ƒëƒÉng nh·∫≠p')
            flash('ƒêƒÉng nh·∫≠p th√†nh c√¥ng!', 'success')
            login_attempts[user_ip] = 0
            last_login_time[user_ip] = now
            return redirect(url_for('main.about'))
        elif staff and check_password_hash(staff.password, password):
            session['user_id'] = staff.id
            session['role'] = 'teacher'
            session['name'] = staff.name
            log_activity('login', description=f'Gi√°o vi√™n {staff.name} ƒëƒÉng nh·∫≠p')
            flash('ƒêƒÉng nh·∫≠p th√†nh c√¥ng!', 'success')
            login_attempts[user_ip] = 0
            last_login_time[user_ip] = now
            return redirect(url_for('main.about'))
        else:
            login_attempts[user_ip] = login_attempts.get(user_ip, 0) + 1
            if login_attempts[user_ip] >= MAX_LOGIN_ATTEMPTS:
                lockout_until[user_ip] = now + timedelta(minutes=LOCKOUT_TIME_MINUTES)
                flash(f'B·∫°n ƒë√£ nh·∫≠p sai qu√° s·ªë l·∫ßn cho ph√©p. ƒêƒÉng nh·∫≠p b·ªã kh√≥a {LOCKOUT_TIME_MINUTES} ph√∫t!', 'danger')
            else:
                flash('Sai th√¥ng tin ƒëƒÉng nh·∫≠p!', 'danger')
            return render_template('login.html', title='ƒêƒÉng nh·∫≠p')
    return render_template('login.html', title='ƒêƒÉng nh·∫≠p')

@main.route('/logout')
def logout():
    log_activity('logout', description=f'{session.get("name", "User")} ƒëƒÉng xu·∫•t')
    session.clear()
    flash('ƒê√£ ƒëƒÉng xu·∫•t!', 'success')
    return redirect(url_for('main.about'))

@main.route('/create-test-account')
def create_test_account():
    """Create test account for debugging purposes"""
    # Check if gv1@gmail.com already exists
    existing_staff = Staff.query.filter_by(email='gv1@gmail.com').first()
    
    if existing_staff:
        return jsonify({
            'status': 'exists',
            'message': f'Account gv1@gmail.com already exists with position: {existing_staff.position}',
            'staff_id': existing_staff.id
        })
    
    # Create new staff account
    hashed_password = generate_password_hash('123456')
    
    new_staff = Staff(
        name='Gi√°o vi√™n 1',
        position='teacher',
        contact_info='gv1@gmail.com',
        email='gv1@gmail.com', 
        phone='0123456789',
        password=hashed_password
    )
    
    db.session.add(new_staff)
    db.session.commit()
    
    return jsonify({
        'status': 'created',
        'message': 'Test account created successfully',
        'email': 'gv1@gmail.com',
        'password': '123456',
        'position': 'teacher',
        'staff_id': new_staff.id
    })

@main.route('/accounts', methods=['GET', 'POST'])
def accounts():
    from app.models import Staff, Child
    admin = Staff.query.filter_by(position='admin').first()
    # N·∫øu ch∆∞a c√≥ admin, cho ph√©p t·∫°o admin l·∫ßn ƒë·∫ßu
    if not admin:
        if request.method == 'POST':
            username = request.form.get('admin_username')
            email = request.form.get('admin_email')
            password = request.form.get('admin_password')
            password_confirm = request.form.get('admin_password_confirm')
            if password != password_confirm:
                flash('M·∫≠t kh·∫©u nh·∫≠p l·∫°i kh√¥ng kh·ªõp!', 'danger')
                return render_template('accounts.html', show_admin_create=True, title='Kh·ªüi t·∫°o Admin')
            if Staff.query.filter_by(name=username).first() or Staff.query.filter_by(email=email).first():
                flash('T√™n ƒëƒÉng nh·∫≠p ho·∫∑c email ƒë√£ t·ªìn t·∫°i!', 'danger')
                return render_template('accounts.html', show_admin_create=True, title='Kh·ªüi t·∫°o Admin')
            hashed_pw = generate_password_hash(password)
            new_admin = Staff(name=username, email=email, password=hashed_pw, position='admin', contact_info=email)
            db.session.add(new_admin)
            db.session.commit()
            flash('T·∫°o t√†i kho·∫£n admin th√†nh c√¥ng! H√£y ƒëƒÉng nh·∫≠p.', 'success')
            return redirect(url_for('main.login'))
        return render_template('accounts.html', show_admin_create=True, title='Kh·ªüi t·∫°o Admin')
    # N·∫øu ƒë√£ c√≥ admin, ch·ªâ cho ph√©p truy c·∫≠p n·∫øu ƒë√£ ƒëƒÉng nh·∫≠p v·ªõi vai tr√≤ admin
    if session.get('role') != 'admin':
        flash('B·∫°n kh√¥ng c√≥ quy·ªÅn truy c·∫≠p trang n√†y!', 'danger')
        return redirect(url_for('main.login'))
    parents = Child.query.filter_by(is_active=True).all()
    teachers = Staff.query.filter(Staff.position != 'admin').all()
    mobile = is_mobile()
    def mask_user(u):
        return {
            'id': u.id,
            'name': u.name,
            'email': u.email,
            'phone': u.phone,
            'student_code': getattr(u, 'student_code', None),
            'class_name': getattr(u, 'class_name', None),
            'parent_contact': getattr(u, 'parent_contact', None),
            'position': getattr(u, 'position', None),
        }
    masked_parents = [mask_user(p) for p in parents]
    masked_teachers = [mask_user(t) for t in teachers]
    return render_template('accounts.html', parents=masked_parents, teachers=masked_teachers, show_admin_create=False, title='Qu·∫£n l√Ω t√†i kho·∫£n', mobile=mobile)

@main.route('/analytics')
def analytics():
    """Dashboard th·ªëng k√™ ho·∫°t ƒë·ªông ng∆∞·ªùi d√πng v·ªõi filter"""
    if session.get('role') != 'admin':
        return redirect_no_permission()
    
    mobile = is_mobile()
    
    try:
        # Get filter parameters
        filter_user_type = request.args.get('user_type', '')
        filter_action = request.args.get('action', '')
        filter_user_name = request.args.get('user_name', '')
        filter_resource = request.args.get('resource', '')
        filter_days = request.args.get('days', '7')
        # Ensure filter_days is integer
        try:
            filter_days = int(filter_days)
        except (ValueError, TypeError):
            filter_days = 7
        page = request.args.get('page', 1, type=int)
        per_page = 50
        
        # Th·ªëng k√™ theo role
        from sqlalchemy import func, desc, or_
        from datetime import datetime, timedelta, timezone
        
        # Build base filtered query (√°p d·ª•ng cho t·∫•t c·∫£ th·ªëng k√™)
        # S·ª≠ d·ª•ng timezone Vi·ªát Nam (UTC+7)
        vietnam_tz = timezone(timedelta(hours=7))
        cutoff_date = datetime.now(vietnam_tz) - timedelta(days=filter_days)
        base_query = UserActivity.query.filter(UserActivity.timestamp >= cutoff_date)
        
        # √Åp d·ª•ng t·∫•t c·∫£ filters
        if filter_user_type:
            base_query = base_query.filter(UserActivity.user_type == filter_user_type)
        if filter_action:
            base_query = base_query.filter(UserActivity.action == filter_action)
        if filter_user_name:
            base_query = base_query.filter(UserActivity.user_name.like(f'%{filter_user_name}%'))
        if filter_resource:
            base_query = base_query.filter(UserActivity.resource_type.like(f'%{filter_resource}%'))
        
        # T·ªïng l∆∞·ª£t truy c·∫≠p theo user_type (ƒë√£ filter)
        stats_by_role = db.session.query(
            UserActivity.user_type,
            func.count(UserActivity.id).label('count')
        ).filter(UserActivity.timestamp >= cutoff_date)
        
        # √Åp d·ª•ng filters v√†o stats
        if filter_user_type:
            stats_by_role = stats_by_role.filter(UserActivity.user_type == filter_user_type)
        if filter_action:
            stats_by_role = stats_by_role.filter(UserActivity.action == filter_action)
        if filter_user_name:
            stats_by_role = stats_by_role.filter(UserActivity.user_name.like(f'%{filter_user_name}%'))
        if filter_resource:
            stats_by_role = stats_by_role.filter(UserActivity.resource_type.like(f'%{filter_resource}%'))
        
        stats_by_role = stats_by_role.group_by(UserActivity.user_type).all()
        
        # Paginated results (d√πng base_query ƒë√£ filter)
        pagination = base_query.order_by(desc(UserActivity.timestamp)).paginate(
            page=page, per_page=per_page, error_out=False
        )
        recent_activities = pagination.items
        
        # User ho·∫°t ƒë·ªông nhi·ªÅu nh·∫•t (ƒë√£ filter)
        top_users_query = db.session.query(
            UserActivity.user_name,
            UserActivity.user_type,
            func.count(UserActivity.id).label('count')
        ).filter(
            UserActivity.timestamp >= cutoff_date,
            UserActivity.user_id.isnot(None)
        )
        
        # √Åp d·ª•ng filters v√†o top users
        if filter_user_type:
            top_users_query = top_users_query.filter(UserActivity.user_type == filter_user_type)
        if filter_action:
            top_users_query = top_users_query.filter(UserActivity.action == filter_action)
        if filter_user_name:
            top_users_query = top_users_query.filter(UserActivity.user_name.like(f'%{filter_user_name}%'))
        if filter_resource:
            top_users_query = top_users_query.filter(UserActivity.resource_type.like(f'%{filter_resource}%'))
        
        top_users = top_users_query.group_by(
            UserActivity.user_name,
            UserActivity.user_type
        ).order_by(desc('count')).limit(10).all()
        
        # S·ªë l∆∞·ª£t truy c·∫≠p kh√°ch v√£ng lai (ƒë√£ filter)
        guest_query = UserActivity.query.filter(
            UserActivity.timestamp >= cutoff_date
        )
        if not filter_user_type or filter_user_type == 'guest':
            if filter_action:
                guest_query = guest_query.filter(UserActivity.action == filter_action)
            if filter_resource:
                guest_query = guest_query.filter(UserActivity.resource_type.like(f'%{filter_resource}%'))
            guest_visits = guest_query.filter(UserActivity.user_type == 'guest').count()
        else:
            guest_visits = 0
        
        # S·ªë ph·ª• huynh ƒëƒÉng nh·∫≠p (ƒë√£ filter)
        parent_query = UserActivity.query.filter(
            UserActivity.timestamp >= cutoff_date
        )
        if not filter_user_type or filter_user_type == 'parent':
            if not filter_action or filter_action == 'login':
                if filter_resource:
                    parent_query = parent_query.filter(UserActivity.resource_type.like(f'%{filter_resource}%'))
                parent_logins = parent_query.filter(
                    UserActivity.user_type == 'parent',
                    UserActivity.action == 'login'
                ).count()
            else:
                parent_logins = 0
        else:
            parent_logins = 0
        
        # Action ph·ªï bi·∫øn nh·∫•t (ƒë√£ filter)
        top_actions_query = db.session.query(
            UserActivity.action,
            func.count(UserActivity.id).label('count')
        ).filter(UserActivity.timestamp >= cutoff_date)
        
        if filter_user_type:
            top_actions_query = top_actions_query.filter(UserActivity.user_type == filter_user_type)
        if filter_action:
            top_actions_query = top_actions_query.filter(UserActivity.action == filter_action)
        if filter_user_name:
            top_actions_query = top_actions_query.filter(UserActivity.user_name.like(f'%{filter_user_name}%'))
        if filter_resource:
            top_actions_query = top_actions_query.filter(UserActivity.resource_type.like(f'%{filter_resource}%'))
        
        top_actions = top_actions_query.group_by(UserActivity.action).order_by(desc('count')).limit(10).all()
        
        # Get distinct values for filters
        all_user_types = db.session.query(UserActivity.user_type).distinct().all()
        all_actions = db.session.query(UserActivity.action).distinct().all()
        all_resources = db.session.query(UserActivity.resource_type).filter(
            UserActivity.resource_type.isnot(None)
        ).distinct().all()
        
        return render_template('analytics.html',
                             stats_by_role=stats_by_role,
                             recent_activities=recent_activities,
                             activities=pagination,
                             top_users=top_users,
                             guest_visits=guest_visits,
                             parent_logins=parent_logins,
                             top_actions=top_actions,
                             all_user_types=[t[0] for t in all_user_types if t[0]],
                             all_actions=[a[0] for a in all_actions if a[0]],
                             all_resources=[r[0] for r in all_resources if r[0]],
                             filter_user_type=filter_user_type,
                             filter_action=filter_action,
                             filter_user_name=filter_user_name,
                             filter_resource=filter_resource,
                             filter_days=filter_days,
                             title='Th·ªëng k√™ ho·∫°t ƒë·ªông',
                             mobile=mobile)
    except Exception as e:
        import traceback
        print(f"[ERROR] Analytics error: {str(e)}")
        print(traceback.format_exc())
        # Kh√¥ng redirect v·ªÅ about v√¨ about c√≥ th·ªÉ c≈©ng b·ªã l·ªói
        # Hi·ªÉn th·ªã trang l·ªói ƒë∆°n gi·∫£n thay v√¨ redirect
        return render_template('base.html', 
                             title='L·ªói Th·ªëng k√™',
                             error_message=f'B·∫£ng UserActivity ch∆∞a t·ªìn t·∫°i. Vui l√≤ng ch·∫°y: flask db upgrade tr√™n server.')

@main.route('/analytics/clear', methods=['POST'])
def clear_activities():
    """X√≥a ho·∫°t ƒë·ªông theo kho·∫£ng th·ªùi gian"""
    if session.get('role') != 'admin':
        flash('Ch·ªâ admin m·ªõi c√≥ quy·ªÅn x√≥a log ho·∫°t ƒë·ªông!', 'danger')
        return redirect(url_for('main.analytics'))
    
    from datetime import datetime
    start_date = request.form.get('start_date')
    end_date = request.form.get('end_date')
    
    if not start_date or not end_date:
        flash('Vui l√≤ng ch·ªçn ng√†y b·∫Øt ƒë·∫ßu v√† ng√†y k·∫øt th√∫c!', 'danger')
        return redirect(url_for('main.analytics'))
    
    try:
        # Parse dates
        start_dt = datetime.strptime(start_date, '%Y-%m-%d')
        end_dt = datetime.strptime(end_date, '%Y-%m-%d')
        # Set end_dt to end of day
        end_dt = end_dt.replace(hour=23, minute=59, second=59)
        
        if start_dt > end_dt:
            flash('Ng√†y b·∫Øt ƒë·∫ßu ph·∫£i tr∆∞·ªõc ng√†y k·∫øt th√∫c!', 'danger')
            return redirect(url_for('main.analytics'))
        
        # Delete activities in range
        count = UserActivity.query.filter(
            UserActivity.timestamp >= start_dt,
            UserActivity.timestamp <= end_dt
        ).delete()
        
        db.session.commit()
        log_activity('delete', 'activity_log', None, f'X√≥a {count} log t·ª´ {start_date} ƒë·∫øn {end_date}')
        flash(f'ƒê√£ x√≥a {count} ho·∫°t ƒë·ªông t·ª´ {start_date} ƒë·∫øn {end_date}!', 'success')
        
    except ValueError as e:
        flash(f'ƒê·ªãnh d·∫°ng ng√†y kh√¥ng h·ª£p l·ªá: {str(e)}', 'danger')
    except Exception as e:
        db.session.rollback()
        flash(f'L·ªói khi x√≥a: {str(e)}', 'danger')
    
    return redirect(url_for('main.analytics'))

@main.route('/curriculum/<int:week_number>/delete', methods=['POST'])
def delete_curriculum(week_number):
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    # Get class_id from form data. If empty string or missing, treat as None (global curriculum)
    class_id_raw = request.form.get('class_id')
    if class_id_raw is None or class_id_raw == '':
        class_id = None
    else:
        try:
            class_id = int(class_id_raw)
        except Exception:
            class_id = None

    # Filter by both week_number AND class_id (including None) to avoid deleting wrong curriculum
    week = Curriculum.query.filter_by(week_number=week_number, class_id=class_id).first()
    if week:
        # Read related data before deleting to avoid DetachedInstanceError (lazy load after detach)
        class_name = week.class_obj.name if week.class_obj else ''
        db.session.delete(week)
        db.session.commit()
        log_activity('delete', 'curriculum', week_number, f'X√≥a ch∆∞∆°ng tr√¨nh tu·∫ßn {week_number} l·ªõp {class_name}')
        flash(f'ƒê√£ xo√° ch∆∞∆°ng tr√¨nh h·ªçc tu·∫ßn {week_number} c·ªßa l·ªõp {class_name}!', 'success')
    else:
        flash('Kh√¥ng t√¨m th·∫•y ch∆∞∆°ng tr√¨nh h·ªçc ƒë·ªÉ xo√°!', 'danger')
    return redirect(url_for('main.curriculum'))

@main.route('/curriculum/<int:week_number>/edit', methods=['GET', 'POST'])
def edit_curriculum(week_number):
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    
    # Get class_id from URL parameter to ensure we edit the right curriculum
    class_id = request.args.get('class_id', type=int)
    if not class_id:
        flash('C·∫ßn ch·ªâ ƒë·ªãnh l·ªõp h·ªçc ƒë·ªÉ ch·ªânh s·ª≠a ch∆∞∆°ng tr√¨nh!', 'danger')
        return redirect(url_for('main.curriculum'))
        
    # Filter by both week_number AND class_id to avoid editing wrong curriculum
    week = Curriculum.query.filter_by(week_number=week_number, class_id=class_id).first()
    classes = Class.query.order_by(Class.name).all()
    if not week:
        flash('Kh√¥ng t√¨m th·∫•y ch∆∞∆°ng tr√¨nh h·ªçc ƒë·ªÉ ch·ªânh s·ª≠a!', 'danger')
        return redirect(url_for('main.curriculum'))
    import json
    if request.method == 'POST':
        days = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat']
        morning_slots = ['morning_0', 'morning_1', 'morning_2', 'morning_3', 'morning_4', 'morning_5', 'morning_6']
        afternoon_slots = ['afternoon_1', 'afternoon_2', 'afternoon_3', 'afternoon_4']
        curriculum_data = {}
        for day in days:
            curriculum_data[day] = {}
            for slot in morning_slots:
                curriculum_data[day][slot] = request.form.get(f'{day}_{slot}')
            for slot in afternoon_slots:
                curriculum_data[day][slot] = request.form.get(f'{day}_{slot}')
        class_id = request.form.get('class_id')
        week.class_id = class_id
        week.content = json.dumps(curriculum_data, ensure_ascii=False)
        db.session.commit()
        log_activity('edit', 'curriculum', week_number, f'C·∫≠p nh·∫≠t ch∆∞∆°ng tr√¨nh tu·∫ßn {week_number}')
        flash(f'ƒê√£ c·∫≠p nh·∫≠t ch∆∞∆°ng tr√¨nh h·ªçc tu·∫ßn {week_number}!', 'success')
        return redirect(url_for('main.curriculum'))
    data = json.loads(week.content)
    mobile = is_mobile()
    return render_template('edit_curriculum.html', week=week, data=data, title=f'Ch·ªânh s·ª≠a ch∆∞∆°ng tr√¨nh tu·∫ßn {week_number}', mobile=mobile, classes=classes)

@main.route('/profile')
def profile():
    user = None
    role = session.get('role')
    user_id = session.get('user_id')
    info = {}
    if role == 'parent':
        user = Child.query.get(user_id)
        if user:
            info = {
                'full_name': user.parent_contact,
                'email': user.email,
                'phone': user.phone,
                'role_display': 'Ph·ª• huynh',
                'student_code': user.student_code,
                'class_name': user.class_name,
                'birth_date': user.birth_date,
                'parent_contact': user.parent_contact,
            }
    elif role == 'teacher':
        user = Staff.query.get(user_id)
        if user:
            info = {
                'full_name': user.name,
                'email': user.email,
                'phone': user.phone,
                'role_display': 'Gi√°o vi√™n',
                'student_code': '',
                'class_name': user.position,
                'birth_date': user.birth_date,
                'parent_contact': '',
            }
    elif role == 'admin':
        user = Staff.query.get(user_id)
        if user:
            info = {
                'full_name': user.name,
                'email': user.email,
                'phone': user.phone,
                'role_display': 'Admin',
                'student_code': '',
                'class_name': user.position,
                'birth_date': user.birth_date if hasattr(user, 'birth_date') else '',
                'parent_contact': '',
            }
    else:
        flash('Kh√¥ng t√¨m th·∫•y th√¥ng tin t√†i kho·∫£n!', 'danger')
        return redirect(url_for('main.about'))
    mobile = is_mobile()
    return render_template('profile.html', user=info, mobile=mobile)

@main.route('/profile/edit', methods=['GET', 'POST'])
def edit_profile():
    role = session.get('role')
    user_id = session.get('user_id')
    if role == 'parent':
        flash('Ph·ª• huynh kh√¥ng c√≥ quy·ªÅn ch·ªânh s·ª≠a th√¥ng tin!', 'danger')
        return redirect(url_for('main.profile'))
    elif role == 'teacher':
        user = Staff.query.get(user_id)
        full_name = user.name
    else:
        flash('Admin kh√¥ng th·ªÉ ch·ªânh s·ª≠a th√¥ng tin!', 'danger')
        return redirect(url_for('main.profile'))
    if not user:
        flash('Kh√¥ng t√¨m th·∫•y th√¥ng tin t√†i kho·∫£n!', 'danger')
        return redirect(url_for('main.profile'))
    form = EditProfileForm(full_name=full_name, email=user.email, phone=user.phone)
    if form.validate_on_submit():
        user.name = form.full_name.data
        user.email = form.email.data
        user.phone = form.phone.data
        if form.password.data:
            if not form.old_password.data:
                flash('Vui l√≤ng nh·∫≠p m·∫≠t kh·∫©u c≈© ƒë·ªÉ ƒë·ªïi m·∫≠t kh·∫©u!', 'danger')
                return render_template('edit_profile.html', form=form)
            if user.password != form.old_password.data:
                flash('M·∫≠t kh·∫©u c≈© kh√¥ng ƒë√∫ng!', 'danger')
                return render_template('edit_profile.html', form=form)
            user.password = form.password.data
        db.session.commit()
        flash('C·∫≠p nh·∫≠t th√¥ng tin th√†nh c√¥ng!', 'success')
        return redirect(url_for('main.profile'))
    mobile = is_mobile()
    return render_template('edit_profile.html', form=form, mobile=mobile)

@main.route('/students')
def student_list():
    show_all = request.args.get('show_all', 'false').lower() == 'true'
    
    if show_all and session.get('role') in ['admin', 'teacher']:
        students = Child.query.all()
    else:
        students = Child.query.filter_by(is_active=True).all()
    
    mobile = is_mobile()
    role = session.get('role')
    user_id = session.get('user_id')
    
    # Filter students based on role
    if role == 'parent':
        # Ph·ª• huynh ch·ªâ xem ƒë∆∞·ª£c th√¥ng tin con m√¨nh
        filtered_students = [s for s in students if s.id == user_id]
    else:
        # Gi√°o vi√™n v√† admin xem ƒë∆∞·ª£c t·∫•t c·∫£
        filtered_students = students
    
    # Create display data with proper masking for sensitive info
    display_students = []
    for s in filtered_students:
        # Mask sensitive data for non-admin roles
        if role == 'admin':
            student_data = s
        else:
            # Create a simple object that can be accessed with dot notation
            class StudentDisplay:
                def __init__(self, student):
                    self.id = student.id
                    self.name = student.name
                    self.student_code = student.student_code if role in ['admin', 'teacher'] else '·∫®n'
                    self.class_name = student.class_name if role in ['admin', 'teacher'] else '·∫®n'
                    self.parent_contact = student.parent_contact if role in ['admin', 'teacher'] else '·∫®n'
                    self.birth_date = student.birth_date
                    self.avatar = student.avatar
                    self.is_active = student.is_active
            
            student_data = StudentDisplay(s)
        
        display_students.append(student_data)
    
    return render_template('student_list.html', students=display_students, title='Danh s√°ch h·ªçc sinh', mobile=mobile, show_all=show_all)

@main.route('/students/<int:student_id>/edit', methods=['GET', 'POST'])
def edit_student(student_id):
    if session.get('role') != 'admin' and session.get('role') != 'teacher':
        return redirect_no_permission()
    student = Child.query.get_or_404(student_id)
    classes = Class.query.order_by(Class.name).all()
    
    if request.method == 'POST':
        class_name = request.form.get('class_name')
        # Validate class against database
        if not any(c.name == class_name for c in classes):
            flash('L·ªõp kh√¥ng h·ª£p l·ªá!', 'danger')
            return redirect(url_for('main.edit_student', student_id=student_id))
        
        # C·∫≠p nh·∫≠t th√¥ng tin h·ªçc sinh tr∆∞·ªõc
        student.name = request.form.get('name')
        student.student_code = request.form.get('student_code')
        student.class_name = class_name
        student.birth_date = request.form.get('birth_date')
        student.parent_contact = request.form.get('parent_contact')
        
        # C·∫≠p nh·∫≠t th√¥ng tin ph·ª• huynh chi ti·∫øt
        student.father_name = request.form.get('father_name')
        student.father_phone = request.form.get('father_phone')
        student.mother_name = request.form.get('mother_name')
        student.mother_phone = request.form.get('mother_phone')
        
        # X·ª≠ l√Ω avatar ri√™ng bi·ªát - Save locally on VPS
        avatar_updated = False
        
        avatar_file = request.files.get('avatar')
        if avatar_file and avatar_file.filename:
            import os
            from werkzeug.utils import secure_filename
            ext = os.path.splitext(avatar_file.filename)[1].lower()
            if ext not in ['.jpg', '.jpeg', '.png', '.gif']:
                flash('Ch·ªâ cho ph√©p upload ·∫£nh jpg, jpeg, png, gif! Th√¥ng tin kh√°c ƒë√£ ƒë∆∞·ª£c l∆∞u.', 'warning')
            else:
                try:
                    # Delete old avatar (local only)
                    if student.avatar and not student.avatar.startswith('http'):
                        old_path = os.path.join('app', 'static', student.avatar)
                        if os.path.exists(old_path):
                            os.remove(old_path)
                    
                    filename = f"student_{student.student_code}_{secure_filename(avatar_file.filename)}"
                    save_dir = os.path.join('app', 'static', 'images', 'students')
                    os.makedirs(save_dir, exist_ok=True)
                    avatar_path = os.path.join(save_dir, filename)
                    avatar_file.save(avatar_path)
                    student.avatar = avatar_path.replace('app/static/', '').replace('\\', '/')
                    
                    avatar_updated = True
                except Exception as e:
                    flash(f'L·ªói khi l∆∞u ·∫£nh ƒë·∫°i di·ªán: {str(e)}. Th√¥ng tin kh√°c ƒë√£ ƒë∆∞·ª£c l∆∞u.', 'warning')
        
        # Commit t·∫•t c·∫£ thay ƒë·ªïi
        try:
            db.session.commit()
            log_activity('edit', 'student', student_id, f'S·ª≠a h·ªçc sinh: {student.name}')
            if avatar_updated:
                flash('ƒê√£ l∆∞u th√¥ng tin v√† ·∫£nh ƒë·∫°i di·ªán th√†nh c√¥ng!', 'success')
                # Redirect back to edit page to see new avatar
                return redirect(url_for('main.edit_student', student_id=student_id))
            else:
                flash('ƒê√£ l∆∞u th√¥ng tin h·ªçc sinh th√†nh c√¥ng!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'L·ªói khi l∆∞u thay ƒë·ªïi: {str(e)}', 'danger')
            
        return redirect(url_for('main.student_list'))
    mobile = is_mobile()
    return render_template('edit_student.html', student=student, classes=classes, title='Ch·ªânh s·ª≠a h·ªçc sinh', mobile=mobile)

@main.route('/students/<int:student_id>/delete', methods=['POST'])
def delete_student(student_id):
    if session.get('role') != 'admin' and session.get('role') != 'teacher':
        return redirect_no_permission()
    
    try:
        student = Child.query.get_or_404(student_id)
        
        # Xo√° to√†n b·ªô album v√† ·∫£nh li√™n quan tr∆∞·ªõc khi xo√° h·ªçc sinh
        for album in student.albums:
            for photo in album.photos:
                db.session.delete(photo)
            db.session.delete(album)

        # Xo√° to√†n b·ªô b·∫£n ghi ƒëi·ªÉm danh li√™n quan
        attendance_records = AttendanceRecord.query.filter_by(child_id=student.id).all()
        for record in attendance_records:
            db.session.delete(record)

        # Xo√° to√†n b·ªô b·∫£n ghi BMI li√™n quan
        bmi_records = BmiRecord.query.filter_by(student_id=student.id).all()
        for record in bmi_records:
            db.session.delete(record)

        # Xo√° to√†n b·ªô b·∫£n ghi ti·∫øn b·ªô h·ªçc t·∫≠p li√™n quan
        progress_records = StudentProgress.query.filter_by(student_id=student.id).all()
        for record in progress_records:
            db.session.delete(record)

        # Xo√° to√†n b·ªô d·ªãch v·ª• theo th√°ng li√™n quan
        monthly_services = MonthlyService.query.filter_by(child_id=student.id).all()
        for record in monthly_services:
            db.session.delete(record)

        # X√≥a h·ªçc sinh
        student_name = student.name
        db.session.delete(student)
        
        db.session.commit()
        log_activity('delete', 'student', student_id, f'X√≥a h·ªçc sinh: {student_name}')
        flash(f'ƒê√£ xo√° h·ªçc sinh {student_name}!', 'success')
        
    except Exception as e:
        db.session.rollback()
        if "404" in str(e):
            flash(f'Kh√¥ng t√¨m th·∫•y h·ªçc sinh v·ªõi ID {student_id}!', 'danger')
        else:
            flash(f'L·ªói khi xo√° h·ªçc sinh: {str(e)}', 'danger')
        print(f"[ERROR] L·ªói xo√° h·ªçc sinh {student_id}: {str(e)}")
    
    return redirect(url_for('main.student_list'))

@main.route('/students/<int:student_id>/toggle', methods=['POST'])
def toggle_student_status(student_id):
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    
    student = Child.query.get_or_404(student_id)
    student.is_active = not student.is_active
    db.session.commit()
    
    if student.is_active:
        flash(f'ƒê√£ hi·ªán h·ªçc sinh {student.name}!', 'success')
    else:
        flash(f'ƒê√£ ·∫©n h·ªçc sinh {student.name}!', 'warning')
    
    return redirect(url_for('main.student_list'))

@main.route('/students/export')
def export_students():
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    
    if not OPENPYXL_AVAILABLE:
        flash('Kh√¥ng th·ªÉ xu·∫•t file Excel. Vui l√≤ng c√†i ƒë·∫∑t openpyxl!', 'danger')
        return redirect(url_for('main.student_list'))
    
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter
        
        # T·∫°o workbook v√† worksheet
        wb = Workbook()
        ws = wb.active
        ws.title = "Danh s√°ch h·ªçc sinh"
        
        # T·∫°o header
        headers = ['STT', 'H·ªç v√† t√™n', 'M√£ h·ªçc sinh', 'Ng√†y sinh', 'C√¢n n·∫∑ng (kg)', 'Chi·ªÅu cao (cm)', 'BMI', 'ƒê√°nh gi√°']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="4CAF50", end_color="4CAF50", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
        
        # L·∫•y danh s√°ch h·ªçc sinh v√† s·∫Øp x·∫øp theo th·ª© t·ª± l·ªõp, sau ƒë√≥ theo t√™n
        students = Child.query.filter_by(is_active=True).all()
        students = sorted(students, key=lambda x: (get_class_order(x.class_name), x.name))
        
        # Th√™m d·ªØ li·ªáu h·ªçc sinh
        for row, student in enumerate(students, 2):
            # L·∫•y th√¥ng tin BMI m·ªõi nh·∫•t
            latest_bmi = BmiRecord.query.filter_by(student_id=student.id).order_by(BmiRecord.date.desc()).first()
            
            weight = latest_bmi.weight if latest_bmi else ''
            height = latest_bmi.height if latest_bmi else ''
            bmi = round(latest_bmi.bmi, 2) if latest_bmi else ''
            
            # T√≠nh ƒë√°nh gi√° BMI theo tu·ªïi (percentile cho tr·∫ª em)
            assessment = ''
            if latest_bmi and latest_bmi.bmi and student.birth_date:
                try:
                    birth_date_obj = datetime.strptime(str(student.birth_date), '%Y-%m-%d') if isinstance(student.birth_date, str) else student.birth_date
                    age_months = (datetime.now() - birth_date_obj).days // 30
                    assessment = assess_child_bmi(latest_bmi.bmi, age_months)
                except Exception as e:
                    print(f"[ERROR] L·ªói t√≠nh age_months cho {student.name}: {e}")
                    assessment = 'Ch∆∞a c√≥ ƒë·ªß th√¥ng tin'
            
            data = [
                row - 1,  # STT
                student.name,
                student.student_code or '',
                student.birth_date or '',
                weight,
                height,
                bmi,
                assessment
            ]
            
            for col, value in enumerate(data, 1):
                cell = ws.cell(row=row, column=col, value=value)
                cell.alignment = Alignment(horizontal="center" if col == 1 else "left", vertical="center")
                cell.border = Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )
        
        # T·ª± ƒë·ªông ƒëi·ªÅu ch·ªânh ƒë·ªô r·ªông c·ªôt
        for col in range(1, len(headers) + 1):
            column_letter = get_column_letter(col)
            max_length = 0
            for row in ws[column_letter]:
                try:
                    if len(str(row.value)) > max_length:
                        max_length = len(str(row.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # T·∫°o file t·∫°m
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
            wb.save(tmp.name)
            tmp_path = tmp.name
        
        # T·∫°o t√™n file v·ªõi ng√†y gi·ªù hi·ªán t·∫°i
        from datetime import datetime
        filename = f"danh_sach_hoc_sinh_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        
        return send_file(
            tmp_path,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        
    except Exception as e:
        flash(f'L·ªói khi xu·∫•t file Excel: {str(e)}', 'danger')
        return redirect(url_for('main.student_list'))

@main.route('/students/export-subsidized')
def export_subsidized_students():
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    
    if not DOCX_AVAILABLE:
        flash('Kh√¥ng th·ªÉ xu·∫•t file Word. Vui l√≤ng c√†i ƒë·∫∑t python-docx!', 'danger')
        return redirect(url_for('main.student_list'))
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        from docx.shared import RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from datetime import datetime
        
        # T·∫°o document m·ªõi
        doc = Document()
        
        # Thi·∫øt l·∫≠p margin
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        # D√≤ng 1: UBND X√É ƒê·ª®C TR·ªåNG
        p1 = doc.add_paragraph('UBND X√É ƒê·ª®C TR·ªåNG')
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.runs[0].font.size = Pt(12)
        p1.runs[0].font.bold = True
        
        # D√≤ng 2: M·∫¶M NON C√ÇY NH·ªé
        p2 = doc.add_paragraph('M·∫¶M NON C√ÇY NH·ªé')
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.runs[0].font.size = Pt(12)
        p2.runs[0].font.bold = True
        
        # Th√™m kho·∫£ng tr·ªëng
        doc.add_paragraph('')
        
        # D√≤ng 3: Ti√™u ƒë·ªÅ ch√≠nh - d√≤ng 1
        p3 = doc.add_paragraph('DANH S√ÅCH TR·∫∫ M·∫¶M NON ƒê∆Ø·ª¢C MI·ªÑN, GI·∫¢M, H·ªñ TR·ª¢ H·ªåC PH√ç')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.runs[0].font.size = Pt(16)
        p3.runs[0].font.bold = True
        
        # D√≤ng 4: Ti√™u ƒë·ªÅ ch√≠nh - d√≤ng 2
        p4 = doc.add_paragraph('NƒÇM H·ªåC 2025-2026')
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.runs[0].font.size = Pt(16)
        p4.runs[0].font.bold = True
        
        # D√≤ng 5: Ngh·ªã ƒë·ªãnh
        p5 = doc.add_paragraph('(Theo ngh·ªã ƒë·ªãnh 238/2025/Nƒê-CP)')
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p5.runs[0].font.size = Pt(11)
        p5.runs[0].font.italic = True
        
        # Th√™m kho·∫£ng tr·ªëng
        doc.add_paragraph('')
        
        # L·∫•y danh s√°ch h·ªçc sinh v√† s·∫Øp x·∫øp theo th·ª© t·ª± l·ªõp, sau ƒë√≥ theo t√™n
        students = Child.query.filter_by(is_active=True).all()
        students = sorted(students, key=lambda x: (get_class_order(x.class_name), x.name))
        
        # T·∫°o table v·ªõi 5 c·ªôt
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        headers = ['STT', 'H·ªç v√† T√™n H·ªçc Sinh', 'Ng√†y th√°ng nƒÉm sinh', 'H·ªçc l·ªõp', 'Ghi ch√∫']
        
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # ƒê·ªãnh d·∫°ng header
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
        
        # Th√™m d·ªØ li·ªáu h·ªçc sinh
        for idx, student in enumerate(students, 1):
            row_cells = table.add_row().cells
            
            # STT
            row_cells[0].text = str(idx)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # H·ªç v√† t√™n
            row_cells[1].text = student.name or ''
            
            # Ng√†y sinh - format dd/mm/yyyy
            birth_date = ''
            if student.birth_date:
                try:
                    # N·∫øu ng√†y sinh ƒë√£ ·ªü format dd/mm/yyyy th√¨ gi·ªØ nguy√™n
                    if '/' in student.birth_date and len(student.birth_date.split('/')) == 3:
                        birth_date = student.birth_date
                    # N·∫øu ·ªü format yyyy-mm-dd th√¨ chuy·ªÉn ƒë·ªïi sang dd/mm/yyyy
                    elif '-' in student.birth_date and len(student.birth_date.split('-')) == 3:
                        date_parts = student.birth_date.split('-')
                        if len(date_parts) == 3:
                            birth_date = f"{date_parts[2]}/{date_parts[1]}/{date_parts[0]}"
                        else:
                            birth_date = student.birth_date
                    else:
                        birth_date = student.birth_date
                except:
                    birth_date = student.birth_date or ''
            row_cells[2].text = birth_date
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # L·ªõp h·ªçc
            row_cells[3].text = student.class_name or ''
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Ghi ch√∫ (ƒë·ªÉ tr·ªëng cho vi·ªác ƒëi·ªÅn tay sau)
            row_cells[4].text = ''
            
            # ƒê·ªãnh d·∫°ng font cho c√°c cell
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        # Thi·∫øt l·∫≠p ƒë·ªô r·ªông c·ªôt
        for i, width in enumerate([Inches(0.5), Inches(2.5), Inches(1.5), Inches(1.2), Inches(1.8)]):
            for row in table.rows:
                row.cells[i].width = width
        
        # Th√™m kho·∫£ng tr·ªëng v√† ch·ªØ k√Ω
        doc.add_paragraph('')
        doc.add_paragraph('')
        
        # Ch·ªØ k√Ω
        signature_p = doc.add_paragraph('Ch·ªß c∆° s·ªü')
        signature_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        signature_p.runs[0].font.size = Pt(12)
        signature_p.runs[0].font.bold = True
        
        # Th√™m 2 d√≤ng tr·ªëng gi·ªØa "Ch·ªß c∆° s·ªü" v√† t√™n
        doc.add_paragraph('')
        
        name_p = doc.add_paragraph('Nguy·ªÖn Th·ªã V√¢n')
        name_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        name_p.runs[0].font.size = Pt(12)
        name_p.runs[0].font.bold = True
        
        # L∆∞u file t·∫°m
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        # T·∫°o t√™n file
        filename = "danh_sach_mien_giam_hoc_phi_2025_2026.docx"
        
        return send_file(
            tmp_path,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        flash(f'L·ªói khi xu·∫•t file Word: {str(e)}', 'danger')
        return redirect(url_for('main.student_list'))

@main.route('/admin/change-password', methods=['GET', 'POST'])
def change_admin_password():
    if session.get('role') != 'admin':
        flash('B·∫°n kh√¥ng c√≥ quy·ªÅn ƒë·ªïi m·∫≠t kh·∫©u admin!', 'danger')
        return redirect(url_for('main.login'))
    admin = Staff.query.filter_by(name='admin').first()
    if not admin:
        flash('Kh√¥ng t√¨m th·∫•y t√†i kho·∫£n admin!', 'danger')
        return redirect(url_for('main.accounts'))
    if request.method == 'POST':
        old_password = request.form.get('old_password')
        new_password = request.form.get('new_password')
        confirm_password = request.form.get('confirm_password')
        if old_password != admin.password:
            flash('M·∫≠t kh·∫©u hi·ªán t·∫°i kh√¥ng ƒë√∫ng!', 'danger')
        elif new_password != confirm_password:
            flash('M·∫≠t kh·∫©u m·ªõi nh·∫≠p l·∫°i kh√¥ng kh·ªõp!', 'danger')
        else:
            admin.password = generate_password_hash(new_password)
            db.session.commit()
            flash('ƒê·ªïi m·∫≠t kh·∫©u admin th√†nh c√¥ng!', 'success')
            return redirect(url_for('main.accounts'))
    return render_template('change_admin_password.html', title='ƒê·ªïi m·∫≠t kh·∫©u Admin')

@main.route('/activities')
def activities():
    user_role = session.get('role')
    user_id = session.get('user_id')
    posts = None
    if user_role == 'parent':
        child = Child.query.filter_by(id=user_id).first()
        class_name = child.class_name if child else None
        class_obj = Class.query.filter_by(name=class_name).first() if class_name else None
        class_id = class_obj.id if class_obj else None
        # Ch·ªâ l·∫•y b√†i vi·∫øt c·ªßa l·ªõp con m√¨nh ho·∫∑c b√†i cho kh√°ch v√£ng lai
        posts = Activity.query.filter(
            (Activity.class_id == class_id) | (Activity.class_id == None)
        ).order_by(Activity.date.desc()).all()
    else:
        # Gi√°o vi√™n, admin xem t·∫•t c·∫£
        posts = Activity.query.order_by(Activity.date.desc()).all()
    activities = [
        {
            'id': post.id,
            'title': post.title,
            'content': post.description,
            'image': post.image,
            'date_posted': post.date.strftime('%Y-%m-%d'),
            'images': post.images  # Truy·ªÅn danh s√°ch ·∫£nh gallery
        } for post in posts
    ]
    mobile = is_mobile()
    from app.forms import DeleteActivityForm
    form = DeleteActivityForm()
    return render_template('activities.html', activities=activities, title='Ho·∫°t ƒë·ªông', mobile=mobile, form=form)

@main.route('/accounts/create', methods=['GET', 'POST'])
def create_account():
    if session.get('role') != 'admin':
        return redirect_no_permission()
    if request.method == 'POST':
        role = request.form.get('role')
        name = request.form.get('name')
        email = request.form.get('email')
        phone = request.form.get('phone')
        password = request.form.get('password')
        if role == 'admin':
            flash('Kh√¥ng th·ªÉ t·∫°o t√†i kho·∫£n admin qua form n√†y!', 'danger')
            return render_template('create_account.html', title='T·∫°o t√†i kho·∫£n m·ªõi')
        # Ki·ªÉm tra tr√πng t√™n/email
        if not name or not email or not phone or not password:
            flash('Vui l√≤ng nh·∫≠p ƒë·∫ßy ƒë·ªß th√¥ng tin b·∫Øt bu·ªôc!', 'danger')
            return render_template('create_account.html', title='T·∫°o t√†i kho·∫£n m·ªõi')
        if (Child.query.filter_by(name=name).first() or Staff.query.filter_by(name=name).first() or name == 'admin'):
            flash('T√™n ƒë√£ t·ªìn t·∫°i ho·∫∑c tr√πng v·ªõi t√†i kho·∫£n kh√°c!', 'danger')
            return render_template('create_account.html', title='T·∫°o t√†i kho·∫£n m·ªõi')
        if (Child.query.filter_by(email=email).first() or Staff.query.filter_by(email=email).first() or email == 'admin@smalltree.vn'):
            flash('Email ƒë√£ t·ªìn t·∫°i ho·∫∑c tr√πng v·ªõi t√†i kho·∫£n kh√°c!', 'danger')
            return render_template('create_account.html', title='T·∫°o t√†i kho·∫£n m·ªõi')
        if role == 'parent':
            student_code = request.form.get('student_code')
            class_name = request.form.get('class_name')
            birth_date = request.form.get('birth_date')
            parent_contact = request.form.get('parent_contact')
            if not student_code or not class_name or not birth_date or not parent_contact:
                flash('Vui l√≤ng nh·∫≠p ƒë·∫ßy ƒë·ªß th√¥ng tin h·ªçc sinh/ph·ª• huynh!', 'danger')
                return render_template('create_account.html', title='T·∫°o t√†i kho·∫£n m·ªõi')
            new_child = Child(name=name, age=0, parent_contact=parent_contact, class_name=class_name, birth_date=birth_date, email=email, phone=phone, password=generate_password_hash(password), student_code=student_code)
            db.session.add(new_child)
        elif role == 'teacher':
            position = request.form.get('position')
            if not position:
                flash('Vui l√≤ng nh·∫≠p ch·ª©c v·ª• gi√°o vi√™n!', 'danger')
                return render_template('create_account.html', title='T·∫°o t√†i kho·∫£n m·ªõi')
            new_staff = Staff(name=name, position=position, contact_info=phone, email=email, phone=phone, password=generate_password_hash(password))
            db.session.add(new_staff)
        db.session.commit()
        new_user_id = new_child.id if role == 'parent' else new_staff.id
        log_activity('create', 'account', new_user_id, f'T·∫°o t√†i kho·∫£n {role}: {name}')
        flash('T·∫°o t√†i kho·∫£n th√†nh c√¥ng!', 'success')
        return redirect(url_for('main.accounts'))
    
    classes = Class.query.order_by(Class.name).all()
    return render_template('create_account.html', classes=classes, title='T·∫°o t√†i kho·∫£n m·ªõi')

@main.route('/accounts/<int:user_id>/edit', methods=['GET', 'POST'])
def edit_account(user_id):
    if session.get('role') != 'admin':
        return redirect_no_permission()
    user_type = request.args.get('type', 'parent')
    classes = Class.query.order_by(Class.name).all() if user_type == 'parent' else []
    if user_type == 'teacher':
        user = Staff.query.get_or_404(user_id)
    else:
        user = Child.query.get_or_404(user_id)
    if request.method == 'POST':
        user.name = request.form.get('name')
        user.email = request.form.get('email')
        user.phone = request.form.get('phone')
        if user_type == 'parent':
            user.parent_contact = request.form.get('parent_contact')
            user.student_code = request.form.get('student_code')
            user.class_name = request.form.get('class_name')
            user.birth_date = request.form.get('birth_date')
        elif user_type == 'teacher':
            user.position = request.form.get('position')
        password = request.form.get('password')
        if password:
            user.password = generate_password_hash(password)
        db.session.commit()
        log_activity('edit', 'account', user_id, f'C·∫≠p nh·∫≠t t√†i kho·∫£n: {user.name}')
        flash('ƒê√£ c·∫≠p nh·∫≠t th√¥ng tin t√†i kho·∫£n!', 'success')
        return redirect(url_for('main.accounts'))
    # Hide sensitive info for non-admins (should only be admin here, but for safety)
    show_sensitive = session.get('role') == 'admin'
    masked_user = {
        'id': user.id,
        'name': user.name,
        'email': user.email if show_sensitive else '·∫®n',
        'phone': user.phone if show_sensitive else '·∫®n',
        'student_code': getattr(user, 'student_code', None) if show_sensitive else '·∫®n',
        'class_name': getattr(user, 'class_name', None) if show_sensitive else '·∫®n',
        'parent_contact': getattr(user, 'parent_contact', None) if show_sensitive else '·∫®n',
        'position': getattr(user, 'position', None) if show_sensitive else '·∫®n',
    }
    return render_template('edit_account.html', user=masked_user, type=user_type, title='Ch·ªânh s·ª≠a t√†i kho·∫£n', classes=classes)

@main.route('/accounts/parent/<int:user_id>/delete', methods=['POST'])
def delete_parent_account(user_id):
    if session.get('role') != 'admin':
        return redirect_no_permission()
    user = Child.query.get_or_404(user_id)
    user_name = user.name
    db.session.delete(user)
    db.session.commit()
    log_activity('delete', 'account', user_id, f'X√≥a t√†i kho·∫£n ph·ª• huynh: {user_name}')
    flash('ƒê√£ xo√° t√†i kho·∫£n ph·ª• huynh!', 'success')
    return redirect(url_for('main.accounts'))

@main.route('/accounts/teacher/<int:user_id>/delete', methods=['POST'])
def delete_teacher_account(user_id):
    if session.get('role') != 'admin':
        return redirect_no_permission()
    user = Staff.query.get_or_404(user_id)
    user_name = user.name
    db.session.delete(user)
    db.session.commit()
    log_activity('delete', 'account', user_id, f'X√≥a t√†i kho·∫£n gi√°o vi√™n: {user_name}')
    flash('ƒê√£ xo√° t√†i kho·∫£n gi√°o vi√™n!', 'success')
    return redirect(url_for('main.accounts'))

@main.route('/activities/<int:id>/edit', methods=['GET', 'POST'])
def edit_activity(id):
    try:
        if session.get('role') not in ['admin', 'teacher']:
            return redirect_no_permission()
        post = Activity.query.get_or_404(id)
        if not post:
            flash('Kh√¥ng t√¨m th·∫•y b√†i vi·∫øt ƒë·ªÉ ch·ªânh s·ª≠a!', 'danger')
            return redirect(url_for('main.activities'))
        from app.forms import ActivityEditForm
        classes = Class.query.order_by(Class.name).all()
        class_choices = [(0, 'T·∫•t c·∫£ kh√°ch v√£ng lai')] + [(c.id, c.name) for c in classes]
        form = ActivityEditForm()
        form.class_id.choices = class_choices
        if request.method == 'POST' and form.validate_on_submit():
            post.title = form.title.data
            post.description = form.description.data
            post.class_id = form.class_id.data if form.class_id.data != 0 else None
            background_file = form.background.data
            image_url = post.image
            if background_file and getattr(background_file, 'filename', None):
                allowed_ext = {'.jpg', '.jpeg', '.png', '.gif', '.jfif'}
                ext = os.path.splitext(background_file.filename)[1].lower()
                safe_filename = re.sub(r'[^a-zA-Z0-9_.-]', '', background_file.filename)
                if ext not in allowed_ext:
                    flash('Ch·ªâ cho ph√©p t·∫£i l√™n c√°c file ·∫£nh c√≥ ƒëu√¥i: .jpg, .jpeg, .png, .gif, .jfif!', 'danger')
                    return render_template('edit_activity.html', post=post, form=form, title='Ch·ªânh s·ª≠a ho·∫°t ƒë·ªông', mobile=is_mobile(), classes=classes)
                filename = 'bg_' + datetime.now().strftime('%Y%m%d%H%M%S') + '_' + safe_filename
                save_path = os.path.join('app', 'static', 'images', filename)
                img = Image.open(background_file)
                img.thumbnail((1200, 800))
                img.save(save_path)
                image_url = url_for('static', filename=f'images/{filename}')
                post.image = image_url
            files = request.files.getlist('images')
            activity_dir = os.path.join('app', 'static', 'images', 'activities', str(post.id))
            os.makedirs(activity_dir, exist_ok=True)
            for file in files:
                if file and getattr(file, 'filename', None):
                    ext = os.path.splitext(file.filename)[1].lower()
                    if ext not in ['.jpg', '.jpeg', '.png', '.gif', '.jfif']:
                        flash(f"File {file.filename} kh√¥ng ƒë√∫ng ƒë·ªãnh d·∫°ng ·∫£nh!", 'danger')
                        continue
                    safe_filename = re.sub(r'[^a-zA-Z0-9_.-]', '', file.filename)
                    img_filename = datetime.now().strftime('%Y%m%d%H%M%S%f') + '_' + safe_filename
                    img_path = os.path.join(activity_dir, img_filename)
                    try:
                        file.stream.seek(0)
                        img = Image.open(file.stream)
                        img.thumbnail((1200, 800))
                        img.save(img_path)
                        rel_path = os.path.join('images', 'activities', str(post.id), img_filename).replace('\\', '/')
                        db.session.add(ActivityImage(filename=img_filename, filepath=rel_path, upload_date=datetime.now(), activity_id=post.id))
                    except Exception as e:
                        print(f"[ERROR] L·ªói upload ·∫£nh: {getattr(file, 'filename', 'unknown')} - {e}")
                        import traceback
                        traceback.print_exc()
                        flash(f"L·ªói upload ·∫£nh: {getattr(file, 'filename', 'unknown')} - {e}", 'danger')
                        continue
            db.session.commit()
            log_activity('edit', 'activity', id, f'C·∫≠p nh·∫≠t ho·∫°t ƒë·ªông: {post.title}')
            flash('ƒê√£ c·∫≠p nh·∫≠t b√†i vi·∫øt!', 'success')
            return redirect(url_for('main.activities'))
        mobile = is_mobile()
        # G√°n d·ªØ li·ªáu m·∫∑c ƒë·ªãnh cho form khi GET
        if request.method == 'GET':
            form.title.data = post.title
            form.description.data = post.description
            form.class_id.data = post.class_id if post.class_id is not None else 0
        return render_template('edit_activity.html', post=post, form=form, title='Ch·ªânh s·ª≠a ho·∫°t ƒë·ªông', mobile=mobile, classes=classes)
    except Exception as e:
        print(f"[ERROR] L·ªói khi render edit_activity: {e}")
        import traceback
        traceback.print_exc()
        flash(f"L·ªói h·ªá th·ªëng khi ch·ªânh s·ª≠a ho·∫°t ƒë·ªông: {e}", 'danger')
        return redirect(url_for('main.activities'))

@main.route('/bmi-index', methods=['GET', 'POST'])
def bmi_index():
    students = Child.query.all()
    bmi = None
    bmi_id = None

    if request.method == 'POST':
        student_id = int(request.form['student_id'])
        weight = float(request.form['weight'])
        height = float(request.form['height']) / 100  # ƒë·ªïi cm sang m
        bmi = round(weight / (height * height), 2)
        bmi_id = student_id
        record_date = request.form.get('date', date.today().isoformat())
        # Fix: round height to 2 decimals before saving
        rounded_height_cm = round(height * 100, 2)
        new_record = BmiRecord(
            student_id=student_id,
            date=date.fromisoformat(record_date),
            weight=weight,
            height=rounded_height_cm,  # l∆∞u l·∫°i ƒë∆°n v·ªã cm, ƒë√£ l√†m tr√≤n
            bmi=bmi
        )
        db.session.add(new_record)
        db.session.commit()

    bmi_history = {}
    for student in students:
        records = BmiRecord.query.filter_by(student_id=student.id).order_by(BmiRecord.date.desc(), BmiRecord.id.desc()).all()
        filtered = {}
        for record in records:
            date_str = record.date.strftime('%Y-%m-%d')
            if date_str not in filtered:
                filtered[date_str] = record
        bmi_history[student.id] = list(filtered.values())

    current_date_iso = date.today().isoformat()
    return render_template(
        'bmi_index.html',
        title='Ch·ªâ S·ªë BMI',
        students=students,
        bmi=bmi,
        bmi_id=bmi_id,
        bmi_history=bmi_history,
        current_date_iso=current_date_iso,
        mobile=False
    )

@main.route('/bmi-record/<int:record_id>/edit', methods=['POST'])
def edit_bmi_record(record_id):
    record = BmiRecord.query.get_or_404(record_id)
    date_str = request.form.get('date')
    weight = request.form.get('weight')
    height = request.form.get('height')
    if date_str and weight and height:
        record.date = date.fromisoformat(date_str)
        record.weight = float(weight)
        record.height = request.form.get('height')
        record.bmi = round(float(weight) / ((float(height)/100) ** 2), 2)
        db.session.commit()
        flash('ƒê√£ c·∫≠p nh·∫≠t ch·ªâ s·ªë BMI!', 'success')
    else:
        flash('Vui l√≤ng nh·∫≠p ƒë·∫ßy ƒë·ªß th√¥ng tin!', 'danger')
    return redirect(url_for('main.bmi_index', edit_id=None))

@main.route('/bmi-record/<int:record_id>/delete', methods=['POST'])
def delete_bmi_record(record_id):
    record = BmiRecord.query.get_or_404(record_id)
    db.session.delete(record)
    db.session.commit()
    flash('ƒê√£ xo√° ch·ªâ s·ªë BMI!', 'success')
    return redirect(url_for('main.bmi_index', edit_id=None))

@main.route('/menu')
def menu():
    # Ch·ªâ s·ª≠ d·ª•ng Menu model cho th·ª±c ƒë∆°n
    menus = Menu.query.order_by(Menu.week_number.desc()).all()
    menu = []
    
    for menu_item in menus:
        # Convert structured Menu fields to template format
        data = {
            'mon': {
                'morning': menu_item.monday_morning or '',
                'snack': menu_item.monday_snack or '',
                'dessert': menu_item.monday_dessert or '',
                'lunch': menu_item.monday_lunch or '',
                'afternoon': menu_item.monday_afternoon or '',
                'lateafternoon': menu_item.monday_lateafternoon or ''
            },
            'tue': {
                'morning': menu_item.tuesday_morning or '',
                'snack': menu_item.tuesday_snack or '',
                'dessert': menu_item.tuesday_dessert or '',
                'lunch': menu_item.tuesday_lunch or '',
                'afternoon': menu_item.tuesday_afternoon or '',
                'lateafternoon': menu_item.tuesday_lateafternoon or ''
            },
            'wed': {
                'morning': menu_item.wednesday_morning or '',
                'snack': menu_item.wednesday_snack or '',
                'dessert': menu_item.wednesday_dessert or '',
                'lunch': menu_item.wednesday_lunch or '',
                'afternoon': menu_item.wednesday_afternoon or '',
                'lateafternoon': menu_item.wednesday_lateafternoon or ''
            },
            'thu': {
                'morning': menu_item.thursday_morning or '',
                'snack': menu_item.thursday_snack or '',
                'dessert': menu_item.thursday_dessert or '',
                'lunch': menu_item.thursday_lunch or '',
                'afternoon': menu_item.thursday_afternoon or '',
                'lateafternoon': menu_item.thursday_lateafternoon or ''
            },
            'fri': {
                'morning': menu_item.friday_morning or '',
                'snack': menu_item.friday_snack or '',
                'dessert': menu_item.friday_dessert or '',
                'lunch': menu_item.friday_lunch or '',
                'afternoon': menu_item.friday_afternoon or '',
                'lateafternoon': menu_item.friday_lateafternoon or ''
            },
            'sat': {
                'morning': menu_item.saturday_morning or '',
                'snack': menu_item.saturday_snack or '',
                'dessert': menu_item.saturday_dessert or '',
                'lunch': menu_item.saturday_lunch or '',
                'afternoon': menu_item.saturday_afternoon or '',
                'lateafternoon': menu_item.saturday_lateafternoon or ''
            }
        }
        
        menu.append({
            'week_number': menu_item.week_number,
            'data': data
        })
    
    mobile = is_mobile()
    return render_template('menu.html', menu=menu, title='Th·ª±c ƒë∆°n', mobile=mobile)

@main.route('/menu/new', methods=['GET', 'POST'])
def new_menu():
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    if request.method == 'POST':
        week_number = int(request.form.get('week_number'))
        
        # Check if menu already exists for this week
        existing_menu = Menu.query.filter_by(week_number=week_number, year=2025).first()
        if existing_menu:
            flash(f'Th·ª±c ƒë∆°n tu·∫ßn {week_number}/2025 ƒë√£ t·ªìn t·∫°i! Vui l√≤ng ch·ªçn tu·∫ßn kh√°c ho·∫∑c s·ª≠a th·ª±c ƒë∆°n hi·ªán c√≥.', 'danger')
            # Get all active dishes for re-rendering the form
            dishes = Dish.query.filter_by(is_active=True).all()
            mobile = is_mobile()
            
            # Calculate current week for display
            from datetime import datetime, timedelta
            today = datetime.now().date()
            week_start = today - timedelta(days=today.weekday())
            current_week = week_start.isocalendar()[1]
            current_year = week_start.year
            
            return render_template('new_menu.html', title='T·∫°o th·ª±c ƒë∆°n m·ªõi', mobile=mobile, dishes=dishes, 
                                 error_week=week_number, current_week=current_week, current_year=current_year)
        
        # Create new Menu record
        new_menu = Menu(
            week_number=week_number,
            year=2025,  # Current year
            # Monday
            monday_morning=request.form.get('content_mon_morning', ''),
            monday_snack=request.form.get('content_mon_snack', ''),
            monday_dessert=request.form.get('content_mon_dessert', ''),
            monday_lunch=request.form.get('content_mon_lunch', ''),
            monday_afternoon=request.form.get('content_mon_afternoon', ''),
            monday_lateafternoon=request.form.get('content_mon_lateafternoon', ''),
            # Tuesday
            tuesday_morning=request.form.get('content_tue_morning', ''),
            tuesday_snack=request.form.get('content_tue_snack', ''),
            tuesday_dessert=request.form.get('content_tue_dessert', ''),
            tuesday_lunch=request.form.get('content_tue_lunch', ''),
            tuesday_afternoon=request.form.get('content_tue_afternoon', ''),
            tuesday_lateafternoon=request.form.get('content_tue_lateafternoon', ''),
            # Wednesday
            wednesday_morning=request.form.get('content_wed_morning', ''),
            wednesday_snack=request.form.get('content_wed_snack', ''),
            wednesday_dessert=request.form.get('content_wed_dessert', ''),
            wednesday_lunch=request.form.get('content_wed_lunch', ''),
            wednesday_afternoon=request.form.get('content_wed_afternoon', ''),
            wednesday_lateafternoon=request.form.get('content_wed_lateafternoon', ''),
            # Thursday
            thursday_morning=request.form.get('content_thu_morning', ''),
            thursday_snack=request.form.get('content_thu_snack', ''),
            thursday_dessert=request.form.get('content_thu_dessert', ''),
            thursday_lunch=request.form.get('content_thu_lunch', ''),
            thursday_afternoon=request.form.get('content_thu_afternoon', ''),
            thursday_lateafternoon=request.form.get('content_thu_lateafternoon', ''),
            # Friday
            friday_morning=request.form.get('content_fri_morning', ''),
            friday_snack=request.form.get('content_fri_snack', ''),
            friday_dessert=request.form.get('content_fri_dessert', ''),
            friday_lunch=request.form.get('content_fri_lunch', ''),
            friday_afternoon=request.form.get('content_fri_afternoon', ''),
            friday_lateafternoon=request.form.get('content_fri_lateafternoon', ''),
            # Saturday
            saturday_morning=request.form.get('content_sat_morning', ''),
            saturday_snack=request.form.get('content_sat_snack', ''),
            saturday_dessert=request.form.get('content_sat_dessert', ''),
            saturday_lunch=request.form.get('content_sat_lunch', ''),
            saturday_afternoon=request.form.get('content_sat_afternoon', ''),
            saturday_lateafternoon=request.form.get('content_sat_lateafternoon', '')
        )
        
        db.session.add(new_menu)
        db.session.commit()
        log_activity('create', 'menu', new_menu.id, f'T·∫°o th·ª±c ƒë∆°n tu·∫ßn {week_number}')
        flash('ƒê√£ th√™m th·ª±c ƒë∆°n m·ªõi!', 'success')
        return redirect(url_for('main.menu'))
    
    # Get all active dishes
    dishes = Dish.query.filter_by(is_active=True).all()
    mobile = is_mobile()
    
    # Calculate current week number for default value
    from datetime import datetime, timedelta, date
    today = datetime.now().date()
    week_start = today - timedelta(days=today.weekday())
    current_week = week_start.isocalendar()[1]
    current_year = week_start.year
    
    # Compute the Monday (start) and Sunday (end) dates for the current ISO week
    try:
        first_iso_week_ref = date(current_year, 1, 4)
        week_start_iso = first_iso_week_ref + timedelta(days=(current_week - 1) * 7)
        week_start_iso = week_start_iso - timedelta(days=week_start_iso.isoweekday() - 1)
        week_end_iso = week_start_iso + timedelta(days=6)
        current_week_start = week_start_iso.strftime('%d/%m/%Y')
        current_week_end = week_end_iso.strftime('%d/%m/%Y')
    except Exception:
        current_week_start = week_start.strftime('%d/%m/%Y')
        current_week_end = (week_start + timedelta(days=6)).strftime('%d/%m/%Y')
    
    return render_template('new_menu.html', title='T·∫°o th·ª±c ƒë∆°n m·ªõi', mobile=mobile, dishes=dishes, 
                         current_week=current_week, current_year=current_year,
                         current_week_start=current_week_start, current_week_end=current_week_end)

@main.route('/menu/<int:week_number>/edit', methods=['GET', 'POST'])
def edit_menu(week_number):
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    menu_item = Menu.query.filter_by(week_number=week_number, year=2025).first()
    if not menu_item:
        flash('Kh√¥ng t√¨m th·∫•y th·ª±c ƒë∆°n ƒë·ªÉ ch·ªânh s·ª≠a!', 'danger')
        return redirect(url_for('main.menu'))
    
    if request.method == 'POST':
        new_week_number = request.form.get('week_number', type=int)
        
        # Update menu fields
        menu_item.monday_morning = request.form.get('content_mon_morning', '')
        menu_item.monday_snack = request.form.get('content_mon_snack', '')
        menu_item.monday_dessert = request.form.get('content_mon_dessert', '')
        menu_item.monday_lunch = request.form.get('content_mon_lunch', '')
        menu_item.monday_afternoon = request.form.get('content_mon_afternoon', '')
        menu_item.monday_lateafternoon = request.form.get('content_mon_lateafternoon', '')
        
        menu_item.tuesday_morning = request.form.get('content_tue_morning', '')
        menu_item.tuesday_snack = request.form.get('content_tue_snack', '')
        menu_item.tuesday_dessert = request.form.get('content_tue_dessert', '')
        menu_item.tuesday_lunch = request.form.get('content_tue_lunch', '')
        menu_item.tuesday_afternoon = request.form.get('content_tue_afternoon', '')
        menu_item.tuesday_lateafternoon = request.form.get('content_tue_lateafternoon', '')
        
        menu_item.wednesday_morning = request.form.get('content_wed_morning', '')
        menu_item.wednesday_snack = request.form.get('content_wed_snack', '')
        menu_item.wednesday_dessert = request.form.get('content_wed_dessert', '')
        menu_item.wednesday_lunch = request.form.get('content_wed_lunch', '')
        menu_item.wednesday_afternoon = request.form.get('content_wed_afternoon', '')
        menu_item.wednesday_lateafternoon = request.form.get('content_wed_lateafternoon', '')
        
        menu_item.thursday_morning = request.form.get('content_thu_morning', '')
        menu_item.thursday_snack = request.form.get('content_thu_snack', '')
        menu_item.thursday_dessert = request.form.get('content_thu_dessert', '')
        menu_item.thursday_lunch = request.form.get('content_thu_lunch', '')
        menu_item.thursday_afternoon = request.form.get('content_thu_afternoon', '')
        menu_item.thursday_lateafternoon = request.form.get('content_thu_lateafternoon', '')
        
        menu_item.friday_morning = request.form.get('content_fri_morning', '')
        menu_item.friday_snack = request.form.get('content_fri_snack', '')
        menu_item.friday_dessert = request.form.get('content_fri_dessert', '')
        menu_item.friday_lunch = request.form.get('content_fri_lunch', '')
        menu_item.friday_afternoon = request.form.get('content_fri_afternoon', '')
        menu_item.friday_lateafternoon = request.form.get('content_fri_lateafternoon', '')
        
        menu_item.saturday_morning = request.form.get('content_sat_morning', '')
        menu_item.saturday_snack = request.form.get('content_sat_snack', '')
        menu_item.saturday_dessert = request.form.get('content_sat_dessert', '')
        menu_item.saturday_lunch = request.form.get('content_sat_lunch', '')
        menu_item.saturday_afternoon = request.form.get('content_sat_afternoon', '')
        menu_item.saturday_lateafternoon = request.form.get('content_sat_lateafternoon', '')
        
        # Check for duplicate week number if changed
        if new_week_number != menu_item.week_number:
            existing = Menu.query.filter_by(week_number=new_week_number, year=2025).first()
            if existing:
                flash(f'ƒê√£ t·ªìn t·∫°i th·ª±c ƒë∆°n tu·∫ßn {new_week_number}, kh√¥ng th·ªÉ ƒë·ªïi!', 'danger')
                return redirect(url_for('main.edit_menu', week_number=menu_item.week_number))
            menu_item.week_number = new_week_number
            
        db.session.commit()
        log_activity('edit', 'menu', menu_item.id, f'C·∫≠p nh·∫≠t th·ª±c ƒë∆°n tu·∫ßn {menu_item.week_number}')
        flash(f'ƒê√£ c·∫≠p nh·∫≠t th·ª±c ƒë∆°n tu·∫ßn {menu_item.week_number}!', 'success')
        return redirect(url_for('main.menu'))
    
    # Convert Menu fields to template format for editing
    data = {
        'mon': {
            'morning': menu_item.monday_morning or '',
            'snack': menu_item.monday_snack or '',
            'dessert': menu_item.monday_dessert or '',
            'lunch': menu_item.monday_lunch or '',
            'afternoon': menu_item.monday_afternoon or '',
            'lateafternoon': menu_item.monday_lateafternoon or ''
        },
        'tue': {
            'morning': menu_item.tuesday_morning or '',
            'snack': menu_item.tuesday_snack or '',
            'dessert': menu_item.tuesday_dessert or '',
            'lunch': menu_item.tuesday_lunch or '',
            'afternoon': menu_item.tuesday_afternoon or '',
            'lateafternoon': menu_item.tuesday_lateafternoon or ''
        },
        'wed': {
            'morning': menu_item.wednesday_morning or '',
            'snack': menu_item.wednesday_snack or '',
            'dessert': menu_item.wednesday_dessert or '',
            'lunch': menu_item.wednesday_lunch or '',
            'afternoon': menu_item.wednesday_afternoon or '',
            'lateafternoon': menu_item.wednesday_lateafternoon or ''
        },
        'thu': {
            'morning': menu_item.thursday_morning or '',
            'snack': menu_item.thursday_snack or '',
            'dessert': menu_item.thursday_dessert or '',
            'lunch': menu_item.thursday_lunch or '',
            'afternoon': menu_item.thursday_afternoon or '',
            'lateafternoon': menu_item.thursday_lateafternoon or ''
        },
        'fri': {
            'morning': menu_item.friday_morning or '',
            'snack': menu_item.friday_snack or '',
            'dessert': menu_item.friday_dessert or '',
            'lunch': menu_item.friday_lunch or '',
            'afternoon': menu_item.friday_afternoon or '',
            'lateafternoon': menu_item.friday_lateafternoon or ''
        },
        'sat': {
            'morning': menu_item.saturday_morning or '',
            'snack': menu_item.saturday_snack or '',
            'dessert': menu_item.saturday_dessert or '',
            'lunch': menu_item.saturday_lunch or '',
            'afternoon': menu_item.saturday_afternoon or '',
            'lateafternoon': menu_item.saturday_lateafternoon or ''
        }
    }
    
    # Get all active dishes
    dishes = Dish.query.filter_by(is_active=True).all()
    mobile = is_mobile()
    return render_template('edit_menu.html', week=menu_item, data=data, title=f'Ch·ªânh s·ª≠a th·ª±c ƒë∆°n tu·∫ßn {menu_item.week_number}', mobile=mobile, dishes=dishes)

@main.route('/menu/<int:week_number>/delete', methods=['POST'])
def delete_menu(week_number):
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    menu_item = Menu.query.filter_by(week_number=week_number, year=2025).first()
    if menu_item:
        db.session.delete(menu_item)
        db.session.commit()
        log_activity('delete', 'menu', menu_item.id, f'X√≥a th·ª±c ƒë∆°n tu·∫ßn {week_number}')
        flash(f'ƒê√£ xo√° th·ª±c ƒë∆°n tu·∫ßn {week_number}!', 'success')
    else:
        flash('Kh√¥ng t√¨m th·∫•y th·ª±c ƒë∆°n ƒë·ªÉ xo√°!', 'danger')
    return redirect(url_for('main.menu'))

@main.route('/menu/import', methods=['POST'])
def import_menu():
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    file = request.files.get('excel_file')
    week_number = request.form.get('week_number', type=int)
    if not file:
        flash('Vui l√≤ng ch·ªçn file Excel!', 'danger')
        return redirect(url_for('main.menu'))
    if not week_number:
        flash('Vui l√≤ng nh·∫≠p s·ªë tu·∫ßn!', 'danger')
        return redirect(url_for('main.menu'))

    from openpyxl import load_workbook
    wb = load_workbook(file)
    ws = wb.active

    # ƒê·ªçc d·ªØ li·ªáu t·ª´ d√≤ng 3 ƒë·∫øn 8, c·ªôt B-G (theo m·∫´u: A1:A2 "Th·ª©", B1:G1 "Khung gi·ªù", B2-G2 slot, A3-A8 th·ª©)
    days = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat']
    slots = ['morning', 'snack', 'dessert', 'lunch', 'afternoon', 'lateafternoon']
    menu_data = {}
    for i, day in enumerate(days):
        row = i + 3  # D√≤ng 3-8
        menu_data[day] = {}
        for j, slot in enumerate(slots):
            col = j + 2  # B=2, C=3, ... G=7
            value = ws.cell(row=row, column=col).value
            menu_data[day][slot] = value if value is not None else ""
    
    # Check if menu exists for this week
    menu_item = Menu.query.filter_by(week_number=week_number, year=2025).first()
    if menu_item:
        # Update existing menu
        menu_item.monday_morning = menu_data['mon']['morning']
        menu_item.monday_snack = menu_data['mon']['snack']
        menu_item.monday_dessert = menu_data['mon']['dessert']
        menu_item.monday_lunch = menu_data['mon']['lunch']
        menu_item.monday_afternoon = menu_data['mon']['afternoon']
        menu_item.monday_lateafternoon = menu_data['mon']['lateafternoon']
        
        menu_item.tuesday_morning = menu_data['tue']['morning']
        menu_item.tuesday_snack = menu_data['tue']['snack']
        menu_item.tuesday_dessert = menu_data['tue']['dessert']
        menu_item.tuesday_lunch = menu_data['tue']['lunch']
        menu_item.tuesday_afternoon = menu_data['tue']['afternoon']
        menu_item.tuesday_lateafternoon = menu_data['tue']['lateafternoon']
        
        menu_item.wednesday_morning = menu_data['wed']['morning']
        menu_item.wednesday_snack = menu_data['wed']['snack']
        menu_item.wednesday_dessert = menu_data['wed']['dessert']
        menu_item.wednesday_lunch = menu_data['wed']['lunch']
        menu_item.wednesday_afternoon = menu_data['wed']['afternoon']
        menu_item.wednesday_lateafternoon = menu_data['wed']['lateafternoon']
        
        menu_item.thursday_morning = menu_data['thu']['morning']
        menu_item.thursday_snack = menu_data['thu']['snack']
        menu_item.thursday_dessert = menu_data['thu']['dessert']
        menu_item.thursday_lunch = menu_data['thu']['lunch']
        menu_item.thursday_afternoon = menu_data['thu']['afternoon']
        menu_item.thursday_lateafternoon = menu_data['thu']['lateafternoon']
        
        menu_item.friday_morning = menu_data['fri']['morning']
        menu_item.friday_snack = menu_data['fri']['snack']
        menu_item.friday_dessert = menu_data['fri']['dessert']
        menu_item.friday_lunch = menu_data['fri']['lunch']
        menu_item.friday_afternoon = menu_data['fri']['afternoon']
        menu_item.friday_lateafternoon = menu_data['fri']['lateafternoon']
        
        menu_item.saturday_morning = menu_data['sat']['morning']
        menu_item.saturday_snack = menu_data['sat']['snack']
        menu_item.saturday_dessert = menu_data['sat']['dessert']
        menu_item.saturday_lunch = menu_data['sat']['lunch']
        menu_item.saturday_afternoon = menu_data['sat']['afternoon']
        menu_item.saturday_lateafternoon = menu_data['sat']['lateafternoon']
    else:
        # Create new menu
        new_menu = Menu(
            week_number=week_number,
            year=2025,
            monday_morning=menu_data['mon']['morning'],
            monday_snack=menu_data['mon']['snack'],
            monday_dessert=menu_data['mon']['dessert'],
            monday_lunch=menu_data['mon']['lunch'],
            monday_afternoon=menu_data['mon']['afternoon'],
            monday_lateafternoon=menu_data['mon']['lateafternoon'],
            
            tuesday_morning=menu_data['tue']['morning'],
            tuesday_snack=menu_data['tue']['snack'],
            tuesday_dessert=menu_data['tue']['dessert'],
            tuesday_lunch=menu_data['tue']['lunch'],
            tuesday_afternoon=menu_data['tue']['afternoon'],
            tuesday_lateafternoon=menu_data['tue']['lateafternoon'],
            
            wednesday_morning=menu_data['wed']['morning'],
            wednesday_snack=menu_data['wed']['snack'],
            wednesday_dessert=menu_data['wed']['dessert'],
            wednesday_lunch=menu_data['wed']['lunch'],
            wednesday_afternoon=menu_data['wed']['afternoon'],
            wednesday_lateafternoon=menu_data['wed']['lateafternoon'],
            
            thursday_morning=menu_data['thu']['morning'],
            thursday_snack=menu_data['thu']['snack'],
            thursday_dessert=menu_data['thu']['dessert'],
            thursday_lunch=menu_data['thu']['lunch'],
            thursday_afternoon=menu_data['thu']['afternoon'],
            thursday_lateafternoon=menu_data['thu']['lateafternoon'],
            
            friday_morning=menu_data['fri']['morning'],
            friday_snack=menu_data['fri']['snack'],
            friday_dessert=menu_data['fri']['dessert'],
            friday_lunch=menu_data['fri']['lunch'],
            friday_afternoon=menu_data['fri']['afternoon'],
            friday_lateafternoon=menu_data['fri']['lateafternoon'],
            
            saturday_morning=menu_data['sat']['morning'],
            saturday_snack=menu_data['sat']['snack'],
            saturday_dessert=menu_data['sat']['dessert'],
            saturday_lunch=menu_data['sat']['lunch'],
            saturday_afternoon=menu_data['sat']['afternoon'],
            saturday_lateafternoon=menu_data['sat']['lateafternoon']
        )
        db.session.add(new_menu)
    
    db.session.commit()
    flash('ƒê√£ import th·ª±c ƒë∆°n t·ª´ Excel!', 'success')
    return redirect(url_for('main.menu'))

@main.route('/curriculum/import', methods=['POST'])
def import_curriculum():
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    file = request.files.get('excel_file')

    week_number = request.form.get('week_number')
    class_id = request.form.get('class_id')
    if not class_id:
        flash('Vui l√≤ng ch·ªçn l·ªõp!', 'danger')
        return redirect(url_for('main.curriculum'))
    if not file:
        flash('Vui l√≤ng ch·ªçn file Excel!', 'danger')
        return redirect(url_for('main.curriculum'))
    if not week_number:
        flash('Vui l√≤ng nh·∫≠p s·ªë tu·∫ßn!', 'danger')
        return redirect(url_for('main.curriculum'))
    try:
        class_id = int(class_id)
    except Exception:
        flash('L·ªói l·ªõp h·ªçc kh√¥ng h·ª£p l·ªá!', 'danger')
        return redirect(url_for('main.curriculum'))
    try:
        week_number = int(week_number)
    except Exception:
        flash('L·ªói s·ªë tu·∫ßn kh√¥ng h·ª£p l·ªá!', 'danger')
        return redirect(url_for('main.curriculum'))

    from openpyxl import load_workbook
    wb = load_workbook(file)
    ws = wb.active

    # ƒê·ªçc d·ªØ li·ªáu theo m·∫´u m·ªõi:
    days = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat']
    morning_slots = ['morning_0', 'morning_1', 'morning_2', 'morning_3', 'morning_4', 'morning_5', 'morning_6']
    afternoon_slots = ['afternoon_1', 'afternoon_2', 'afternoon_3', 'afternoon_4']
    curriculum_data = {}

    # S√°ng: d√≤ng 4-10 (A4-A10)
    for col_idx, day in enumerate(days):
        curriculum_data[day] = {}
        for slot_idx, slot in enumerate(morning_slots):
            row = 4 + slot_idx  # d√≤ng 4-10
            col = 2 + col_idx   # B=2, C=3, ... G=7
            value = ws.cell(row=row, column=col).value
            curriculum_data[day][slot] = value if value is not None else ""
        # Chi·ªÅu: d√≤ng 11-14 (A11-A14)
        for slot_idx, slot in enumerate(afternoon_slots):
            row = 11 + slot_idx  # d√≤ng 11-14 (A12=afternoon_1, A13=afternoon_2, ...)
            col = 2 + col_idx
            value = ws.cell(row=row, column=col).value
            if slot == 'afternoon_2':
                # ƒê·∫£m b·∫£o ch·ªâ l·∫•y ƒë√∫ng d√≤ng 13 (A13) cho 15h-15h30
                if col_idx == 0:
                    curriculum_data[day][slot] = ""
                elif col_idx == 1 or col_idx == 3:
                    curriculum_data[day][slot] = value if value is not None else "Ho·∫°t ƒë·ªông v·ªõi gi√°o c·ª•"
                elif col_idx == 2 or col_idx == 4:
                    curriculum_data[day][slot] = value if value is not None else "Lego time"
                else:
                    curriculum_data[day][slot] = value if value is not None else ""
            else:
                curriculum_data[day][slot] = value if value is not None else ""
    import json
    content = json.dumps(curriculum_data, ensure_ascii=False)
    # ƒê·∫£m b·∫£o kh√¥ng b·ªã ƒë√® curriculum c·ªßa l·ªõp kh√°c c√πng tu·∫ßn
    week = Curriculum.query.filter_by(week_number=week_number, class_id=class_id).first()
    if week:
        week.content = content
    else:
        new_week = Curriculum(week_number=week_number, class_id=class_id, content=content, material=None)
        db.session.add(new_week)
    db.session.commit()
    flash('ƒê√£ import ch∆∞∆°ng tr√¨nh h·ªçc t·ª´ Excel!', 'success')
    return redirect(url_for('main.curriculum', class_id=class_id))

@main.route('/curriculum/export', methods=['GET'])
def export_curriculum_template():
    """Export file Excel m·∫´u ch∆∞∆°ng tr√¨nh h·ªçc v·ªõi ƒë·ªãnh d·∫°ng n√¢ng cao (merged cells, header, khung gi·ªù, ...)."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
    from io import BytesIO

    wb = Workbook()
    ws = wb.active
    ws.title = "Curriculum Template"

    # ƒê·ªãnh nghƒ©a style
    bold = Font(bold=True)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    fill = PatternFill(start_color="C8E6C9", end_color="C8E6C9", fill_type="solid")

    # Header
    ws.merge_cells('A1:A2')
    ws['A1'] = "Khung gi·ªù"
    ws['A1'].font = bold
    ws['A1'].alignment = center
    ws['A1'].fill = fill
    ws['A1'].border = border
    ws.merge_cells('B1:G1')
    ws['B1'] = "Th·ª©"
    ws['B1'].font = bold
    ws['B1'].alignment = center
    ws['B1'].fill = fill
    ws['B1'].border = border
    days = ["Th·ª© 2", "Th·ª© 3", "Th·ª© 4", "Th·ª© 5", "Th·ª© 6", "Th·ª© 7"]
    for i, day in enumerate(days):
        cell = ws.cell(row=2, column=2+i, value=day)
        cell.font = bold
        cell.alignment = center
        cell.fill = fill
        cell.border = border

    # Section: Bu·ªïi s√°ng
    ws.merge_cells('A3:G3')
    ws['A3'] = "Bu·ªïi s√°ng"
    ws['A3'].font = bold
    ws['A3'].alignment = center
    ws['A3'].fill = fill
    ws['A3'].border = border

    morning_slots = [
        ("7-17h", 4), ("7-8h", 5), ("8h-8h30", 6), ("8h30-9h", 7), ("9h-9h40", 8), ("9h40-10h30", 9), ("10h30-14h", 10)
    ]
    for idx, (label, row) in enumerate(morning_slots):
        ws.cell(row=row, column=1, value=label).font = bold
        ws.cell(row=row, column=1).alignment = center
        ws.cell(row=row, column=1).border = border
        if idx == 0:
            # Slot 7-17h: ƒëi·ªÅn gi√° tr·ªã m·∫∑c ƒë·ªãnh cho t·ª´ng th·ª©
            for col in range(2, 8):
                ws.cell(row=row, column=col).border = border
                if col == 3:
                    ws.cell(row=row, column=col, value="To√°n h·ªçc")
                elif col == 4:
                    ws.cell(row=row, column=col, value="Ng√¥n Ng·ªØ")
                elif col == 5:
                    ws.cell(row=row, column=col, value="Stemax")
                elif col == 6:
                    ws.cell(row=row, column=col, value="Tr·∫£i Nghi·ªám")
        elif idx == 1:
            # Merge cell cho 7-8h (B5:G5)
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=7)
            merged_cell = ws.cell(row=row, column=2)
            merged_cell.value = "ƒê√≥n tr·∫ª - STEAM (massage k√≠ch th√≠ch gi√°c quan) - ƒÇn s√°ng"
            merged_cell.alignment = center
            merged_cell.font = Font(bold=False)
            merged_cell.border = border
            for col in range(2, 8):
                ws.cell(row=row, column=col).border = border
        elif idx == 2:
            # Merge cell cho 8h-8h30 (B6:G6)
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=7)
            merged_cell = ws.cell(row=row, column=2)
            merged_cell.value = "Th·ªÉ d·ª•c bu·ªïi s√°ng - Tr√≤ chuy·ªán ƒë·∫ßu ng√†y - Ki·ªÉm tra th√¢n th·ªÉ"
            merged_cell.alignment = center
            merged_cell.font = Font(bold=False)
            merged_cell.border = border
            for col in range(2, 8):
                ws.cell(row=row, column=col).border = border
        elif idx == 6:
            # Merge cell cho 10h30-14h (B10:G10)
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=7)
            merged_cell = ws.cell(row=row, column=2)
            merged_cell.value = "V·ªá sinh ƒÉn tr∆∞a - ng·ªß tr∆∞a"
            merged_cell.alignment = center
            merged_cell.font = Font(bold=False)
            merged_cell.border = border
            for col in range(2, 8):
                ws.cell(row=row, column=col).border = border
        else:
            for col in range(2, 8):
                ws.cell(row=row, column=col).border = border

    # Section: Bu·ªïi chi·ªÅu
    ws.merge_cells('A11:G11')
    ws['A11'] = "Bu·ªïi chi·ªÅu"
    ws['A11'].font = bold
    ws['A11'].alignment = center
    ws['A11'].fill = fill
    ws['A11'].border = border

    afternoon_slots = [
        ("14h15-15h", 12), ("15h-15h30", 13), ("15h45-16h", 14), ("16h-17h", 15)
    ]
    for idx, (label, row) in enumerate(afternoon_slots):
        ws.cell(row=row, column=1, value=label).font = bold
        ws.cell(row=row, column=1).alignment = center
        ws.cell(row=row, column=1).border = border
        if idx == 0:
            # Merge cell cho 14h15-15h (B12:G12)
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=7)
            merged_cell = ws.cell(row=row, column=2)
            merged_cell.value = "V·ªá sinh - u·ªëng n∆∞·ªõc - v·∫≠n ƒë·ªông nh·∫π - ƒÉn chi·ªÅu"
            merged_cell.alignment = center
            merged_cell.font = Font(bold=False)
            merged_cell.border = border
            for col in range(2, 8):
                ws.cell(row=row, column=col).border = border
        elif idx == 1:
            # Set default values for 15h-15h30 slot
            # Th·ª© 3 (col=3): Ho·∫°t ƒë·ªông v·ªõi gi√°o c·ª•
            # Th·ª© 4 (col=4): Lego time
            # Th·ª© 5 (col=5): Ho·∫°t ƒë·ªông v·ªõi gi√°o c·ª•
            # Th·ª© 6 (col=6): Lego time
            for col in range(2, 8):
                ws.cell(row=row, column=col).border = border
                if col == 3 or col == 5:
                    ws.cell(row=row, column=col, value="Ho·∫°t ƒë·ªông v·ªõi gi√°o c·ª•")
                elif col == 4 or col == 6:
                    ws.cell(row=row, column=col, value="Lego time")
        elif idx == 2:
            # Merge cell cho 15h45-16h (B14:G14)
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=7)
            merged_cell = ws.cell(row=row, column=2)
            merged_cell.value = "Yoga/dance"
            merged_cell.alignment = center
            merged_cell.font = Font(bold=False)
            merged_cell.border = border
            for col in range(2, 8):
                ws.cell(row=row, column=col).border = border
        elif idx == 3:
            # Merge cell cho 16h-17h (B15:G15)
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=7)
            merged_cell = ws.cell(row=row, column=2)
            merged_cell.value = "Tr·∫£ tr·∫ª - trao ƒë·ªïi v·ªõi ph·ª• huynh"
            merged_cell.alignment = center
            merged_cell.font = Font(bold=False)
            merged_cell.border = border
            for col in range(2, 8):
                ws.cell(row=row, column=col).border = border
        else:
            for col in range(2, 8):
                ws.cell(row=row, column=col).border = border

    # Set column widths
    ws.column_dimensions['A'].width = 16
    for col in ['B','C','D','E','F','G']:
        ws.column_dimensions[col].width = 18

    # Save to BytesIO
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return send_file(output, download_name="curriculum_template.xlsx", as_attachment=True)

@main.route('/menu/export', methods=['GET'])
def export_menu_template():
    """Export file Excel m·∫´u th·ª±c ƒë∆°n v·ªõi ƒë·ªãnh d·∫°ng n√¢ng cao (merged cells, header, khung gi·ªù, ...)."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
    from io import BytesIO

    wb = Workbook()
    ws = wb.active
    ws.title = "Menu Template"

    # ƒê·ªãnh nghƒ©a style
    bold = Font(bold=True)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    fill = PatternFill(start_color="C8E6C9", end_color="C8E6C9", fill_type="solid")

    # Header
    ws.merge_cells('A1:A2')
    ws['A1'] = "Th·ª©"
    ws['A1'].font = bold
    ws['A1'].alignment = center
    ws['A1'].fill = fill
    ws['A1'].border = border
    ws.merge_cells('B1:G1')
    ws['B1'] = "Khung gi·ªù"
    ws['B1'].font = bold
    ws['B1'].alignment = center
    ws['B1'].fill = fill
    ws['B1'].border = border
    slots = ["S√°ng", "Ph·ª• s√°ng", "Tr√°ng mi·ªáng", "Tr∆∞a", "X·∫ø", "X·∫ø chi·ªÅu"]
    for i, slot in enumerate(slots):
        cell = ws.cell(row=2, column=2+i, value=slot)
        cell.font = bold
        cell.alignment = center
        cell.fill = fill
        cell.border = border

    # Fill days
    days = ["Th·ª© 2", "Th·ª© 3", "Th·ª© 4", "Th·ª© 5", "Th·ª© 6", "Th·ª© 7"]
    for i, day in enumerate(days):
        cell = ws.cell(row=3+i, column=1, value=day)
        cell.font = bold
        cell.alignment = center
        cell.border = border
        for col in range(2, 8):
            ws.cell(row=3+i, column=col).border = border

    # Set column widths
    ws.column_dimensions['A'].width = 16
    for col in ['B','C','D','E','F','G']:
        ws.column_dimensions[col].width = 18

    # Save to BytesIO
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return send_file(output, download_name="menu_template.xlsx", as_attachment=True)

@main.route('/activities/<int:id>/delete-image/<int:image_id>', methods=['POST'])
def delete_activity_image(id, image_id):
    try:
        if session.get('role') not in ['admin', 'teacher']:
            print(f"[LOG] Kh√¥ng c√≥ quy·ªÅn xo√° ·∫£nh ho·∫°t ƒë·ªông")
            return redirect_no_permission()
        img = ActivityImage.query.get_or_404(image_id)
        print(f"[LOG] ƒêang xo√° ·∫£nh: id={image_id}, filepath={img.filepath}")
        
        # Ki·ªÉm tra n·∫øu l√† URL R2 (b·∫Øt ƒë·∫ßu b·∫±ng http)
        if img.filepath.startswith('http'):
            # X√≥a tr√™n R2
            try:
                if R2_ENABLED:
                    r2 = get_r2_storage()
                    if r2.enabled:
                        r2.delete_file(img.filepath)
                        print(f"[LOG] ƒê√£ x√≥a R2: {img.filepath}")
            except Exception as e:
                print(f"[LOG] Kh√¥ng th·ªÉ x√≥a R2: {e}")
        else:
            # X√≥a file local
            img_path = os.path.join('app', 'static', img.filepath)
            if os.path.exists(img_path):
                os.remove(img_path)
                print(f"[LOG] ƒê√£ xo√° file v·∫≠t l√Ω: {img_path}")
            else:
                print(f"[LOG] File v·∫≠t l√Ω kh√¥ng t·ªìn t·∫°i: {img_path}")
        
        db.session.delete(img)
        db.session.commit()
        print(f"[LOG] ƒê√£ xo√° b·∫£n ghi ActivityImage id={image_id} kh·ªèi DB")
        flash('ƒê√£ xo√° ·∫£nh ho·∫°t ƒë·ªông!', 'success')
        return redirect(url_for('main.edit_activity', id=id))
    except Exception as e:
        print(f"[ERROR] L·ªói khi xo√° ·∫£nh ho·∫°t ƒë·ªông: {e}")
        import traceback
        traceback.print_exc()
        flash(f"L·ªói khi xo√° ·∫£nh ho·∫°t ƒë·ªông: {e}", 'danger')
        return redirect(url_for('main.edit_activity', id=id))

@main.route('/menu/<int:week_number>/export-food-safety', methods=['GET'])
def export_food_safety_process(week_number):
    """Xu·∫•t quy tr√¨nh an to√†n th·ª±c ph·∫©m 3 b∆∞·ªõc theo ti√™u chu·∫©n chuy√™n nghi·ªáp v·ªõi ƒë·∫ßy ƒë·ªß th√¥ng tin qu·∫£n l√Ω."""
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    
    # L·∫•y th·ª±c ƒë∆°n c·ªßa tu·∫ßn (s·ª≠ d·ª•ng Menu model)
    menu_item = Menu.query.filter_by(week_number=week_number, year=2025).first()
    if not menu_item:
        flash('Kh√¥ng t√¨m th·∫•y th·ª±c ƒë∆°n!', 'danger')
        return redirect(url_for('main.menu'))
    
    import json
    if not OPENPYXL_AVAILABLE:
        flash('Ch·ª©c nƒÉng n√†y c·∫ßn c√†i ƒë·∫∑t openpyxl. Vui l√≤ng li√™n h·ªá qu·∫£n tr·ªã vi√™n.', 'warning')
        return redirect(url_for('main.menu'))
    
    from openpyxl import Workbook
    from openpyxl.styles import Font, Border, Side, Alignment, PatternFill
    from io import BytesIO
    import zipfile
    from datetime import datetime, timedelta
    
    # Convert Menu model data to the same format as old Curriculum JSON
    menu_data = {
        'mon': {
            'morning': menu_item.monday_morning or '',
            'snack': menu_item.monday_snack or '',
            'dessert': menu_item.monday_dessert or '',
            'lunch': menu_item.monday_lunch or '',
            'afternoon': menu_item.monday_afternoon or '',
            'lateafternoon': menu_item.monday_lateafternoon or ''
        },
        'tue': {
            'morning': menu_item.tuesday_morning or '',
            'snack': menu_item.tuesday_snack or '',
            'dessert': menu_item.tuesday_dessert or '',
            'lunch': menu_item.tuesday_lunch or '',
            'afternoon': menu_item.tuesday_afternoon or '',
            'lateafternoon': menu_item.tuesday_lateafternoon or ''
        },
        'wed': {
            'morning': menu_item.wednesday_morning or '',
            'snack': menu_item.wednesday_snack or '',
            'dessert': menu_item.wednesday_dessert or '',
            'lunch': menu_item.wednesday_lunch or '',
            'afternoon': menu_item.wednesday_afternoon or '',
            'lateafternoon': menu_item.wednesday_lateafternoon or ''
        },
        'thu': {
            'morning': menu_item.thursday_morning or '',
            'snack': menu_item.thursday_snack or '',
            'dessert': menu_item.thursday_dessert or '',
            'lunch': menu_item.thursday_lunch or '',
            'afternoon': menu_item.thursday_afternoon or '',
            'lateafternoon': menu_item.thursday_lateafternoon or ''
        },
        'fri': {
            'morning': menu_item.friday_morning or '',
            'snack': menu_item.friday_snack or '',
            'dessert': menu_item.friday_dessert or '',
            'lunch': menu_item.friday_lunch or '',
            'afternoon': menu_item.friday_afternoon or '',
            'lateafternoon': menu_item.friday_lateafternoon or ''
        },
        'sat': {
            'morning': menu_item.saturday_morning or '',
            'snack': menu_item.saturday_snack or '',
            'dessert': menu_item.saturday_dessert or '',
            'lunch': menu_item.saturday_lunch or '',
            'afternoon': menu_item.saturday_afternoon or '',
            'lateafternoon': menu_item.saturday_lateafternoon or ''
        }
    }
    
    # L·∫•y th√¥ng tin suppliers chi ti·∫øt
    from app.models import Supplier
    suppliers = Supplier.query.all()
    supplier_dict = {}
    for supplier in suppliers:
        supplier_dict[supplier.name] = {
            'address': supplier.address or 'Ch∆∞a c·∫≠p nh·∫≠t ƒë·ªãa ch·ªâ',
            'phone': supplier.phone or 'Ch∆∞a c·∫≠p nh·∫≠t SƒêT',
            'contact_person': supplier.contact_person or 'Ch∆∞a c·∫≠p nh·∫≠t ng∆∞·ªùi li√™n h·ªá',
            'food_safety_cert': supplier.food_safety_cert or '',
            'established_date': getattr(supplier, 'established_date', 'Ch∆∞a c·∫≠p nh·∫≠t')
        }
    
    # T√≠nh s·ªë h·ªçc sinh c√≥ m·∫∑t theo t·ª´ng ng√†y trong tu·∫ßn
    def get_daily_attendance_for_week(week_number):
        """Tr·∫£ v·ªÅ dict s·ªë h·ªçc sinh c√≥ m·∫∑t m·ªói ng√†y trong tu·∫ßn"""
        from datetime import date
        year = datetime.now().year
        week_start = date.fromisocalendar(year, int(week_number), 1)
        week_dates = [week_start + timedelta(days=i) for i in range(6)]  # Th·ª© 2 ƒë·∫øn Th·ª© 7
        
        daily_attendance = {}
        days_vn = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat']
        
        for i, day in enumerate(week_dates):
            day_str = day.strftime('%Y-%m-%d')
            # ƒê·∫øm s·ªë h·ªçc sinh C√ì M·∫∂T trong ng√†y n√†y
            count = AttendanceRecord.query.filter(
                AttendanceRecord.date == day_str,
                AttendanceRecord.status == 'C√≥ m·∫∑t'
            ).count()
            # N·∫øu kh√¥ng c√≥ d·ªØ li·ªáu ƒëi·ªÉm danh, d√πng t·ªïng s·ªë h·ªçc sinh active
            if count == 0:
                count = Child.query.filter_by(is_active=True).count()
            daily_attendance[days_vn[i]] = count
        
        return daily_attendance
    
    daily_attendance = get_daily_attendance_for_week(week_number)
    
    

    # Helper: L·∫•y nguy√™n li·ªáu th·ª±c t·∫ø t·ª´ database
    from app.models import Dish, DishIngredient, Product
    def get_dish_ingredients(dish_name):
        dish = Dish.query.filter_by(name=dish_name).first()
        if not dish:
            return []
        return dish.ingredients

    def get_ingredient_info(dish_ingredient, student_count):
        product = dish_ingredient.product
        total_qty = dish_ingredient.quantity * student_count
        return {
            'name': product.name,
            'unit': dish_ingredient.unit,
            'total_qty': total_qty,
            'supplier': product.supplier.name if product.supplier else '',
            'supplier_info': {
                'address': product.supplier.address if product.supplier else '',
                'phone': product.supplier.phone if product.supplier else '',
                'contact_person': product.supplier.contact_person if product.supplier else '',
                'food_safety_cert': product.supplier.food_safety_cert if product.supplier else ''
            }
        }
   
    
    # Refactor: Aggregate all ingredients from the actual weekly menu using real dish/ingredient data
    from collections import defaultdict
    ingredient_totals = defaultdict(lambda: {'total_qty': 0, 'unit': '', 'category': '', 'supplier': None, 'product': None, 'usage_frequency': 0})
    dish_by_day = defaultdict(list)  # L∆∞u m√≥n ƒÉn theo t·ª´ng ng√†y
    dishes = set()

    # 1. L∆∞u m√≥n ƒÉn theo t·ª´ng ng√†y ƒë·ªÉ t√≠nh nguy√™n li·ªáu theo s·ªë h·ªçc sinh c√≥ m·∫∑t m·ªói ng√†y
    days_order = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat']
    for day in days_order:
        if day in menu_data:
            day_data = menu_data[day]
            for slot_dish in day_data.values():
                if slot_dish:
                    # Support both single dish and comma-separated dishes
                    for dish_name in [d.strip() for d in slot_dish.split(',') if d.strip()]:
                        dish_by_day[day].append(dish_name)
                        dishes.add(dish_name)

    # 2. T√≠nh nguy√™n li·ªáu cho t·ª´ng ng√†y v·ªõi s·ªë h·ªçc sinh c√≥ m·∫∑t c·ªßa ng√†y ƒë√≥
    for day, dish_list in dish_by_day.items():
        students_today = daily_attendance.get(day, 0)
        if students_today == 0:
            continue
            
        # ƒê·∫øm s·ªë l·∫ßn m·ªói m√≥n xu·∫•t hi·ªán trong ng√†y n√†y
        dish_count_today = defaultdict(int)
        for dish_name in dish_list:
            dish_count_today[dish_name] += 1
        
        # T√≠nh nguy√™n li·ªáu cho ng√†y n√†y
        for dish_name, count_in_day in dish_count_today.items():
            dish = Dish.query.filter_by(name=dish_name).first()
            if not dish:
                continue
            for di in dish.ingredients:
                product = di.product
                if not product:
                    continue
                key = (di.product.name, di.unit, di.product.category, di.product.supplier)
                # T√≠nh: (l∆∞·ª£ng/h·ªçc sinh) √ó (s·ªë h·ªçc sinh c√≥ m·∫∑t h√¥m nay) √ó (s·ªë l·∫ßn m√≥n xu·∫•t hi·ªán trong ng√†y)
                qty = di.quantity * students_today * count_in_day
                
                if key not in ingredient_totals:
                    ingredient_totals[key] = {'total_qty': 0, 'unit': di.unit, 'category': di.product.category, 'supplier': di.product.supplier, 'product': di.product, 'usage_frequency': 0}
                ingredient_totals[key]['total_qty'] += qty
                ingredient_totals[key]['usage_frequency'] += 1
    
    # 3. Split into fresh, dry, fruit by category
    fresh_ingredients_with_qty = []
    dry_ingredients_with_qty = []
    fruit_ingredients_with_qty = []
    
    def convert_to_kg(quantity, unit):
        """
        Convert quantity to kg/l√≠t based on unit
        - M√≥n ƒÉn th∆∞·ªùng l∆∞u theo gram/ml cho 1 h·ªçc sinh
        - Phi·∫øu ti·∫øp nh·∫≠n c·∫ßn kg/l√≠t
        - N·∫øu ƒë∆°n v·ªã l√† gram/ml ‚Üí chia 1000
        - N·∫øu ƒë∆°n v·ªã ƒë√£ l√† kg/l√≠t ‚Üí gi·ªØ nguy√™n
        """
        unit = unit.lower().strip()
        
        # ƒê∆°n v·ªã nh·ªè (gram, ml) ‚Üí chuy·ªÉn sang kg/l√≠t
        if 'gram' in unit or unit == 'g' or 'gr' in unit:
            return quantity / 1000
        elif 'ml' in unit or 'milliliter' in unit:
            return quantity / 1000
        # ƒê∆°n v·ªã l·ªõn (kg, l√≠t) ‚Üí gi·ªØ nguy√™n
        elif 'kg' in unit or 'kilogram' in unit:
            return quantity
        elif 'l√≠t' in unit or 'liter' in unit or unit == 'l':
            return quantity
        # ƒê∆°n v·ªã kh√°c (g√≥i, h·ªôp, qu·∫£, c·ªß...) ‚Üí gi·ªØ nguy√™n
        else:
            return quantity
    
    for name, info in ingredient_totals.items():
        # Properly convert to kg based on unit
        weight_kg = convert_to_kg(info['total_qty'], info['unit'])
        
        row = {
            'name': name[0].title() if isinstance(name, tuple) else str(name).title(),
            'weight_kg': round(weight_kg, 2),
            'unit': info['unit'],
            'category': info['category'],
            'supplier': info['supplier'],
            'supplier_info': supplier_dict.get(info['supplier'].name if info['supplier'] else '', {
                'address': 'ƒê·ªãa ch·ªâ ch∆∞a c·∫≠p nh·∫≠t',
                'phone': 'SƒêT ch∆∞a c·∫≠p nh·∫≠t', 
                'contact_person': 'Ng∆∞·ªùi li√™n h·ªá ch∆∞a c·∫≠p nh·∫≠t',
                'food_safety_cert': 'Ch∆∞a c√≥ gi·∫•y ch·ª©ng nh·∫≠n'
            }),
            'usage_frequency': info.get('usage_frequency', 0)
        }
        cat = (info['category'] or '').lower()
        if 't∆∞∆°i' in cat or 'rau' in cat or 'th·ªãt' in cat or 'c√°' in cat or 'tr·ª©ng' in cat:
            fresh_ingredients_with_qty.append(row)
        elif 'kh√¥' in cat or 'gia v·ªã' in cat or 'b·ªôt' in cat or 'g·∫°o' in cat or 'ƒë∆∞·ªùng' in cat:
            dry_ingredients_with_qty.append(row)
        elif 'tr√°i c√¢y' in cat or 'hoa qu·∫£' in cat or 'fruit' in cat:
            fruit_ingredients_with_qty.append(row)
        else:
            # Default: fresh if unknown
            fresh_ingredients_with_qty.append(row)

    # Sort by usage frequency
    fresh_ingredients_with_qty.sort(key=lambda x: x['usage_frequency'], reverse=True)
    dry_ingredients_with_qty.sort(key=lambda x: x['usage_frequency'], reverse=True)
    fruit_ingredients_with_qty.sort(key=lambda x: x['usage_frequency'], reverse=True)
    
    zip_buffer = BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w') as zipf:
        
        # B∆Ø·ªöC 1.1: Ti·∫øp nh·∫≠n th·ª±c ph·∫©m t∆∞∆°i - Xu·∫•t m·ªói ng√†y 1 sheet, ƒë√∫ng menu/ng√†y
        from datetime import date, timedelta, datetime
        year = datetime.now().year
        week_start = date.fromisocalendar(year, int(week_number), 1)
        days_vn = ["Th·ª© 2", "Th·ª© 3", "Th·ª© 4", "Th·ª© 5", "Th·ª© 6", "Th·ª© 7"]
        wb1 = Workbook()
        # Style
        from openpyxl.styles import Font, Border, Side, Alignment, PatternFill
        thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
        thick_border = Border(left=Side(style='medium'), right=Side(style='medium'), top=Side(style='medium'), bottom=Side(style='medium'))
        for day_offset in range(6):
            day_date = week_start + timedelta(days=day_offset)
            day_key = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat'][day_offset]
            safe_date = day_date.strftime('%d-%m')  # Kh√¥ng d√πng d·∫•u '/'
            sheet_title = f"{days_vn[day_offset]} ({safe_date})"
            if day_offset == 0:
                ws1 = wb1.active
                ws1.title = sheet_title
            else:
                ws1 = wb1.create_sheet(title=sheet_title)
            # L·∫•y menu ng√†y
            menu_today = menu_data.get(day_key, {})
            # L·∫•y s·ªë h·ªçc sinh c√≥ m·∫∑t ng√†y n√†y
            students_today = daily_attendance.get(day_key, 0)
            # T√≠nh nguy√™n li·ªáu th·ª±c t·∫ø cho ng√†y n√†y
            daily_ingredients = {}
            for meal in menu_today.values():
                if not meal: continue
                for dish_name in [d.strip() for d in meal.split(',') if d.strip()]:
                    dish = Dish.query.filter_by(name=dish_name).first()
                    if dish:
                        for di in dish.ingredients:
                            key = (di.product.name, di.unit, di.product.category, di.product.supplier)
                            qty = di.quantity * students_today
                            if key not in daily_ingredients:
                                daily_ingredients[key] = {'total_qty': 0, 'unit': di.unit, 'category': di.product.category, 'supplier': di.product.supplier, 'product': di.product}
                            daily_ingredients[key]['total_qty'] += qty
            # Ph√¢n lo·∫°i t∆∞∆°i (s·ª≠ d·ª•ng logic gi·ªëng nh∆∞ t√≠nh to√°n tu·∫ßn)
            fresh_ingredients = []
            for (name, unit, category, supplier), info in daily_ingredients.items():
                cat = (category or '').lower()
                if 't∆∞∆°i' in cat or 'rau' in cat or 'th·ªãt' in cat or 'c√°' in cat or 'tr·ª©ng' in cat or cat == 'fresh':
                    # supplier c√≥ th·ªÉ l√† object, c·∫ßn l·∫•y t√™n ho·∫∑c chu·ªói
                    if hasattr(supplier, 'name'):
                        supplier_name = supplier.name
                    elif isinstance(supplier, str):
                        supplier_name = supplier
                    else:
                        supplier_name = ''
                    
                    # Use the same convert_to_kg function for consistency
                    weight_kg = convert_to_kg(info['total_qty'], info['unit'])
                    
                    fresh_ingredients.append({
                        'name': name,
                        'weight_kg': round(weight_kg, 2),
                        'unit': unit,
                        'category': category,
                        'supplier': supplier_name,
                        'supplier_info': supplier_dict.get(supplier_name, {}),
                    })
            # --- Ghi d·ªØ li·ªáu v√† style sheet nh∆∞ c≈© ---
            ws1['A1'] = "T√äN C∆† S·ªû: MNƒêL C√¢y Nh·ªè"
            ws1['A1'].font = Font(bold=True, size=12)
            ws1['A1'].fill = PatternFill(start_color="FFE6CC", end_color="FFE6CC", fill_type="solid")
            ws1.merge_cells('A1:P1')
            ws1['D2'] = "BI·ªÇU M·∫™U KI·ªÇM TRA TR∆Ø·ªöC KHI CH·∫æ BI·∫æN TH·ª®C ƒÇN"
            ws1['D2'].font = Font(bold=True, size=14, color="FF0000")
            ws1['D2'].alignment = Alignment(horizontal='center', vertical='center')
            ws1.merge_cells('D2:M2')
            ws1['O2'] = "S·ªë: 1246/Qƒê - B·ªô Y T·∫ø"
            ws1['O2'].font = Font(bold=True, size=10)
            ws1['O2'].fill = PatternFill(start_color="FFCCCC", end_color="FFCCCC", fill_type="solid")
            info_data = [
                (3, 'A', f"Ng∆∞·ªùi ki·ªÉm tra: Nguy·ªÖn Th·ªã V√¢n", 'O', "M·∫´u s·ªë 1.1"),
                (4, 'A', f"Ng√†y ki·ªÉm tra: {day_date.strftime('%d/%m/%Y')} - {days_vn[day_offset]}", 'O', f"S·ªë h·ªçc sinh: {students_today}"),
                (5, 'A', "ƒê·ªãa ƒëi·ªÉm: B·∫øp ƒÉn Tr∆∞·ªùng MNƒêL C√¢y Nh·ªè", 'O', "Phi√™n b·∫£n: v2.0")
            ]
            for row, col_a, text_a, col_o, text_o in info_data:
                ws1[f'{col_a}{row}'] = text_a
                ws1[f'{col_a}{row}'].font = Font(bold=True, size=10)
                ws1[f'{col_o}{row}'] = text_o
                ws1[f'{col_o}{row}'].font = Font(bold=True, size=10)
                ws1[f'{col_o}{row}'].fill = PatternFill(start_color="E6F3FF", end_color="E6F3FF", fill_type="solid")
            ws1['A7'] = "PH·∫¶N I: TH·ª∞C PH·∫®M T∆Ø∆†I S·ªêNG, ƒê√îNG L·∫†NH (Th·ªãt, c√°, rau, c·ªß, qu·∫£...)"
            ws1['A7'].font = Font(bold=True, size=12, color="0066CC")
            ws1['A7'].fill = PatternFill(start_color="E6F3FF", end_color="E6F3FF", fill_type="solid")
            ws1.merge_cells('A7:M7')
            ws1['O7'] = "B∆Ø·ªöC 1.1"
            ws1['O7'].font = Font(bold=True, size=12, color="FF0000")
            ws1['O7'].fill = PatternFill(start_color="FFEEEE", end_color="FFEEEE", fill_type="solid")
            headers_main = [
                'STT', 'T√äN TH·ª∞C PH·∫®M', 'TH·ªúI GIAN NH·∫¨P\n(Ng√†y/Gi·ªù)', 
                'KH·ªêI L∆Ø·ª¢NG\n(kg/l√≠t)', 'N∆†I CUNG C·∫§P', '', '', 'S·ªê CH·ª®NG T·ª™/S·ªê HO√Å ƒê∆†N',
                'GI·∫§Y ƒêƒÇNG K√ù V·ªöI TH√ö Y', 'GI·∫§Y KI·ªÇM D·ªäCH',
                'KI·ªÇM TRA C·∫¢M QUAN', '',
                'X√âT NGHI·ªÜM NHANH', '',
                'BI·ªÜN PH√ÅP X·ª¨ L√ù/ GHI CH√ö'
            ]
            for i, header in enumerate(headers_main, 1):
                cell = ws1.cell(row=8, column=i, value=header)
                cell.font = Font(bold=True, size=9, color="FFFFFF")
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                cell.border = thick_border
            sub_headers = [
                '', '', '', '', 'T√™n c∆° s·ªü', 'SƒêT/ƒê·ªãa ch·ªâ', 'Ng∆∞·ªùi Giao H√†ng', '','', '',
                'ƒê·∫°t', 'Kh√¥ng ƒë·∫°t', 'ƒê·∫°t', 'Kh√¥ng ƒë·∫°t', ''
            ]
            for i, header in enumerate(sub_headers, 1):
                cell = ws1.cell(row=9, column=i, value=header)
                cell.font = Font(bold=True, size=8)
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.fill = PatternFill(start_color="B4C6E7", end_color="B4C6E7", fill_type="solid")
                cell.border = thin_border
            ws1.merge_cells('E8:G8')
            ws1.merge_cells('K8:L8')
            ws1.merge_cells('M8:N8')

            # S·ªë th·ª© t·ª± c·ªôt
            for i in range(1, 16):
                cell = ws1.cell(row=10, column=i, value=i)
                cell.font = Font(bold=True, size=8)
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.fill = PatternFill(start_color="B4C6E7", end_color="B4C6E7", fill_type="solid")
                cell.border = thin_border
            # Ghi d·ªØ li·ªáu th·ª±c ph·∫©m t∆∞∆°i t·ª´ng ng√†y
            for i, ingredient_info in enumerate(fresh_ingredients[:25], 1):
                row_num = 10 + i
                supplier_info = ingredient_info.get('supplier_info', {})
                supplier_name = ingredient_info.get('supplier', '') or 'CTY TNHH Th·ª±c ph·∫©m An to√†n'
                phone = supplier_info.get('phone', '0902.xxx.xxx')
                address = supplier_info.get('address', 'ƒê√† L·∫°t')
                contact_person = supplier_info.get('contact_person', 'Ch∆∞a c·∫≠p nh·∫≠t')
                data_row = [
                    i,
                    ingredient_info['name'].upper(),
                    f"{day_date.strftime('%d/%m/%Y')}\n6:00-7:00",
                    f"{ingredient_info['weight_kg']} kg",
                    supplier_name,
                    f"{phone}\n{address[:30]}...",
                    contact_person,
                    '',  # ƒê·ªÉ tr·ªëng S·ªê CH·ª®NG T·ª™/S·ªê HO√Å ƒê∆†N
                    supplier_info.get('food_safety_cert', ''),
                    "",
                    '‚úì',
                    '',
                    '‚úì',
                    '',
                    ""
                ]
                for j, value in enumerate(data_row, 1):
                    cell = ws1.cell(row=row_num, column=j, value=value)
                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    cell.border = thin_border
                    if j == 1:
                        cell.font = Font(bold=True, color="0066CC")
                        cell.fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
                    elif j == 2:
                        cell.font = Font(bold=True, size=10)
                        cell.fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
                    elif j in [11, 13] and value == '‚úì':
                        cell.font = Font(bold=True, size=12, color="00AA00")
                        cell.fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")
                    elif j == 5:
                        cell.font = Font(bold=True, color="CC6600")
            # Th·ªëng k√™
            stats_row = len(fresh_ingredients) + 12
            ws1[f'A{stats_row}'] = "TH·ªêNG K√ä T·ªîNG QUAN:"
            ws1[f'A{stats_row}'].font = Font(bold=True, size=11, color="0066CC")
            ws1[f'A{stats_row}'].fill = PatternFill(start_color="E6F3FF", end_color="E6F3FF", fill_type="solid")
            total_weight = sum(item['weight_kg'] for item in fresh_ingredients)
            total_items = len(fresh_ingredients)
            stats_info = [
                f"‚Ä¢ T·ªïng s·ªë lo·∫°i th·ª±c ph·∫©m t∆∞∆°i: {total_items} lo·∫°i",
                f"‚Ä¢ T·ªïng kh·ªëi l∆∞·ª£ng ∆∞·ªõc t√≠nh: {total_weight:.1f} kg",
                f"‚Ä¢ S·ªë h·ªçc sinh ph·ª•c v·ª•: {students_today} em",
                f"‚Ä¢ Kh·ªëi l∆∞·ª£ng trung b√¨nh/h·ªçc sinh: {(total_weight/students_today):.2f} kg/em/ng√†y" if students_today else "‚Ä¢ Kh·ªëi l∆∞·ª£ng trung b√¨nh/h·ªçc sinh: N/A"
            ]
            for i, stat in enumerate(stats_info, 1):
                ws1[f'A{stats_row + i}'] = stat
                ws1[f'A{stats_row + i}'].font = Font(size=10)
            note_row = stats_row + 6
            ws1[f'A{note_row}'] = "GHI CH√ö QUAN TR·ªåNG:"
            ws1[f'A{note_row}'].font = Font(bold=True, size=11, color="FF0000")
            notes = [
                "‚Ä¢ Ki·ªÉm tra nhi·ªát ƒë·ªô b·∫£o qu·∫£n: Th·ª±c ph·∫©m t∆∞∆°i <4¬∞C, ƒë√¥ng l·∫°nh <-18¬∞C",
                "‚Ä¢ Th·ªùi gian s·ª≠ d·ª•ng: Th·ª±c ph·∫©m t∆∞∆°i trong ng√†y, ƒë√¥ng l·∫°nh theo h·∫°n s·ª≠ d·ª•ng",  
                "‚Ä¢ X√©t nghi·ªám nhanh: ∆Øu ti√™n th·ª±c ph·∫©m c√≥ ngu·ªìn g·ªëc kh√¥ng r√µ r√†ng",
                "‚Ä¢ B√°o c√°o ngay n·∫øu ph√°t hi·ªán b·∫•t th∆∞·ªùng v·ªÅ m√†u s·∫Øc, m√πi v·ªã, bao b√¨"
            ]
            for i, note in enumerate(notes, 1):
                ws1[f'A{note_row + i}'] = note
                ws1[f'A{note_row + i}'].font = Font(size=9, color="CC0000")
            signature_row = note_row + 7
            signature_data = [
                (signature_row, 'D', "B·∫æP TR∆Ø·ªûNG", 'K', "HI·ªÜU TR∆Ø·ªûNG"),
                (signature_row + 1, 'D', "(K√Ω, ghi r√µ h·ªç t√™n)", 'K', "(K√Ω, ghi r√µ h·ªç t√™n)"),
                (signature_row + 5, 'D', "Ho√†ng Thanh Tu·∫•n", 'K', "Nguy·ªÖn Th·ªã V√¢n"),
                (signature_row + 6, 'D', f"Ng√†y {day_date.day}/{day_date.month}/{day_date.year}",
                 'K',
                 f"Ng√†y {day_date.day}/{day_date.month}/{day_date.year}")
            ]
            for row, col_d, text_d, col_k, text_k in signature_data:
                ws1[f'{col_d}{row}'] = text_d
                ws1[f'{col_k}{row}'] = text_k
                for col, text in [(col_d, text_d), (col_k, text_k)]:
                    cell = ws1[f'{col}{row}']
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    if row == signature_row:
                        cell.font = Font(bold=True, size=12, color="0066CC")
                        cell.fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
                    elif row == signature_row + 1:
                        cell.font = Font(italic=True, size=9)
                    elif row == signature_row + 5:
                        cell.font = Font(bold=True, size=11)
                    else:
                        cell.font = Font(size=9)
            
            # Page setup cho in A4
            ws1.page_setup.orientation = ws1.ORIENTATION_LANDSCAPE
            ws1.page_setup.paperSize = ws1.PAPERSIZE_A4
            ws1.page_setup.fitToPage = True
            ws1.page_setup.fitToWidth = 1
            ws1.page_setup.fitToHeight = 0
            ws1.print_options.horizontalCentered = True
            ws1.page_margins.left = 0.5
            ws1.page_margins.right = 0.5
            ws1.page_margins.top = 0.75
            ws1.page_margins.bottom = 0.75
        
        if 'Sheet' in wb1.sheetnames:
            wb1.remove(wb1['Sheet'])
        file1_buffer = BytesIO()
        wb1.save(file1_buffer)
        file1_buffer.seek(0)
        zipf.writestr(f"B∆∞·ªõc 1.1 - Ti·∫øp nh·∫≠n th·ª±c ph·∫©m t∆∞∆°i - Tu·∫ßn {week_number}.xlsx", file1_buffer.read())
        
        # B∆Ø·ªöC 1.2: Ti·∫øp nh·∫≠n th·ª±c ph·∫©m kh√¥ - Format chuy√™n nghi·ªáp 

        # B∆Ø·ªöC 1.2: Ti·∫øp nh·∫≠n th·ª±c ph·∫©m kh√¥ - m·ªói ng√†y 1 sheet, ch·ªâ t·∫°o wb2 1 l·∫ßn
        wb2 = Workbook()
        days_vn = ["Th·ª© 2", "Th·ª© 3", "Th·ª© 4", "Th·ª© 5", "Th·ª© 6", "Th·ª© 7"]
        for day_offset in range(6):
            day_date = week_start + timedelta(days=day_offset)
            day_key = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat'][day_offset]
            safe_date = day_date.strftime('%d-%m')
            sheet_title = f"{days_vn[day_offset]} ({safe_date})"
            if day_offset == 0:
                ws2 = wb2.active
                ws2.title = sheet_title
            else:
                ws2 = wb2.create_sheet(title=sheet_title)
            # L·∫•y menu ng√†y
            menu_today = menu_data.get(day_key, {})
            # L·∫•y s·ªë h·ªçc sinh c√≥ m·∫∑t ng√†y n√†y
            students_today = daily_attendance.get(day_key, 0)
            # T√≠nh nguy√™n li·ªáu th·ª±c t·∫ø cho ng√†y n√†y
            daily_ingredients = {}
            for meal in menu_today.values():
                if not meal: continue
                for dish_name in [d.strip() for d in meal.split(',') if d.strip()]:
                    dish = Dish.query.filter_by(name=dish_name).first()
                    if dish:
                        for di in dish.ingredients:
                            key = (di.product.name, di.unit, di.product.category, di.product.supplier)
                            qty = di.quantity * students_today
                            if key not in daily_ingredients:
                                daily_ingredients[key] = {'total_qty': 0, 'unit': di.unit, 'category': di.product.category, 'supplier': di.product.supplier, 'product': di.product}
                            daily_ingredients[key]['total_qty'] += qty
            # Ph√¢n lo·∫°i kh√¥
            dry_ingredients = []
            for (name, unit, category, supplier), info in daily_ingredients.items():
                cat = (category or '').lower()
                if cat == 'dry' or 'kh√¥' in cat or 'gia v·ªã' in cat or 'b·ªôt' in cat or 'g·∫°o' in cat or 'ƒë∆∞·ªùng' in cat:
                    if hasattr(supplier, 'name'):
                        supplier_name = supplier.name
                    elif isinstance(supplier, str):
                        supplier_name = supplier
                    else:
                        supplier_name = ''
                    dry_ingredients.append({
                        'name': name,
                        'weight_kg': round(convert_to_kg(info['total_qty'], info['unit']), 2),
                        'unit': unit,
                        'category': category,
                        'supplier': supplier_name,
                        'supplier_info': supplier_dict.get(supplier_name, {}),
                    })
            # --- Ghi d·ªØ li·ªáu v√† style sheet nh∆∞ B∆∞·ªõc 1.1 ---
            ws2['A1'] = "T√äN C∆† S·ªû: MNƒêL C√¢y Nh·ªè"
            ws2['A1'].font = Font(bold=True, size=12)
            ws2['A1'].fill = PatternFill(start_color="FFE6CC", end_color="FFE6CC", fill_type="solid")
            ws2.merge_cells('A1:P1')
            ws2['D2'] = "BI·ªÇU M·∫™U KI·ªÇM TRA TH·ª∞C PH·∫®M KH√î V√Ä BAO G√ìI"
            ws2['D2'].font = Font(bold=True, size=14, color="FF0000")
            ws2['D2'].alignment = Alignment(horizontal='center', vertical='center')
            ws2.merge_cells('D2:L2')
            ws2['N2'] = "S·ªë: 1246/Qƒê - B·ªô Y T·∫ø"
            ws2['N2'].font = Font(bold=True, size=10)
            ws2['N2'].fill = PatternFill(start_color="FFCCCC", end_color="FFCCCC", fill_type="solid")
            info_data2 = [
                (3, 'A', f"Ng∆∞·ªùi ki·ªÉm tra: Nguy·ªÖn Th·ªã V√¢n", 'N', "M·∫´u s·ªë 1.2"),
                (4, 'A', f"Ng√†y ki·ªÉm tra: {day_date.strftime('%d/%m/%Y')} - {days_vn[day_offset]}", 'N', f"S·ªë h·ªçc sinh: {students_today}"),
                (5, 'A', "ƒê·ªãa ƒëi·ªÉm: Kho th·ª±c ph·∫©m kh√¥ - MNƒêL C√¢y Nh·ªè", 'N', "")
            ]
            for row, col_a, text_a, col_n, text_n in info_data2:
                ws2[f'{col_a}{row}'] = text_a
                ws2[f'{col_a}{row}'].font = Font(bold=True, size=10)
                ws2[f'{col_n}{row}'] = text_n
                ws2[f'{col_n}{row}'].font = Font(bold=True, size=10)
                ws2[f'{col_n}{row}'].fill = PatternFill(start_color="F0F8FF", end_color="F0F8FF", fill_type="solid")
            ws2['A7'] = "PH·∫¶N II: TH·ª∞C PH·∫®M KH√î, BAO G√ìI S·∫¥N V√Ä PH·ª§ GIA TH·ª∞C PH·∫®M"
            ws2['A7'].font = Font(bold=True, size=12, color="FF6600")
            ws2['A7'].fill = PatternFill(start_color="FFF2E6", end_color="FFF2E6", fill_type="solid")
            ws2.merge_cells('A7:M7')
            ws2['N7'] = "B∆Ø·ªöC 1.2"
            ws2['N7'].font = Font(bold=True, size=12, color="FF0000")
            ws2['N7'].fill = PatternFill(start_color="FFEEEE", end_color="FFEEEE", fill_type="solid")
            headers2_main = [
                'STT', 'T√äN TH·ª∞C PH·∫®M', 'T√äN C∆† S·ªû S·∫¢N XU·∫§T', 
                'ƒê·ªäA CH·ªà S·∫¢N XU·∫§T', 'TH·ªúI GIAN NH·∫¨P\n(Ng√†y/Gi·ªù)', 'KH·ªêI L∆Ø·ª¢NG (KG/L√çT)', 'N∆†I CUNG C·∫§P', '', '',
                'H·∫†N S·ª¨ D·ª§NG', 'ƒêI·ªÄU KI·ªÜN B·∫¢O QU·∫¢N', 'CH·ª®NG T·ª™, HO√Å ƒê∆†N', 'KI·ªÇM TRA C·∫¢M QUAN', '', 'BI·ªÜN PH√ÅP X·ª¨ L√ù / GHI CH√ö'
            ]
            for i, header in enumerate(headers2_main, 1):
                cell = ws2.cell(row=8, column=i, value=header)
                cell.font = Font(bold=True, size=9, color="FFFFFF")
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                cell.fill = PatternFill(start_color="E67E22", end_color="E67E22", fill_type="solid")
                cell.border = thick_border
            sub_headers2 = [
                '', '', '', '', '', '', '', 'T√™n c∆° s·ªü', '', '', '', '', 
                'ƒê·∫°t', 'Kh√¥ng ƒë·∫°t', ''
            ]
            for i, header in enumerate(sub_headers2, 1):
                cell = ws2.cell(row=9, column=i, value=header)
                cell.font = Font(bold=True, size=8)
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.fill = PatternFill(start_color="F8C471", end_color="F8C471", fill_type="solid")
                cell.border = thin_border
            ws2.merge_cells('G8:I8')  # N∆°i cung c·∫•p
            ws2.merge_cells('M8:N8')  # Ki·ªÉm tra c·∫£m quan
            for i in range(1, 16):
                cell = ws2.cell(row=10, column=i, value=i)
                cell.font = Font(bold=True, size=8)
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.fill = PatternFill(start_color="FADBD8", end_color="FADBD8", fill_type="solid")
                cell.border = thin_border
            # Ghi d·ªØ li·ªáu th·ª±c ph·∫©m kh√¥ t·ª´ng ng√†y
            for i, ingredient_info in enumerate(dry_ingredients[:25], 1):
                row_num = 10 + i
                supplier_info = ingredient_info.get('supplier_info', {})
                supplier_name = ingredient_info.get('supplier', '') or 'Si√™u th·ªã Co.opmart'
                phone = supplier_info.get('phone', '0902.xxx.xxx')
                address = supplier_info.get('address', 'ƒê√† L·∫°t')
                contact_person = supplier_info.get('contact_person', 'Ch∆∞a c·∫≠p nh·∫≠t')
                expiry_date = (day_date + timedelta(days=180)).strftime('%d/%m/%Y')
                # ƒê·ªìng b·ªô v·ªõi headers2_main m·ªõi
                data_row2 = [
                    i,  # STT
                    ingredient_info['name'].upper(),  # T√äN TH·ª∞C PH·∫®M
                    supplier_name,  # T√äN C∆† S·ªû S·∫¢N XU·∫§T (gi·∫£ ƒë·ªãnh l√† supplier)
                    address,  # ƒê·ªäA CH·ªà S·∫¢N XU·∫§T (gi·∫£ ƒë·ªãnh l√† address supplier)
                    f"{day_date.strftime('%d/%m/%Y')}\n8:00-9:00",  # TH·ªúI GIAN NH·∫¨P
                    f"{ingredient_info['weight_kg']} kg",  # KH·ªêI L∆Ø·ª¢NG
                    supplier_name,  # N∆†I CUNG C·∫§P
                    '',  # c·ªôt ph·ª• (merge)
                    '',  # c·ªôt ph·ª• (merge)
                    "c√≤n HDS",  # H·∫†N S·ª¨ D·ª§NG
                    "Kh√¥ r√°o, tho√°ng m√°t\n<25¬∞C",  # ƒêI·ªÄU KI·ªÜN B·∫¢O QU·∫¢N
                    '',  # CH·ª®NG T·ª™, HO√Å ƒê∆†N (ch∆∞a c√≥)
                    '‚úì',  # KI·ªÇM TRA C·∫¢M QUAN (ƒê·∫°t)
                    '',  # Kh√¥ng ƒë·∫°t
                    '',  # BI·ªÜN PH√ÅP X·ª¨ L√ù / GHI CH√ö
                ]
                for j, value in enumerate(data_row2, 1):
                    cell = ws2.cell(row=row_num, column=j, value=value)
                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    cell.border = thin_border
                    if j == 1:
                        cell.font = Font(bold=True, color="E67E22")
                        cell.fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
                    elif j == 2:
                        cell.font = Font(bold=True, size=10)
                        cell.fill = PatternFill(start_color="FEF9E7", end_color="FEF9E7", fill_type="solid")
                    elif j == 13 and value == '‚úì':
                        cell.font = Font(bold=True, size=12, color="27AE60")
                        cell.fill = PatternFill(start_color="E8F5E8", end_color="E8F5E8", fill_type="solid")
                    elif j == 4:
                        cell.font = Font(bold=True, color="D35400")
                    elif j == 11:
                        cell.font = Font(bold=True, color="8E44AD")
            # Th·ªëng k√™
            stats_row2 = len(dry_ingredients) + 12
            ws2[f'A{stats_row2}'] = "TH·ªêNG K√ä TH·ª∞C PH·∫®M KH√î:"
            ws2[f'A{stats_row2}'].font = Font(bold=True, size=11, color="E67E22")
            ws2[f'A{stats_row2}'].fill = PatternFill(start_color="FFF2E6", end_color="FFF2E6", fill_type="solid")
            total_weight2 = sum(item['weight_kg'] for item in dry_ingredients)
            total_items2 = len(dry_ingredients)
            stats_info2 = [
                f"‚Ä¢ T·ªïng s·ªë lo·∫°i th·ª±c ph·∫©m kh√¥: {total_items2} lo·∫°i",
                f"‚Ä¢ T·ªïng kh·ªëi l∆∞·ª£ng ∆∞·ªõc t√≠nh: {total_weight2:.1f} kg",
                f"‚Ä¢ S·ªë h·ªçc sinh ph·ª•c v·ª•: {students_today} em",
                f"‚Ä¢ Kh·ªëi l∆∞·ª£ng trung b√¨nh/h·ªçc sinh: {(total_weight2/students_today):.2f} kg/em/ng√†y" if students_today else "‚Ä¢ Kh·ªëi l∆∞·ª£ng trung b√¨nh/h·ªçc sinh: N/A"
            ]
            for i, stat in enumerate(stats_info2, 1):
                ws2[f'A{stats_row2 + i}'] = stat
                ws2[f'A{stats_row2 + i}'].font = Font(size=10)
            note_row2 = stats_row2 + 6
            ws2[f'A{note_row2}'] = "GHI CH√ö QUAN TR·ªåNG:"
            ws2[f'A{note_row2}'].font = Font(bold=True, size=11, color="FF0000")
            notes2 = [
                "‚Ä¢ B·∫£o qu·∫£n n∆°i kh√¥ r√°o, tho√°ng m√°t, tr√°nh √°nh n·∫Øng tr·ª±c ti·∫øp",
                "‚Ä¢ Ki·ªÉm tra h·∫°n s·ª≠ d·ª•ng, bao b√¨ nguy√™n v·∫πn tr∆∞·ªõc khi nh·∫≠p kho",
                "‚Ä¢ S·ª≠ d·ª•ng theo nguy√™n t·∫Øc FIFO (nh·∫≠p tr∆∞·ªõc xu·∫•t tr∆∞·ªõc)",
                "‚Ä¢ B√°o c√°o ngay n·∫øu ph√°t hi·ªán b·∫•t th∆∞·ªùng v·ªÅ m√†u s·∫Øc, m√πi v·ªã, bao b√¨"
            ]
            for i, note in enumerate(notes2, 1):
                ws2[f'A{note_row2 + i}'] = note
                ws2[f'A{note_row2 + i}'].font = Font(size=9, color="CC0000")
            signature_row2 = note_row2 + 7
            signature_data2 = [
                (signature_row2, 'D', "TH·ª¶ KHO", 'K', "HI·ªÜU TR∆Ø·ªûNG"),
                (signature_row2 + 1, 'D', "(K√Ω, ghi r√µ h·ªç t√™n)", 'K', "(K√Ω, ghi r√µ h·ªç t√™n)"),
                (signature_row2 + 5, 'D', "Ho√†ng Thanh Tu·∫•n", 'K', "Nguy·ªÖn Th·ªã V√¢n"),
                (signature_row2 + 6, 'D', f"Ng√†y {day_date.day}/{day_date.month}/{day_date.year}", 'K', f"Ng√†y {day_date.day}/{day_date.month}/{day_date.year}")
            ]
            for row, col_d, text_d, col_k, text_k in signature_data2:
                ws2[f'{col_d}{row}'] = text_d
                ws2[f'{col_k}{row}'] = text_k
                for col, text in [(col_d, text_d), (col_k, text_k)]:
                    cell = ws2[f'{col}{row}']
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    if row == signature_row2:
                        cell.font = Font(bold=True, size=12, color="E67E22")
                        cell.fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
                    elif row == signature_row2 + 1:
                        cell.font = Font(italic=True, size=9)
                    elif row == signature_row2 + 5:
                        cell.font = Font(bold=True, size=11)
                    else:
                        cell.font = Font(size=9)
            
            # Page setup cho in A4
            ws2.page_setup.orientation = ws2.ORIENTATION_LANDSCAPE
            ws2.page_setup.paperSize = ws2.PAPERSIZE_A4
            ws2.page_setup.fitToPage = True
            ws2.page_setup.fitToWidth = 1
            ws2.page_setup.fitToHeight = 0
            ws2.print_options.horizontalCentered = True
            ws2.page_margins.left = 0.5
            ws2.page_margins.right = 0.5
            ws2.page_margins.top = 0.75
            ws2.page_margins.bottom = 0.75
        
        if 'Sheet' in wb2.sheetnames:
            wb2.remove(wb2['Sheet'])
        file2_buffer = BytesIO()
        wb2.save(file2_buffer)
        file2_buffer.seek(0)
        zipf.writestr(f"B∆∞·ªõc 1.2 - Ti·∫øp nh·∫≠n th·ª±c ph·∫©m kh√¥ - Tu·∫ßn {week_number}.xlsx", file2_buffer.read())

        # B∆Ø·ªöC 2: Ki·ªÉm tra khi ch·∫ø bi·∫øn th·ª©c ƒÉn - m·ªói ng√†y 1 sheet
        wb3 = Workbook()
        days_vn = ["Th·ª© 2", "Th·ª© 3", "Th·ª© 4", "Th·ª© 5", "Th·ª© 6", "Th·ª© 7"]
        for day_offset in range(6):
            day_date = week_start + timedelta(days=day_offset)
            day_key = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat'][day_offset]
            safe_date = day_date.strftime('%d-%m')
            sheet_title = f"{days_vn[day_offset]} ({safe_date})"
            if day_offset == 0:
                ws3 = wb3.active
                ws3.title = sheet_title
            else:
                ws3 = wb3.create_sheet(title=sheet_title)
            # L·∫•y s·ªë h·ªçc sinh c√≥ m·∫∑t ng√†y n√†y
            students_today = daily_attendance.get(day_key, 0)
            # Header ch√≠nh t∆∞∆°ng t·ª± c√°c b∆∞·ªõc tr∆∞·ªõc
            ws3['A1'] = "T√äN C∆† S·ªû: MNƒêL C√¢y Nh·ªè"
            ws3['A1'].font = Font(bold=True, size=12)
            ws3['A1'].fill = PatternFill(start_color="E6FFE6", end_color="E6FFE6", fill_type="solid")
            ws3.merge_cells('A1:O1')
            ws3['D2'] = "BI·ªÇU M·∫™U KI·ªÇM TRA KHI CH·∫æ BI·∫æN TH·ª®C ƒÇN"
            ws3['D2'].font = Font(bold=True, size=14, color="006600")
            ws3['D2'].alignment = Alignment(horizontal='center', vertical='center')
            ws3.merge_cells('D2:K2')
            ws3['M2'] = "S·ªë: 1246/Qƒê - B·ªô Y T·∫ø"
            ws3['M2'].font = Font(bold=True, size=10)
            ws3['M2'].fill = PatternFill(start_color="CCFFCC", end_color="CCFFCC", fill_type="solid")
            info_data3 = [
                (3, 'A', f"Ng∆∞·ªùi ki·ªÉm tra: Nguy·ªÖn Th·ªã V√¢n", 'M', "M·∫´u s·ªë 2.0"),
                (4, 'A', f"Ng√†y ki·ªÉm tra: {day_date.strftime('%d/%m/%Y')} - {days_vn[day_offset]}", 'M', f"S·ªë h·ªçc sinh: {students_today}"),
                (5, 'A', "ƒê·ªãa ƒëi·ªÉm: B·∫øp ch·∫ø bi·∫øn - MNƒêL C√¢y Nh·ªè", 'M', "")
            ]
            for row, col_a, text_a, col_m, text_m in info_data3:
                ws3[f'{col_a}{row}'] = text_a
                ws3[f'{col_a}{row}'].font = Font(bold=True, size=10)
                ws3[f'{col_m}{row}'] = text_m
                ws3[f'{col_m}{row}'].font = Font(bold=True, size=10)
                ws3[f'{col_m}{row}'].fill = PatternFill(start_color="F0F8FF", end_color="F0F8FF", fill_type="solid")
            ws3['A7'] = "PH·∫¶N II: KI·ªÇM TRA QUY TR√åNH CH·∫æ BI·∫æN TH·ª®C ƒÇN"
            ws3['A7'].font = Font(bold=True, size=12, color="8B0000")
            ws3['A7'].fill = PatternFill(start_color="FFE6E6", end_color="FFE6E6", fill_type="solid")
            ws3.merge_cells('A7:L7')
            ws3['M7'] = "B∆Ø·ªöC 2"
            ws3['M7'].font = Font(bold=True, size=12, color="FF0000")
            ws3['M7'].fill = PatternFill(start_color="FFEEEE", end_color="FFEEEE", fill_type="solid")
            headers3_main = [
                'STT', 'CA/B·ªÆA ƒÇN', 'T√äN M√ìN ƒÇN', 'NGUY√äN LI·ªÜU CH√çNH', 'S·ªê SU·∫§T\n(ph·∫ßn)', 
                'TH·ªúI GIAN S∆† CH·∫æ XONG\n(ng√†y, gi·ªù)', 'TH·ªúI GIAN CH·∫æ BI·∫æN XONG\n(ng√†y, gi·ªù)', 'KI·ªÇM TRA V·ªÜ SINH', '', '',
                'KI·ªÇM TRA C·∫¢M QUAN TH·ª®C ƒÇN', '', 'BI·ªÜN PH√ÅP X·ª¨ L√ù\nGHI CH√ö'
            ]
            for i, header in enumerate(headers3_main, 1):
                cell = ws3.cell(row=8, column=i, value=header)
                cell.font = Font(bold=True, size=9, color="FFFFFF")
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                cell.fill = PatternFill(start_color="8B0000", end_color="8B0000", fill_type="solid")
                cell.border = thick_border
            sub_headers3 = [
                '', '', '', '', '', '', '', 'Ng∆∞·ªùi tham gia\n ch·∫ø bi·∫øn', 'Trang thi·∫øt b·ªã\n d·ª•ng c·ª•', 'Khu v·ª±c ch·∫ø bi·∫øn\n v√† ph·ª• tr·ª£',
                'ƒê·∫°t', 'Kh√¥ng ƒë·∫°t', ''
            ]
            for i, header in enumerate(sub_headers3, 1):
                cell = ws3.cell(row=9, column=i, value=header)
                cell.font = Font(bold=True, size=8)
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.fill = PatternFill(start_color="CD5C5C", end_color="CD5C5C", fill_type="solid")
                cell.border = thin_border

            # ƒê·∫∑t ƒë·ªô r·ªông v√† cƒÉn gi·ªØa cho H9, I9, J9
            ws3.column_dimensions['H'].width = 18
            ws3.column_dimensions['I'].width = 18
            ws3.column_dimensions['J'].width = 18
            for col in ['H', 'I', 'J']:
                cell = ws3[f'{col}9']
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            ws3.merge_cells('H8:J8')  # V·ªá sinh, C·∫£m quan
            ws3.merge_cells('K8:L8')  # ƒê√°nh gi√° ph·ª•c v·ª•
            for i in range(1, 14):
                cell = ws3.cell(row=10, column=i, value=i)
                cell.font = Font(bold=True, size=8)
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.fill = PatternFill(start_color="F0F8FF", end_color="F0F8FF", fill_type="solid")
                cell.border = thin_border
            # Ghi d·ªØ li·ªáu m√≥n ƒÉn t·ª´ng ng√†y (t∆∞∆°ng t·ª± logic c≈©, ch·ªâ cho ng√†y n√†y)
            row_num = 11
            stt = 1
            meal_times = {
                'morning': 'B·ªØa s√°ng',
                'snack': 'ƒÇn ph·ª• s√°ng',
                'dessert': 'Tr√°ng mi·ªáng',
                'lunch': 'B·ªØa tr∆∞a',
                'afternoon': 'ƒÇn ph·ª• chi·ªÅu',
                'lateafternoon': 'B·ªØa x·∫ø',
            }
            # Gi·ªù chu·∫©n cho t·ª´ng ca (gi·ªù_s∆°_ch·∫ø, gi·ªù_ch·∫ø_bi·∫øn)
            meal_time_hours = {
                'morning':   ('07:00', '07:25'),
                'snack':     ('09:00', '10:00'),
                'dessert':   ('09:00', '10:00'),
                'lunch':     ('09:00', '10:00'),
                'afternoon': ('09:00', '10:00'),
                'lateafternoon': ('14:00', '14:25')
            }
            # L·∫∑p ƒë·ªß 5 ca d·ª±a tr√™n meal_times
            for meal_key, meal_name in meal_times.items():
                dishes = []
                if menu_data[day_key].get(meal_key):
                    dishes = [d.strip() for d in menu_data[day_key][meal_key].split(',') if d.strip()]
                # Kh√¥ng l·ªçc, l·∫•y to√†n b·ªô m√≥n trong menu t·ª´ng b·ªØa
                if dishes:
                    dish_names = ', '.join([dish.title() for dish in dishes])
                    all_ingredients = set()
                    for dish in dishes:
                        dish_obj = Dish.query.filter_by(name=dish).first()
                        if dish_obj and dish_obj.ingredients:
                            for di in dish_obj.ingredients:
                                all_ingredients.add(di.product.name)
                    main_ingredients = ', '.join(sorted(all_ingredients))
                else:
                    dish_names = ''
                    main_ingredients = ''
                # T·∫°o gi√° tr·ªã th·ªùi gian s∆° ch·∫ø xong, ch·∫ø bi·∫øn xong
                date_str = day_date.strftime('%d/%m/%Y')
                time_so_che = meal_time_hours[meal_key][0] if meal_key in meal_time_hours else ''
                time_che_bien = meal_time_hours[meal_key][1] if meal_key in meal_time_hours else ''
                so_che_str = f"{date_str} {time_so_che}" if time_so_che else ''
                che_bien_str = f"{date_str} {time_che_bien}" if time_che_bien else ''
                data_row3 = [
                    stt,  # STT
                    meal_name,  # CA/B·ªÆA ƒÇN ch·ªâ t√™n b·ªØa
                    dish_names,  # T√äN M√ìN ƒÇN (danh s√°ch m√≥n)
                    main_ingredients,  # NGUY√äN LI·ªÜU CH√çNH (to√†n b·ªô nguy√™n li·ªáu c√°c m√≥n)
                    students_today,  # S·ªê SU·∫§T (ph·∫ßn)
                    so_che_str,  # TH·ªúI GIAN S∆† CH·∫æ XONG (ng√†y, gi·ªù)
                    che_bien_str,  # TH·ªúI GIAN CH·∫æ BI·∫æN XONG (ng√†y, gi·ªù)
                    "Trang ph·ª•c g·ªçn g√†ng, v·ªá sinh c√° nh√¢n s·∫°ch s·∫Ω",  # Ng∆∞·ªùi tham gia ch·∫ø bi·∫øn
                    "ƒê·∫£m b·∫£o v·ªá sinh",  # Trang thi·∫øt b·ªã d·ª•ng c·ª•
                    "ƒê·∫£m b·∫£o v·ªá sinh",  # Khu v·ª±c ch·∫ø bi·∫øn v√† ph·ª• tr·ª£
                    "",  # KI·ªÇM TRA C·∫¢M QUAN TH·ª®C ƒÇN - ƒê·∫°t
                    "",  # KI·ªÇM TRA C·∫¢M QUAN TH·ª®C ƒÇN - Kh√¥ng ƒë·∫°t
                    ""   # BI·ªÜN PH√ÅP X·ª¨ L√ù GHI CH√ö
                ]
                for j, value in enumerate(data_row3, 1):
                    cell = ws3.cell(row=row_num, column=j, value=value)
                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    cell.border = thin_border
                    if j == 1:
                        cell.font = Font(bold=True, color="006600")
                        cell.fill = PatternFill(start_color="F0FFF0", end_color="F0FFF0", fill_type="solid")
                    elif j == 3:
                        cell.font = Font(bold=True, size=10)
                        cell.fill = PatternFill(start_color="F0FFF0", end_color="F0FFF0", fill_type="solid")
                row_num += 1
                stt += 1
            # Th·ªëng k√™ ph·ª•c v·ª•
            stats_row3 = row_num + 2
            ws3[f'A{stats_row3}'] = "TH·ªêNG K√ä PH·ª§C V·ª§ TH·ª®C ƒÇN:"
            ws3[f'A{stats_row3}'].font = Font(bold=True, size=11, color="006600")
            ws3[f'A{stats_row3}'].fill = PatternFill(start_color="E6FFE6", end_color="E6FFE6", fill_type="solid")
            total_servings = stt - 1
            total_portions = total_servings * students_today
            stats_info3 = [
                f"‚Ä¢ T·ªïng s·ªë l·∫ßn ph·ª•c v·ª•: {total_servings} l·∫ßn",
                f"‚Ä¢ T·ªïng s·ªë su·∫•t ƒÉn ph·ª•c v·ª•: {total_portions} su·∫•t",
                f"‚Ä¢ Trung b√¨nh su·∫•t/l·∫ßn: {(total_portions/total_servings):.1f} su·∫•t/l·∫ßn" if total_servings else "‚Ä¢ Trung b√¨nh su·∫•t/l·∫ßn: N/A",
                f"‚Ä¢ Th·ªùi gian trung b√¨nh t·ª´ ch·∫ø bi·∫øn xong ƒë·∫øn ph·ª•c v·ª•: <30 ph√∫t"
            ]
            for i, stat in enumerate(stats_info3, 1):
                ws3[f'A{stats_row3 + i}'] = stat
                ws3[f'A{stats_row3 + i}'].font = Font(size=10)
            # Nguy√™n t·∫Øc b·∫£o qu·∫£n v√† ph·ª•c v·ª•
            principles_row = stats_row3 + 6
            ws3[f'A{principles_row}'] = "NGUY√äN T·∫ÆC B·∫¢O QU·∫¢N V√Ä PH·ª§C V·ª§ AN TO√ÄN:"
            ws3[f'A{principles_row}'].font = Font(bold=True, size=11, color="004000")
            principles_notes = [
                "‚Ä¢ Th·ªùi gian: T·ª´ ch·∫ø bi·∫øn xong ƒë·∫øn ph·ª•c v·ª• kh√¥ng qu√° 2 gi·ªù",
                "‚Ä¢ Nhi·ªát ƒë·ªô: M√≥n n√≥ng >60¬∞C, m√≥n l·∫°nh <10¬∞C khi ph·ª•c v·ª•",
                "‚Ä¢ Thi·∫øt b·ªã: S·ª≠ d·ª•ng t·ªß gi·ªØ nhi·ªát, n·ªìi c∆°m ƒëi·ªán, b√¨nh gi·ªØ nhi·ªát",
                "‚Ä¢ V·ªá sinh: Kh·ª≠ tr√πng d·ª•ng c·ª• tr∆∞·ªõc m·ªói b·ªØa ƒÉn",
                "‚Ä¢ Ki·ªÉm tra: Nhi·ªát ƒë·ªô th·ª©c ƒÉn tr∆∞·ªõc khi ph·ª•c v·ª• cho tr·∫ª"
            ]
            for i, note in enumerate(principles_notes, 1):
                ws3[f'A{principles_row + i}'] = note
                ws3[f'A{principles_row + i}'].font = Font(size=9, color="004000")
            # Ch·ªØ k√Ω
            signature_row3 = principles_row + 8
            signature_data3 = [
                (signature_row3,     'D', "B·∫æP TR∆Ø·ªûNG",  'H', "NV. Y T·∫æ",   'K', "HI·ªÜU TR∆Ø·ªûNG"),
                (signature_row3 + 1, 'D', "(K√Ω, ghi r√µ h·ªç t√™n)",'H', "(K√Ω, ghi r√µ h·ªç t√™n)", 'K', "(K√Ω, ghi r√µ h·ªç t√™n)"),
                (signature_row3 + 5, 'D', "Ho√†ng Thanh Tu·∫•n",'H', "(K√Ω, ghi r√µ h·ªç t√™n)",  'K', "Nguy·ªÖn Th·ªã V√¢n")
            ]
                # --- Ghi ch·ªØ k√Ω ---
            for item in signature_data3:
                    
                    row, col_d, text_d, col_h, text_h, col_k, text_k = item
                    ws3[f'{col_d}{row}'] = text_d
                    ws3[f'{col_h}{row}'] = text_h
                    ws3[f'{col_k}{row}'] = text_k
                    cols = [(col_d, text_d), (col_h, text_h), (col_k, text_k)]

                    for col, text in cols:
                        cell = ws3[f'{col}{row}']
                        cell.alignment = Alignment(horizontal='center', vertical='center')
                        if row == signature_row3:
                            cell.font = Font(bold=True, size=12, color="006600")
                            cell.fill = PatternFill(start_color="F0FFF0", end_color="F0FFF0", fill_type="solid")
                        elif row == signature_row3 + 1:
                            cell.font = Font(italic=True, size=9)
                        elif row == signature_row3 + 5:
                            cell.font = Font(bold=True, size=11)
                        else:
                            cell.font = Font(size=9)
            
            # Page setup cho in A4
            ws3.page_setup.orientation = ws3.ORIENTATION_LANDSCAPE
            ws3.page_setup.paperSize = ws3.PAPERSIZE_A4
            ws3.page_setup.fitToPage = True
            ws3.page_setup.fitToWidth = 1
            ws3.page_setup.fitToHeight = 0
            ws3.print_options.horizontalCentered = True
            ws3.page_margins.left = 0.5
            ws3.page_margins.right = 0.5
            ws3.page_margins.top = 0.75
            ws3.page_margins.bottom = 0.75
        
        if 'Sheet' in wb3.sheetnames:
            wb3.remove(wb3['Sheet'])

        file3_buffer = BytesIO()
        wb3.save(file3_buffer)
        file3_buffer.seek(0)
        zipf.writestr(f"B∆∞·ªõc 2 - Ki·ªÉm tra khi ch·∫ø bi·∫øn th·ª©c ƒÉn - Tu·∫ßn {week_number}.xlsx", file3_buffer.read())

        # B∆Ø·ªöC 3: Ki·ªÉm tra tr∆∞·ªõc khi ƒÉn - m·ªói ng√†y 1 sheet,
        wb4 = Workbook()
        for day_offset in range(6):
            day_date = week_start + timedelta(days=day_offset)
            day_key = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat'][day_offset]
            safe_date = day_date.strftime('%d-%m')
            sheet_title = f"{days_vn[day_offset]} ({safe_date})"
            if day_offset == 0:
                ws4 = wb4.active
                ws4.title = sheet_title
            else:
                ws4 = wb4.create_sheet(title=sheet_title)
            # L·∫•y s·ªë h·ªçc sinh c√≥ m·∫∑t ng√†y n√†y
            students_today = daily_attendance.get(day_key, 0)
            # Header ch√≠nh t∆∞∆°ng t·ª± ws3
            ws4['A1'] = "T√äN C∆† S·ªû: MNƒêL C√¢y Nh·ªè"
            ws4['A1'].font = Font(bold=True, size=12)
            ws4['A1'].fill = PatternFill(start_color="E6FFE6", end_color="E6FFE6", fill_type="solid")
            ws4.merge_cells('A1:O1')
            ws4['D2'] = "BI·ªÇU M·∫™U KI·ªÇM TRA TR∆Ø·ªöC KHI ƒÇN"
            ws4['D2'].font = Font(bold=True, size=14, color="006600")
            ws4['D2'].alignment = Alignment(horizontal='center', vertical='center')
            ws4.merge_cells('D2:I2')
            ws4['J2'] = "S·ªë: 1246/Qƒê - B·ªô Y T·∫ø"
            ws4['J2'].font = Font(bold=True, size=10)
            ws4['J2'].fill = PatternFill(start_color="CCFFCC", end_color="CCFFCC", fill_type="solid")
            info_data4 = [
                (3, 'A', f"Ng∆∞·ªùi ki·ªÉm tra: Nguy·ªÖn Th·ªã V√¢n", 'J', "M·∫´u s·ªë 3.0"),
                (4, 'A', f"Ng√†y ki·ªÉm tra: {day_date.strftime('%d/%m/%Y')} - {days_vn[day_offset]}", 'J', f"S·ªë h·ªçc sinh: {students_today}"),
                (5, 'A', "ƒê·ªãa ƒëi·ªÉm: Ph√≤ng ƒÉn - MNƒêL C√¢y Nh·ªè", 'J', "")
            ]
            for row, col_a, text_a, col_m, text_m in info_data4:
                ws4[f'{col_a}{row}'] = text_a
                ws4[f'{col_a}{row}'].font = Font(bold=True, size=10)
                ws4[f'{col_m}{row}'] = text_m
                ws4[f'{col_m}{row}'].font = Font(bold=True, size=10)
                ws4[f'{col_m}{row}'].fill = PatternFill(start_color="F0F8FF", end_color="F0F8FF", fill_type="solid")

            ws4['J7'] = "B∆Ø·ªöC 3"
            ws4['J7'].font = Font(bold=True, size=12, color="FF0000")
            ws4['J7'].fill = PatternFill(start_color="FFEEEE", end_color="FFEEEE", fill_type="solid")
            headers4_main = [
                'STT', 'CA/B·ªÆA ƒÇN', 'T√äN M√ìN ƒÇN', 'S·ªê SU·∫§T\n(ph·∫ßn)',
                'TH·ªúI GIAN CHIA M√ìN ƒÇN XONG\n(ng√†y, gi·ªù)', 'TH·ªúI GIAN B·∫ÆT ƒê·∫¶U ƒÇN\n(ng√†y, gi·ªù)',
                'D·ª§NG C·ª§ CHIA, CH·ª®A ƒê·ª∞NG\n, CHE ƒê·∫¨Y, B·∫¢O QU·∫¢N TH·ª®C ƒÇN',
                'KI·ªÇM TRA C·∫¢M QUAN TH·ª®C ƒÇN', '', 'BI·ªÜN PH√ÅP X·ª¨ L√ù\nGHI CH√ö'
            ]
            for i, header in enumerate(headers4_main, 1):
                cell = ws4.cell(row=8, column=i, value=header)
                cell.font = Font(bold=True, size=9, color="FFFFFF")
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                cell.fill = PatternFill(start_color="8B0000", end_color="8B0000", fill_type="solid")
                cell.border = thick_border
            sub_headers4 = [
                '', '', '', '', 
                '', '', 
                '', 
                'ƒê·∫°t', 'Kh√¥ng ƒë·∫°t', ''
            ]
            for i, header in enumerate(sub_headers4, 1):
                cell = ws4.cell(row=9, column=i, value=header)
                cell.font = Font(bold=True, size=8)
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.fill = PatternFill(start_color="CD5C5C", end_color="CD5C5C", fill_type="solid")
                cell.border = thin_border
            ws4.column_dimensions['G'].width = 25
            ws4.column_dimensions['H'].width = 18
            ws4.column_dimensions['I'].width = 18
            for col in ['G','H', 'I']:
                cell = ws4[f'{col}8']
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            ws4.merge_cells('H8:I8')
            for i in range(1, 11):
                cell = ws4.cell(row=10, column=i, value=i)
                cell.font = Font(bold=True, size=8)
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.fill = PatternFill(start_color="F0F8FF", end_color="F0F8FF", fill_type="solid")
                cell.border = thin_border
            # Ghi d·ªØ li·ªáu m√≥n ƒÉn t·ª´ng ng√†y (gi·ªëng ws3)
            row_num = 11
            stt = 1
            for meal_key, meal_name in meal_times.items():
                dishes = []
                if menu_data[day_key].get(meal_key):
                    dishes = [d.strip() for d in menu_data[day_key][meal_key].split(',') if d.strip()]
                if dishes:
                    dish_names = ', '.join([dish.title() for dish in dishes])
                else:
                    dish_names = ''
                date_str = day_date.strftime('%d/%m/%Y')
                time_chia_xong = '10:15'  # Gi·∫£ ƒë·ªãnh gi·ªù chia xong
                time_bat_dau_an = '10:30'  # Gi·∫£ ƒë·ªãnh gi·ªù b·∫Øt ƒë·∫ßu ƒÉn
                chia_xong_str = f"{date_str} {time_chia_xong}"
                bat_dau_an_str = f"{date_str} {time_bat_dau_an}"
                data_row4 = [
                    stt,  # STT
                    meal_name,  # CA/B·ªÆA ƒÇN
                    dish_names,  # T√äN M√ìN ƒÇN
                    students_today,  # S·ªê SU·∫§T
                    chia_xong_str,  # TH·ªúI GIAN CHIA M√ìN ƒÇN XONG
                    bat_dau_an_str,  # TH·ªúI GIAN B·∫ÆT ƒê·∫¶U ƒÇN
                    "ƒê·∫£m b·∫£o v·ªá sinh",  # D·ª§NG C·ª§ CHIA, CH·ª®A ƒê·ª∞NG
                    "",  # KI·ªÇM TRA C·∫¢M QUAN TH·ª®C ƒÇN
                    '',  # Kh√¥ng ƒë·∫°t
                    ''   # Ghi ch√∫
                ]
                for j, value in enumerate(data_row4, 1):
                    cell = ws4.cell(row=row_num, column=j, value=value)
                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    cell.border = thin_border
                    if j == 1:
                        cell.font = Font(bold=True, color="006600")
                        cell.fill = PatternFill(start_color="F0FFF0", end_color="F0FFF0", fill_type="solid")
                    elif j == 3:
                        cell.font = Font(bold=True, size=10)
                        cell.fill = PatternFill(start_color="F0FFF0", end_color="F0FFF0", fill_type="solid")
                row_num += 1
                stt += 1
            # Th·ªëng k√™ ph·ª•c v·ª•
            stats_row4 = row_num + 2
            ws4[f'A{stats_row4}'] = "TH·ªêNG K√ä PH·ª§C V·ª§ TH·ª®C ƒÇN:"
            ws4[f'A{stats_row4}'].font = Font(bold=True, size=11, color="006600")
            ws4[f'A{stats_row4}'].fill = PatternFill(start_color="E6FFE6", end_color="E6FFE6", fill_type="solid")
            total_servings = stt - 1
            total_portions = total_servings * students_today
            stats_info4 = [
                f"‚Ä¢ T·ªïng s·ªë l·∫ßn ph·ª•c v·ª•: {total_servings} l·∫ßn",
                f"‚Ä¢ T·ªïng s·ªë su·∫•t ƒÉn ph·ª•c v·ª•: {total_portions} su·∫•t",
                f"‚Ä¢ Trung b√¨nh su·∫•t/l·∫ßn: {(total_portions/total_servings):.1f} su·∫•t/l·∫ßn" if total_servings else "‚Ä¢ Trung b√¨nh su·∫•t/l·∫ßn: N/A",
                f"‚Ä¢ Th·ªùi gian trung b√¨nh t·ª´ ph·ª•c v·ª• ƒë·∫øn ƒÉn: <15 ph√∫t"
            ]
            for i, stat in enumerate(stats_info4, 1):
                ws4[f'A{stats_row4 + i}'] = stat
                ws4[f'A{stats_row4 + i}'].font = Font(size=10)
            # Nguy√™n t·∫Øc ph·ª•c v·ª•
            principles_row4 = stats_row4 + 6
            ws4[f'A{principles_row4}'] = "NGUY√äN T·∫ÆC PH·ª§C V·ª§ AN TO√ÄN:"
            ws4[f'A{principles_row4}'].font = Font(bold=True, size=11, color="004000")
            principles_notes4 = [
                "‚Ä¢ ƒê·∫£m b·∫£o v·ªá sinh d·ª•ng c·ª•, khu v·ª±c ƒÉn tr∆∞·ªõc khi ph·ª•c v·ª•",
                "‚Ä¢ Ki·ªÉm tra nhi·ªát ƒë·ªô th·ª©c ƒÉn tr∆∞·ªõc khi cho tr·∫ª ƒÉn",
                "‚Ä¢ ƒê·∫£m b·∫£o tr·∫ª r·ª≠a tay s·∫°ch s·∫Ω tr∆∞·ªõc khi ƒÉn",
                "‚Ä¢ B√°o c√°o ngay n·∫øu ph√°t hi·ªán b·∫•t th∆∞·ªùng v·ªÅ th·ª©c ƒÉn ho·∫∑c s·ª©c kh·ªèe tr·∫ª"
            ]
            for i, note in enumerate(principles_notes4, 1):
                ws4[f'A{principles_row4 + i}'] = note
                ws4[f'A{principles_row4 + i}'].font = Font(size=9, color="004000")
            # Ch·ªØ k√Ω
            signature_row4 = principles_row4 + 8
            signature_data4 = [
                (signature_row4,     'D', "B·∫æP TR∆Ø·ªûNG",  'H', "NV. Y T·∫æ",   'K', "HI·ªÜU TR∆Ø·ªûNG"),
                (signature_row4 + 1, 'D', "(K√Ω, ghi r√µ h·ªç t√™n)", 'H', "(K√Ω, ghi r√µ h·ªç t√™n)", 'K', "(K√Ω, ghi r√µ h·ªç t√™n)"),
                (signature_row4 + 5, 'D', "Ho√†ng Thanh Tu·∫•n",'H', "",  'K', "Nguy·ªÖn Th·ªã V√¢n")
            ]
            for item in signature_data4:
                row, col_d, text_d, col_h, text_h, col_k, text_k = item
                ws4[f'{col_d}{row}'] = text_d
                ws4[f'{col_h}{row}'] = text_h
                ws4[f'{col_k}{row}'] = text_k
                cols = [(col_d, text_d), (col_h, text_h), (col_k, text_k)]
                for col, text in cols:
                    cell = ws4[f'{col}{row}']
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    if row == signature_row4:
                        cell.font = Font(bold=True, size=12, color="006600")
                        cell.fill = PatternFill(start_color="F0FFF0", end_color="F0FFF0", fill_type="solid")
                    elif row == signature_row4 + 1:
                        cell.font = Font(italic=True, size=9)
                    elif row == signature_row4 + 5:
                        cell.font = Font(bold=True, size=11)
                    else:
                        cell.font = Font(size=9)
            
            # Page setup cho in A4
            ws4.page_setup.orientation = ws4.ORIENTATION_LANDSCAPE
            ws4.page_setup.paperSize = ws4.PAPERSIZE_A4
            ws4.page_setup.fitToPage = True
            ws4.page_setup.fitToWidth = 1
            ws4.page_setup.fitToHeight = 0
            ws4.print_options.horizontalCentered = True
            ws4.page_margins.left = 0.5
            ws4.page_margins.right = 0.5
            ws4.page_margins.top = 0.75
            ws4.page_margins.bottom = 0.75
        
        if 'Sheet' in wb4.sheetnames:
            wb4.remove(wb4['Sheet'])
        file4_buffer = BytesIO()
        wb4.save(file4_buffer)
        file4_buffer.seek(0)
        zipf.writestr(f"B∆∞·ªõc 3 - Ki·ªÉm tra tr∆∞·ªõc khi ƒÉn - Tu·∫ßn {week_number}.xlsx", file4_buffer.read())
        
        # B∆Ø·ªöC 5: Ki·ªÉm tra tr∆∞·ªõc khi ƒÉn - m·ªói ng√†y 1 sheet,
        wb5 = Workbook()
        for day_offset in range(6):
            day_date = week_start + timedelta(days=day_offset)
            day_key = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat'][day_offset]
            safe_date = day_date.strftime('%d-%m')
            sheet_title = f"{days_vn[day_offset]} ({safe_date})"
            if day_offset == 0:
                ws5 = wb5.active
                ws5.title = sheet_title
            else:
                ws5 = wb5.create_sheet(title=sheet_title)
            # L·∫•y s·ªë h·ªçc sinh c√≥ m·∫∑t ng√†y n√†y
            students_today = daily_attendance.get(day_key, 0)
            # Header ch√≠nh t∆∞∆°ng t·ª± ws3
            ws5['A1'] = "T√äN C∆† S·ªû: MNƒêL C√¢y Nh·ªè"
            ws5['A1'].font = Font(bold=True, size=12)
            ws5['A1'].fill = PatternFill(start_color="E6FFE6", end_color="E6FFE6", fill_type="solid")
            ws5.merge_cells('A1:O1')
            ws5['D2'] = "BI·ªÇU M·∫™U THEO D√ïI L∆ØU V√Ä HU·ª∂ L∆ØU M·∫™U TH·ª®C ƒÇN L∆ØU"
            ws5['D2'].font = Font(bold=True, size=14, color="006600")
            ws5['D2'].alignment = Alignment(horizontal='center', vertical='center')
            ws5.merge_cells('D2:I2')
            ws5['J2'] = "S·ªë: 1246/Qƒê - B·ªô Y T·∫ø"
            ws5['J2'].font = Font(bold=True, size=10)
            ws5['J2'].fill = PatternFill(start_color="CCFFCC", end_color="CCFFCC", fill_type="solid")
            info_data4 = [
                (3, 'A', f"Ng∆∞·ªùi ki·ªÉm tra: Nguy·ªÖn Th·ªã V√¢n", 'J', "M·∫´u s·ªë 5"),
                (4, 'A', f"Ng√†y ki·ªÉm tra: {day_date.strftime('%d/%m/%Y')} - {days_vn[day_offset]}", 'J', f"S·ªë h·ªçc sinh: {students_today}"),
                (5, 'A', "ƒê·ªãa ƒëi·ªÉm: Ph√≤ng ƒÉn - MNƒêL C√¢y Nh·ªè", 'F', f"Ng√†y ti·∫øp ph·∫©m: {day_date.strftime('%d/%m/%Y')} - {days_vn[day_offset]}")
            ]
            for row, col_a, text_a, col_m, text_m in info_data4:
                ws5[f'{col_a}{row}'] = text_a
                ws5[f'{col_a}{row}'].font = Font(bold=True, size=10)
                ws5[f'{col_m}{row}'] = text_m
                ws5[f'{col_m}{row}'].font = Font(bold=True, size=10)
                ws5[f'{col_m}{row}'].fill = PatternFill(start_color="F0F8FF", end_color="F0F8FF", fill_type="solid")

            
            headers5_main = [
                'STT', 'CA/B·ªÆA ƒÇN', 'T√äN M·∫™U TH·ª®C ƒÇN', 'S·ªê SU·∫§T ƒÇN\n(ph·∫ßn)',
                'KH·ªêI L∆Ø·ª¢NG\n/ TH·ªÇ T√çCH M·∫™U(GRAM/ML)', 'D·ª§NG C·ª§ CH·ª®A\n M·∫™U TH·ª®C ƒÇN L∆ØU',
                'NHI·ªÜT ƒê·ªò B·∫¢O QU·∫¢N M·∫™U',
                'TH·ªúI GIAN L·∫§Y M·∫™U\n (gi·ªù, ng√†y, th√°ng, nƒÉm)', 'TH·ªúI GIAN HU·ª∂ M·∫™U\n (gi·ªù, ng√†y, th√°ng, nƒÉm)', 
                'GHI CH√ö', "NG∆Ø·ªúI L∆ØU M·∫™U", "NG∆Ø·ªúI HU·ª∂ M·∫™U"
            ]
            for i, header in enumerate(headers5_main, 1):
                cell = ws5.cell(row=8, column=i, value=header)
                cell.font = Font(bold=True, size=9, color="FFFFFF")
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                cell.fill = PatternFill(start_color="8B0000", end_color="8B0000", fill_type="solid")
                cell.border = thick_border
            sub_headers5 = [
                '', '', '', '', 
                '', '', 
                '', 
                '', '', '',
                '',''
            ]
            for i, header in enumerate(sub_headers5, 1):
                cell = ws5.cell(row=9, column=i, value=header)
                cell.font = Font(bold=True, size=8)
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.fill = PatternFill(start_color="CD5C5C", end_color="CD5C5C", fill_type="solid")
                cell.border = thin_border
            ws5.column_dimensions['E'].width = 18
            ws5.column_dimensions['F'].width = 18
            ws5.column_dimensions['G'].width = 18
            ws5.column_dimensions['H'].width = 18
            ws5.column_dimensions['I'].width = 18
            for col in ['E', 'F', 'G','H', 'I']:
                cell = ws5[f'{col}8']
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            # ws5.merge_cells('H8:I8')
            for i in range(1, 13):
                cell = ws5.cell(row=10, column=i, value=i)
                cell.font = Font(bold=True, size=8)
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.fill = PatternFill(start_color="F0F8FF", end_color="F0F8FF", fill_type="solid")
                cell.border = thin_border
            # Ghi d·ªØ li·ªáu m√≥n ƒÉn t·ª´ng ng√†y (gi·ªëng ws3)
            row_num = 11
            stt = 1
            for meal_key, meal_name in meal_times.items():
                dishes = []
                if menu_data[day_key].get(meal_key):
                    dishes = [d.strip() for d in menu_data[day_key][meal_key].split(',') if d.strip()]
                if dishes:
                    dish_names = ', '.join([dish.title() for dish in dishes])
                else:
                    dish_names = ''
                date_str = day_date.strftime('%d/%m/%Y')
                time_chia_xong = '10:15'  # Gi·∫£ ƒë·ªãnh gi·ªù chia xong
                time_bat_dau_an = '10:30'  # Gi·∫£ ƒë·ªãnh gi·ªù b·∫Øt ƒë·∫ßu ƒÉn
                chia_xong_str = f"{date_str} {time_chia_xong}"
                bat_dau_an_str = f"{date_str} {time_bat_dau_an}"
                # T√≠nh th·ªùi gian hu·ª∑ m·∫´u: 15:00 ng√†y h√¥m sau c·ªßa bat_dau_an_str
                from datetime import datetime, timedelta
                dt_batdau = datetime.strptime(bat_dau_an_str.split()[0], "%d/%m/%Y")
                dt_huy = dt_batdau + timedelta(days=1)
                huy_mau_str = f"15:00, {dt_huy.strftime('%d/%m/%Y')}"
                
                data_row5 = [
                    stt,  # STT
                    meal_name,  # CA/B·ªÆA ƒÇN
                    dish_names,  # T√äN M·∫™U TH·ª®C ƒÇN
                    students_today,  # S·ªê SU·∫§T ƒÇN
                    150,  # KH·ªêI L∆Ø·ª¢NG/TH·ªÇ T√çCH M·∫™U
                    "H·ªôp Inox chuy√™n d·ª•ng",  # D·ª§NG C·ª§ CH·ª®A M·∫™U
                    "2-4¬∞C",  # NHI·ªÜT ƒê·ªò B·∫¢O QU·∫¢N
                    bat_dau_an_str,  # TH·ªúI GIAN L·∫§Y M·∫™U
                    huy_mau_str,  # TH·ªúI GIAN HU·ª∂ M·∫™U
                    'Ngon',  # GHI CH√ö
                    "Ho√†ng Thanh Tu·∫•n",  # NG∆Ø·ªúI L∆ØU M·∫™U
                    "Ho√†ng Thanh Tu·∫•n"  # NG∆Ø·ªúI HU·ª∂ M·∫™U
                ]
                for j, value in enumerate(data_row5, 1):
                    cell = ws5.cell(row=row_num, column=j, value=value)
                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    cell.border = thin_border
                    if j == 1:
                        cell.font = Font(bold=True, color="006600")
                        cell.fill = PatternFill(start_color="F0FFF0", end_color="F0FFF0", fill_type="solid")
                    elif j == 3:
                        cell.font = Font(bold=True, size=10)
                        cell.fill = PatternFill(start_color="F0FFF0", end_color="F0FFF0", fill_type="solid")
                row_num += 1
                stt += 1
            # Th·ªëng k√™ ph·ª•c v·ª•
            stats_row5 = row_num + 2
            ws5[f'A{stats_row5}'] = "TH·ªêNG K√ä PH·ª§C V·ª§ TH·ª®C ƒÇN:"
            ws5[f'A{stats_row5}'].font = Font(bold=True, size=11, color="006600")
            ws5[f'A{stats_row5}'].fill = PatternFill(start_color="E6FFE6", end_color="E6FFE6", fill_type="solid")
            total_servings = stt - 1
            total_portions = total_servings * students_today
            stats_info5 = [
                f"‚Ä¢ T·ªïng s·ªë l·∫ßn ph·ª•c v·ª•: {total_servings} l·∫ßn",
                f"‚Ä¢ T·ªïng s·ªë su·∫•t ƒÉn ph·ª•c v·ª•: {total_portions} su·∫•t",
                f"‚Ä¢ Trung b√¨nh su·∫•t/l·∫ßn: {(total_portions/total_servings):.1f} su·∫•t/l·∫ßn" if total_servings else "‚Ä¢ Trung b√¨nh su·∫•t/l·∫ßn: N/A",
                f"‚Ä¢ Th·ªùi gian trung b√¨nh t·ª´ ph·ª•c v·ª• ƒë·∫øn ƒÉn: <15 ph√∫t"
            ]
            for i, stat in enumerate(stats_info5, 1):
                ws5[f'A{stats_row5 + i}'] = stat
                ws5[f'A{stats_row5 + i}'].font = Font(size=10)
            # Nguy√™n t·∫Øc ph·ª•c v·ª•
            principles_row5 = stats_row5 + 6
            ws5[f'A{principles_row5}'] = "NGUY√äN T·∫ÆC PH·ª§C V·ª§ AN TO√ÄN:"
            ws5[f'A{principles_row5}'].font = Font(bold=True, size=11, color="004000")
            principles_notes5 = [
                "‚Ä¢ ƒê·∫£m b·∫£o v·ªá sinh d·ª•ng c·ª•, khu v·ª±c ƒÉn tr∆∞·ªõc khi ph·ª•c v·ª•",
                "‚Ä¢ Ki·ªÉm tra nhi·ªát ƒë·ªô th·ª©c ƒÉn tr∆∞·ªõc khi cho tr·∫ª ƒÉn",
                "‚Ä¢ ƒê·∫£m b·∫£o tr·∫ª r·ª≠a tay s·∫°ch s·∫Ω tr∆∞·ªõc khi ƒÉn",
                "‚Ä¢ B√°o c√°o ngay n·∫øu ph√°t hi·ªán b·∫•t th∆∞·ªùng v·ªÅ th·ª©c ƒÉn ho·∫∑c s·ª©c kh·ªèe tr·∫ª"
            ]
            for i, note in enumerate(principles_notes5, 1):
                ws5[f'A{principles_row5 + i}'] = note
                ws5[f'A{principles_row5 + i}'].font = Font(size=9, color="004000")
            # Ch·ªØ k√Ω
            signature_row5 = principles_row5 + 8
            signature_data5 = [
                (signature_row5,     'D', "NG∆Ø·ªúI TH·ª∞C HI·ªÜN L∆ØU M·∫™U",  'H', "NG∆Ø·ªúI TH·ª∞C HI·ªÜN HU·ª∂ M·∫™U",   'K', "HI·ªÜU TR∆Ø·ªûNG"),
                (signature_row5 + 1, 'D', "(K√Ω, ghi r√µ h·ªç t√™n)", 'H', "(K√Ω, ghi r√µ h·ªç t√™n)", 'K', "(K√Ω, ghi r√µ h·ªç t√™n)"),
                (signature_row5 + 5, 'D', "Ho√†ng Thanh Tu·∫•n",'H', "Ho√†ng Thanh Tu·∫•n",  'K', "Nguy·ªÖn Th·ªã V√¢n")
            ]
            for item in signature_data5:
                row, col_d, text_d, col_h, text_h, col_k, text_k = item
                ws5[f'{col_d}{row}'] = text_d
                ws5[f'{col_h}{row}'] = text_h
                ws5[f'{col_k}{row}'] = text_k
                cols = [(col_d, text_d), (col_h, text_h), (col_k, text_k)]
                for col, text in cols:
                    cell = ws5[f'{col}{row}']
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    if row == signature_row5:
                        cell.font = Font(bold=True, size=12, color="006600")
                        cell.fill = PatternFill(start_color="F0FFF0", end_color="F0FFF0", fill_type="solid")
                    elif row == signature_row5 + 1:
                        cell.font = Font(italic=True, size=9)
                    elif row == signature_row5 + 5:
                        cell.font = Font(bold=True, size=11)
                    else:
                        cell.font = Font(size=9)
            
            # Page setup cho in A4
            ws5.page_setup.orientation = ws5.ORIENTATION_LANDSCAPE
            ws5.page_setup.paperSize = ws5.PAPERSIZE_A4
            ws5.page_setup.fitToPage = True
            ws5.page_setup.fitToWidth = 1
            ws5.page_setup.fitToHeight = 0
            ws5.print_options.horizontalCentered = True
            ws5.page_margins.left = 0.5
            ws5.page_margins.right = 0.5
            ws5.page_margins.top = 0.75
            ws5.page_margins.bottom = 0.75
        
        if 'Sheet' in wb5.sheetnames:
            wb5.remove(wb5['Sheet'])
        file5_buffer = BytesIO()
        wb5.save(file5_buffer)
        file5_buffer.seek(0)
        zipf.writestr(f"B∆∞·ªõc 4 - Theo d√µi l∆∞u v√† hu·ª∑ m·∫´u th·ª©c ƒÉn l∆∞u - Tu·∫ßn {week_number}.xlsx", file5_buffer.read())
        
        # B∆Ø·ªöC 6: PHI·∫æU TI·∫æP NH·∫¨N V√Ä KI·ªÇM TRA CH·∫§T L∆Ø·ª¢NG TH·ª∞C PH·∫®M
        wb6 = Workbook()
        for day_offset in range(6):
            day_date = week_start + timedelta(days=day_offset)
            day_key = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat'][day_offset]
            safe_date = day_date.strftime('%d-%m')
            sheet_title = f"{days_vn[day_offset]} ({safe_date})"
            if day_offset == 0:
                ws6 = wb6.active
                ws6.title = sheet_title
            else:
                ws6 = wb6.create_sheet(title=sheet_title)
            # L·∫•y s·ªë h·ªçc sinh c√≥ m·∫∑t ng√†y n√†y
            students_today = daily_attendance.get(day_key, 0)
            # Header ch√≠nh t∆∞∆°ng t·ª± ws6
            
            ws6['D2'] = "PHI·∫æU TI·∫æP NH·∫¨N V√Ä KI·ªÇM TRA CH·∫§T L∆Ø·ª¢NG TH·ª∞C PH·∫®M"
            ws6['D2'].font = Font(bold=True, size=14, color="006600")
            ws6['D2'].alignment = Alignment(horizontal='center', vertical='center')
            ws6.merge_cells('D2:I2')
            ws6['J2'] = ""
            ws6['J2'].font = Font(bold=True, size=10)
            ws6['J2'].fill = PatternFill(start_color="CCFFCC", end_color="CCFFCC", fill_type="solid")
            info_data4 = [
                (1, 'A', f"Ph√≤ng GD&ƒêT: X√É ƒê·ª®C TR·ªåNG", 'J', f"Ng√†y: {day_date.strftime('%d/%m/%Y')}"),
                (2, 'A', f"ƒê∆°n v·ªã: M·∫¶M NON C√ÇY NH·ªé", 'J', f"Th·ª©: {days_vn[day_offset]}"),
                (3, 'A', f"S·ªë su·∫•t: {students_today}", 'F', "")
            ]
            for row, col_a, text_a, col_m, text_m in info_data4:
                ws6[f'{col_a}{row}'] = text_a
                ws6[f'{col_a}{row}'].font = Font(bold=True, size=10)
                ws6[f'{col_m}{row}'] = text_m
                ws6[f'{col_m}{row}'].font = Font(bold=True, size=10)
                ws6[f'{col_m}{row}'].fill = PatternFill(start_color="F0F8FF", end_color="F0F8FF", fill_type="solid")
            # Th√™m d√≤ng ti√™u ƒë·ªÅ l·ªõn ph√≠a tr√™n b·∫£ng
            ws6['A5'] = "I. Ti·∫øp nh·∫≠n, ki·ªÉm tra ch·∫•t l∆∞·ª£ng th·ª±c ph·∫©m v√† ch·∫ø bi·∫øn"
            ws6['A5'].font = Font(bold=True, size=12, color="8B0000")
            ws6['A5'].fill = PatternFill(start_color="FFF2E6", end_color="FFF2E6", fill_type="solid")
            ws6.merge_cells('A5:G5')

            headers6_main = [
                'STT', 'T√äN TH·ª∞C PH·∫®M',
                'ƒê∆†N V·ªä T√çNH', 'S·ªê L∆Ø·ª¢NG D·ª∞ KI·∫æN MUA',
                'TH·ª∞C T·∫æ TI·∫æP NH·∫¨N', 'GI√Å TI·ªÄN (VNƒê)',
                'NH·∫¨N X√âT'
            ]
            for i, header in enumerate(headers6_main, 1):
                cell = ws6.cell(row=6, column=i, value=header)
                cell.font = Font(bold=True, size=9, color="FFFFFF")
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                cell.fill = PatternFill(start_color="8B0000", end_color="8B0000", fill_type="solid")
                cell.border = thick_border
            sub_headers6 = [
                '', '', '',
                '', '', '', ''
            ]
            for i, header in enumerate(sub_headers6, 1):
                cell = ws6.cell(row=7, column=i, value=header)
                cell.font = Font(bold=True, size=8)
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.fill = PatternFill(start_color="CD5C5C", end_color="CD5C5C", fill_type="solid")
                cell.border = thin_border
            ws6.column_dimensions['B'].width = 18
            ws6.column_dimensions['C'].width = 18
            ws6.column_dimensions['D'].width = 18
            ws6.column_dimensions['E'].width = 18
            ws6.column_dimensions['F'].width = 15
            ws6.column_dimensions['G'].width = 20
            for col in ['B', 'C', 'D', 'E', 'F', 'G']:
                cell = ws6[f'{col}6']
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            # ws6.merge_cells('H6:I6')
            for i in range(1, 8):   
                cell = ws6.cell(row=8, column=i, value=i)
                cell.font = Font(bold=True, size=8)
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.fill = PatternFill(start_color="F0F8FF", end_color="F0F8FF", fill_type="solid")
                cell.border = thin_border
            # Ghi d·ªØ li·ªáu m√≥n ƒÉn t·ª´ng ng√†y
            row_num = 9
            stt = 1
            daily_total_cost = 0  # T·ªïng chi ph√≠ trong ng√†y
            # T·ªïng h·ª£p nguy√™n li·ªáu trong ng√†y (ch·ªâ 1 l·∫ßn cho ng√†y hi·ªán t·∫°i)
            daily_ingredients = {}
            for meal in menu_data[day_key].values():
                if not meal: continue
                for dish_name in [d.strip() for d in meal.split(',') if d.strip()]:
                    dish = Dish.query.filter_by(name=dish_name).first()
                    if dish:
                        for di in dish.ingredients:
                            key = (di.product.name, di.unit, di.product.category, di.product.supplier)
                            qty = di.quantity * students_today
                            if key not in daily_ingredients:
                                daily_ingredients[key] = {'total_qty': 0, 'unit': di.unit, 'category': di.product.category, 'supplier': di.product.supplier, 'product': di.product}
                            daily_ingredients[key]['total_qty'] += qty

            for (name, unit, category, supplier), info in daily_ingredients.items():
                # Quy ƒë·ªïi ƒë∆°n v·ªã n·∫øu c·∫ßn ƒë·ªÉ hi·ªÉn th·ªã
                qty = info['total_qty']  # S·ªë l∆∞·ª£ng t√≠nh b·∫±ng ƒë∆°n v·ªã g·ªëc (gram/ml t·ª´ m√≥n ƒÉn)
                display_unit = unit
                if unit and unit.lower() in ['g', 'gram', 'gr']:
                    display_unit = 'kg'
                    display_qty = round(qty / 1000, 2)
                elif unit and unit.lower() in ['ml', 'milil√≠t', 'milliliter']:
                    display_unit = 'l√≠t' 
                    display_qty = round(qty / 1000, 2)
                else:
                    display_qty = round(qty, 2)

                # T√≠nh gi√° ti·ªÅn: product.price t∆∞∆°ng ·ª©ng v·ªõi product.unit
                # N·∫øu product.unit l√† kg/l√≠t ‚Üí d√πng display_qty (ƒë√£ quy ƒë·ªïi)
                # N·∫øu product.unit l√† gram/ml ‚Üí d√πng qty (ch∆∞a quy ƒë·ªïi)
                product = info['product']  # Product object
                total_price = 0
                if product and product.price:
                    product_unit_lower = (product.unit or '').lower()
                    # Ki·ªÉm tra xem product.unit l√† ƒë∆°n v·ªã l·ªõn (kg/l√≠t) hay nh·ªè (gram/ml)
                    if product_unit_lower in ['kg', 'kilogram', 'l√≠t', 'lit', 'liter', 'l']:
                        # product.price l√† gi√°/kg ho·∫∑c gi√°/l√≠t ‚Üí d√πng display_qty
                        total_price = round(product.price * display_qty, 0)
                    else:
                        # product.price l√† gi√°/gram ho·∫∑c gi√°/ml ‚Üí d√πng qty
                        total_price = round(product.price * qty, 0)
                    price_display = f"{total_price:,.0f} ƒë"
                    daily_total_cost += total_price  # C·ªông v√†o t·ªïng chi ph√≠
                else:
                    price_display = "Ch∆∞a c√≥ gi√°"

                # name ·ªü ƒë√¢y l√† T√äN TH·ª∞C PH·∫®M ch·ªâ l·∫•y theo ng√†y hi·ªán t·∫°i, kh√¥ng ph·∫£i c·∫£ tu·∫ßn
                data_row6 = [
                    stt,  # STT
                    name,  # T√äN TH·ª∞C PH·∫®M
                    display_unit,  # ƒê∆†N V·ªä
                    display_qty,  # S·ªê L∆Ø·ª¢NG D·ª∞ KI·∫æN MUA
                    '',  # TH·ª∞C T·∫æ TI·∫æP NH·∫¨N (ƒë·ªÉ tr·ªëng)
                    price_display,  # GI√Å TI·ªÄN (VNƒê)
                    "",  # Nh·∫≠n x√©t
                ]
                for j, value in enumerate(data_row6, 1):
                    cell = ws6.cell(row=row_num, column=j, value=value)
                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    cell.border = thin_border
                    if j == 1:
                        cell.font = Font(bold=True, color="006600")
                        cell.fill = PatternFill(start_color="F0FFF0", end_color="F0FFF0", fill_type="solid")
                    elif j == 2:
                        cell.font = Font(bold=True, size=10)
                        cell.fill = PatternFill(start_color="F0FFF0", end_color="F0FFF0", fill_type="solid")
                    elif j == 6:  # C·ªôt gi√° ti·ªÅn
                        if total_price > 0:
                            cell.font = Font(bold=True, color="FF6600")
                            cell.fill = PatternFill(start_color="FFF8E1", end_color="FFF8E1", fill_type="solid")
                        else:
                            cell.font = Font(italic=True, color="999999")
                row_num += 1
                stt += 1
            
            # Th√™m d√≤ng t·ªïng c·ªông
            if daily_total_cost > 0:
                ws6.merge_cells(f'A{row_num}:E{row_num}')
                ws6[f'A{row_num}'] = "T·ªîNG CHI PH√ç D·ª∞ KI·∫æN"
                ws6[f'A{row_num}'].font = Font(bold=True, size=11, color="8B0000")
                ws6[f'A{row_num}'].alignment = Alignment(horizontal='right', vertical='center')
                ws6[f'A{row_num}'].fill = PatternFill(start_color="FFE6E6", end_color="FFE6E6", fill_type="solid")
                
                ws6[f'F{row_num}'] = f"{daily_total_cost:,.0f} ƒë"
                ws6[f'F{row_num}'].font = Font(bold=True, size=12, color="8B0000")
                ws6[f'F{row_num}'].alignment = Alignment(horizontal='center', vertical='center')
                ws6[f'F{row_num}'].fill = PatternFill(start_color="FFE6E6", end_color="FFE6E6", fill_type="solid")
                
                ws6[f'G{row_num}'] = ""
                ws6[f'G{row_num}'].fill = PatternFill(start_color="FFE6E6", end_color="FFE6E6", fill_type="solid")
                
                # Add borders
                for col in range(1, 8):
                    ws6.cell(row=row_num, column=col).border = thick_border
                
                row_num += 1
            # N·ªôi Dung kh√°c
            stats_row6 = row_num + 2
            ws6[f'A{stats_row6}'] = "II. N·ªôi Dung Kh√°c"
            ws6[f'A{stats_row6}'].font = Font(bold=True, size=12, color="8B0000")
            ws6[f'A{stats_row6}'].fill = PatternFill(start_color="FFF2E6", end_color="FFF2E6", fill_type="solid")
            # Ch·ªØ k√Ω
            signature_row6 = stats_row6 + 8
            signature_data6 = [
                (signature_row6,     'A', "NG∆Ø·ªúI GIAO H√ÄNG", "C",'NG∆Ø·ªúI TI·∫æP NH·∫¨N', "E", "NV. Y T·∫æ",   'H', "HI·ªÜU TR∆Ø·ªûNG"),
                (signature_row6 + 1, 'A', "(K√Ω, ghi r√µ h·ªç t√™n)", "C",'(K√Ω, ghi r√µ h·ªç t√™n)', 'E', "(K√Ω, ghi r√µ h·ªç t√™n)", 'H', "(K√Ω, ghi r√µ h·ªç t√™n)"),
                (signature_row6 + 5, 'A', "","C","","E", "",  'H', "Nguy·ªÖn Th·ªã V√¢n")
            ]
            for item in signature_data6:
                # unpack 9 values: row, col_d, text_d, col_f, text_f, col_h, text_h, col_k, text_k
                row, col_d, text_d, col_f, text_f, col_h, text_h, col_k, text_k = item
                ws6[f'{col_d}{row}'] = text_d
                ws6[f'{col_f}{row}'] = text_f
                ws6[f'{col_h}{row}'] = text_h
                ws6[f'{col_k}{row}'] = text_k
                cols = [(col_d, text_d), (col_f, text_f), (col_h, text_h), (col_k, text_k)]
                for col, text in cols:
                    cell = ws6[f'{col}{row}']
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    if row == signature_row6:
                        cell.font = Font(bold=True, size=12, color="006600")
                        cell.fill = PatternFill(start_color="F0FFF0", end_color="F0FFF0", fill_type="solid")
                    elif row == signature_row6 + 1:
                        cell.font = Font(italic=True, size=9)
                    elif row == signature_row6 + 5:
                        cell.font = Font(bold=True, size=11)
                    else:
                        cell.font = Font(size=9)
            
            # Page setup cho in A4
            ws6.page_setup.orientation = ws6.ORIENTATION_PORTRAIT
            ws6.page_setup.paperSize = ws6.PAPERSIZE_A4
            ws6.page_setup.fitToPage = True
            ws6.page_setup.fitToWidth = 1
            ws6.page_setup.fitToHeight = 0
            ws6.print_options.horizontalCentered = True
            ws6.page_margins.left = 0.5
            ws6.page_margins.right = 0.5
            ws6.page_margins.top = 0.75
            ws6.page_margins.bottom = 0.75
        
        if 'Sheet' in wb6.sheetnames:
            wb6.remove(wb6['Sheet'])
        file6_buffer = BytesIO()
        wb6.save(file6_buffer)
        file6_buffer.seek(0)
        zipf.writestr(f"B∆∞·ªõc 6 - PHI·∫æU TI·∫æP NH·∫¨N V√Ä KI·ªÇM TRA CH·∫§T L∆Ø·ª¢NG TH·ª∞C PH·∫®M - Tu·∫ßn {week_number}.xlsx", file6_buffer.read())
        
    # ƒê√≥ng zipfile v√† tr·∫£ v·ªÅ
    zip_buffer.seek(0)
    
    # T·∫°o response ƒë·ªÉ download
    response = send_file(
        zip_buffer,
        as_attachment=True,
        download_name=f"Quy_trinh_an_toan_thuc_pham_3_buoc_Tuan_{week_number}.zip",
        mimetype='application/zip'
    )
    
    flash(f'ƒê√£ xu·∫•t th√†nh c√¥ng quy tr√¨nh an to√†n th·ª±c ph·∫©m 3 b∆∞·ªõc cho tu·∫ßn {week_number}!', 'success')
    return response

# ================== QU·∫¢N L√ù NH√Ä CUNG C·∫§P V√Ä S·∫¢N PH·∫®M ==================

@main.route('/suppliers')
def suppliers():
    """Danh s√°ch nh√† cung c·∫•p"""
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    
    suppliers = Supplier.query.filter_by(is_active=True).order_by(Supplier.name).all()
    return render_template('suppliers.html', suppliers=suppliers)

@main.route('/suppliers/new', methods=['GET', 'POST'])
def new_supplier():
    """Th√™m nh√† cung c·∫•p m·ªõi"""
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    
    form = SupplierForm()
    if form.validate_on_submit():
        supplier = Supplier(
            name=form.name.data,
            address=form.address.data,
            phone=form.phone.data,
            contact_person=form.contact_person.data,
            supplier_type=form.supplier_type.data,
            registration_number=form.registration_number.data,
            food_safety_cert=form.food_safety_cert.data,
            created_date=datetime.utcnow()
        )
        db.session.add(supplier)
        db.session.commit()
        log_activity('create', 'supplier', supplier.id, f'T·∫°o nh√† cung c·∫•p: {form.name.data}')
        flash('Th√™m nh√† cung c·∫•p th√†nh c√¥ng!', 'success')
        return redirect(url_for('main.suppliers'))
    
    return render_template('new_supplier.html', form=form)

@main.route('/suppliers/<int:supplier_id>/edit', methods=['GET', 'POST'])
def edit_supplier(supplier_id):
    """S·ª≠a th√¥ng tin nh√† cung c·∫•p"""
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    
    supplier = Supplier.query.get_or_404(supplier_id)
    form = SupplierForm(obj=supplier)
    
    if form.validate_on_submit():
        form.populate_obj(supplier)
        db.session.commit()
        log_activity('edit', 'supplier', supplier_id, f'C·∫≠p nh·∫≠t nh√† cung c·∫•p: {supplier.name}')
        flash('C·∫≠p nh·∫≠t nh√† cung c·∫•p th√†nh c√¥ng!', 'success')
        return redirect(url_for('main.suppliers'))
    
    return render_template('edit_supplier.html', form=form, supplier=supplier)

@main.route('/suppliers/<int:supplier_id>/delete', methods=['POST'])
def delete_supplier(supplier_id):
    """X√≥a nh√† cung c·∫•p"""
    if session.get('role') != 'admin':
        return redirect_no_permission()
    
    supplier = Supplier.query.get_or_404(supplier_id)
    supplier_name = supplier.name
    supplier.is_active = False
    db.session.commit()
    log_activity('delete', 'supplier', supplier_id, f'X√≥a nh√† cung c·∫•p: {supplier_name}')
    flash('X√≥a nh√† cung c·∫•p th√†nh c√¥ng!', 'success')
    return redirect(url_for('main.suppliers'))

@main.route('/products')
def products():
    """Danh s√°ch s·∫£n ph·∫©m"""
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    
    products = Product.query.filter_by(is_active=True).join(Supplier).order_by(Product.category, Product.name).all()
    return render_template('products.html', products=products)

@main.route('/products/new', methods=['GET', 'POST'])
def new_product():
    """Th√™m s·∫£n ph·∫©m m·ªõi"""
    current_role = session.get('role')
    
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    
    form = ProductForm()
    # L·∫•y danh s√°ch ƒë∆°n v·ªã duy nh·∫•t t·ª´ Product
    product_units = sorted(list(set([p.unit for p in Product.query.all() if p.unit])))
    # L·∫•y danh s√°ch nh√† cung c·∫•p cho dropdown
    suppliers = Supplier.query.filter_by(is_active=True).order_by(Supplier.name).all()
    
    # N·∫øu ch∆∞a c√≥ supplier n√†o, t·∫°o m·ªôt supplier m·∫´u
    if not suppliers:
        default_supplier = Supplier(
            name="Nh√† cung c·∫•p m·∫∑c ƒë·ªãnh",
            address="ƒê·ªãa ch·ªâ c·∫ßn c·∫≠p nh·∫≠t",
            phone="0123456789",
            contact_person="Ng∆∞·ªùi li√™n h·ªá",
            supplier_type="fresh",
            registration_number="",
            food_safety_cert="",
            created_date=datetime.utcnow()
        )
        db.session.add(default_supplier)
        db.session.commit()
        suppliers = [default_supplier]
        flash('ƒê√£ t·∫°o nh√† cung c·∫•p m·∫∑c ƒë·ªãnh. Vui l√≤ng c·∫≠p nh·∫≠t th√¥ng tin sau!', 'info')
    
    form.supplier_id.choices = [(s.id, s.name) for s in suppliers]
    
    if form.validate_on_submit():
        product = Product(
            name=form.name.data,
            category=form.category.data,
            supplier_id=form.supplier_id.data,
            unit=form.unit.data,
            price=form.price.data
        )
        db.session.add(product)
        db.session.commit()
        log_activity('create', 'product', product.id, f'T·∫°o s·∫£n ph·∫©m: {form.name.data}')
        flash('Th√™m s·∫£n ph·∫©m th√†nh c√¥ng!', 'success')
        return redirect(url_for('main.products'))
    return render_template('new_product.html', form=form, product_units=product_units)

@main.route('/products/<int:product_id>/edit', methods=['GET', 'POST'])
def edit_product(product_id):
    """S·ª≠a th√¥ng tin s·∫£n ph·∫©m"""
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    
    product = Product.query.get_or_404(product_id)
    form = ProductForm(obj=product)
    product_units = sorted(list(set([p.unit for p in Product.query.all() if p.unit])))
    
    # L·∫•y danh s√°ch nh√† cung c·∫•p cho dropdown
    suppliers = Supplier.query.filter_by(is_active=True).order_by(Supplier.name).all()
    form.supplier_id.choices = [(s.id, s.name) for s in suppliers]
    
    if form.validate_on_submit():
        form.populate_obj(product)
        db.session.commit()
        log_activity('edit', 'product', product_id, f'C·∫≠p nh·∫≠t s·∫£n ph·∫©m: {product.name}')
        flash('C·∫≠p nh·∫≠t s·∫£n ph·∫©m th√†nh c√¥ng!', 'success')
        return redirect(url_for('main.products'))
    
    return render_template('edit_product.html', form=form, product=product, suppliers=suppliers, product_units=product_units)

@main.route('/products/<int:product_id>/delete', methods=['POST'])
def delete_product(product_id):
    """X√≥a s·∫£n ph·∫©m"""
    if session.get('role') != 'admin' and session.get('role') != 'teacher':
        return redirect_no_permission()
    
    product = Product.query.get_or_404(product_id)
    product_name = product.name
    product.is_active = False
    db.session.commit()
    log_activity('delete', 'product', product_id, f'X√≥a s·∫£n ph·∫©m: {product_name}')
    flash('X√≥a s·∫£n ph·∫©m th√†nh c√¥ng!', 'success')
    return redirect(url_for('main.products'))

# ============== AI Routes v·ªõi LLM Farm ==============

@main.route('/ai/menu-suggestions', methods=['POST'])
def ai_menu_suggestions():
    """API endpoint ƒë·ªÉ l·∫•y g·ª£i √Ω th·ª±c ƒë∆°n t·ª´ Gemini AI - SECURED & OPTIMIZED"""
    
    user_role = session.get('role')
    if user_role not in ['admin', 'teacher']:
        return jsonify({'success': False, 'error': 'Kh√¥ng c√≥ quy·ªÅn truy c·∫≠p. Vui l√≤ng ƒëƒÉng nh·∫≠p v·ªõi t√†i kho·∫£n admin ho·∫∑c gi√°o vi√™n.'}), 403

    try:
        from app.models import Dish
        import random
        meal_types = ["morning", "snack", "dessert", "lunch", "afternoon", "lateafternoon"]
        days = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat']
        dishes = Dish.query.filter_by(is_active=True).all()
        dishes_by_meal = {meal: [] for meal in meal_types}
        for d in dishes:
            if d.meal_times:
                for meal in d.meal_times:
                    if meal in meal_types:
                        dishes_by_meal[meal].append(d.name)

        menu = {}
        used_dishes = {meal: set() for meal in meal_types}  # Theo d√µi m√≥n ƒë√£ d√πng trong tu·∫ßn cho t·ª´ng b·ªØa
        for day in days:
            menu[day] = {}
            for meal in meal_types:
                meal_dishes = dishes_by_meal[meal][:]
                if not meal_dishes:
                    menu[day][meal] = "[Kh√¥ng c√≥]"
                    continue
                # Lo·∫°i c√°c m√≥n ƒë√£ d√πng h·∫øt l∆∞·ª£t trong tu·∫ßn (∆∞u ti√™n kh√¥ng l·∫∑p)
                available = [d for d in meal_dishes if d not in used_dishes[meal]]
                if meal == "lunch":
                    # Tr∆∞a: 2 m√≥n kh√°c nhau, tr√°nh tr√πng trong tu·∫ßn
                    if len(available) >= 2:
                        selected = random.sample(available, 2)
                    elif len(meal_dishes) >= 2:
                        # N·∫øu ƒë√£ d√πng h·∫øt, cho ph√©p l·∫∑p l·∫°i nh∆∞ng v·∫´n ch·ªçn 2 m√≥n kh√°c nhau
                        selected = random.sample(meal_dishes, 2)
                    elif len(meal_dishes) == 1:
                        selected = [meal_dishes[0], meal_dishes[0]]
                    else:
                        selected = ["[Kh√¥ng c√≥]", "[Kh√¥ng c√≥]"]
                    menu[day][meal] = ", ".join(selected)
                    for s in selected:
                        used_dishes[meal].add(s)
                else:
                    # C√°c b·ªØa kh√°c: 1 m√≥n, tr√°nh tr√πng trong tu·∫ßn
                    if available:
                        selected = random.choice(available)
                    else:
                        selected = random.choice(meal_dishes)
                    menu[day][meal] = selected
                    used_dishes[meal].add(selected)
        return jsonify({
            'success': True,
            'menu': menu,
            'suggestions': {
                'nutrition_tips': [
                    f"Th·ª±c ƒë∆°n ƒë∆∞·ª£c t·ªëi ∆∞u cho tr·∫ª {request.json.get('age_group', '1-3 tu·ªïi')}",
                    "ƒê·∫£m b·∫£o c√¢n b·∫±ng dinh d∆∞·ª°ng v·ªõi ƒë·∫ßy ƒë·ªß nh√≥m th·ª±c ph·∫©m",
                    "Tr√°nh l·∫∑p l·∫°i m√≥n ƒÉn trong c√πng b·ªØa trong tu·∫ßn",
                    "B·ªØa tr∆∞a c√≥ 2 m√≥n ƒë·ªÉ tƒÉng ƒëa d·∫°ng dinh d∆∞·ª°ng",
                    "Khuy·∫øn kh√≠ch tr·∫ª th·ª≠ nhi·ªÅu lo·∫°i th·ª±c ph·∫©m kh√°c nhau"
                ],
                'generated_from': 'SmallTree AI - Nutrition optimized menu system'
            },
            'security_info': f"Optimized menu for {user_role} with nutrition balance",
        })
    except Exception as e:
        db.session.rollback()
        print(f"[ERROR] {e}")
        return jsonify({'success': False, 'error': str(e)})

def extract_weekly_menu_from_suggestions(suggestions):
    """Tr√≠ch xu·∫•t v√† chuy·ªÉn ƒë·ªïi suggestions th√†nh format menu database"""
    menu_data = {}
    days = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat']
    slots = ['morning', 'snack', 'dessert', 'lunch', 'afternoon', 'lateafternoon']
    
    # Initialize empty menu
    for day in days:
        menu_data[day] = {}
        for slot in slots:
            menu_data[day][slot] = "M√≥n ƒÉn dinh d∆∞·ª°ng"
    
    current_day = None
    current_day_index = -1
    
    for suggestion in suggestions:
        suggestion = suggestion.strip()
        
        # T√¨m ng√†y
        if '**Th·ª©' in suggestion:
            if 'Th·ª© 2' in suggestion:
                current_day = 'mon'
                current_day_index = 0
            elif 'Th·ª© 3' in suggestion:
                current_day = 'tue' 
                current_day_index = 1
            elif 'Th·ª© 4' in suggestion:
                current_day = 'wed'
                current_day_index = 2
            elif 'Th·ª© 5' in suggestion:
                current_day = 'thu'
                current_day_index = 3
            elif 'Th·ª© 6' in suggestion:
                current_day = 'fri'
                current_day_index = 4
            elif 'Th·ª© 7' in suggestion:
                current_day = 'sat'
                current_day_index = 5
            continue
            
        # T√¨m m√≥n ƒÉn theo khung gi·ªù
        if current_day and suggestion.startswith('‚Ä¢'):
            suggestion = suggestion[1:].strip()  # B·ªè bullet point
            
            if suggestion.startswith('S√°ng:'):
                menu_data[current_day]['morning'] = suggestion[5:].strip()
            elif suggestion.startswith('Ph·ª• s√°ng:'):
                menu_data[current_day]['snack'] = suggestion[9:].strip()
            elif suggestion.startswith('Tr√°ng mi·ªáng:'):
                menu_data[current_day]['dessert'] = suggestion[12:].strip()
            elif suggestion.startswith('Tr∆∞a:'):
                menu_data[current_day]['lunch'] = suggestion[5:].strip()
            elif suggestion.startswith('X·∫ø:'):
                menu_data[current_day]['afternoon'] = suggestion[3:].strip()
            elif suggestion.startswith('X·∫ø chi·ªÅu:'):
                menu_data[current_day]['lateafternoon'] = suggestion[9:].strip()
    
    return menu_data

@main.route('/ai/create-menu-from-suggestions', methods=['POST'])
def create_menu_from_suggestions():
    """API endpoint ƒë·ªÉ t·∫°o th·ª±c ƒë∆°n t·ª´ AI suggestions"""
    
    user_role = session.get('role')
    if user_role not in ['admin', 'teacher']:
        return jsonify({'success': False, 'error': 'Kh√¥ng c√≥ quy·ªÅn truy c·∫≠p'}), 403
    
    # Import modules outside try block to avoid reference errors
    from datetime import datetime, timedelta
    import json

    try:
        
        # L·∫•y d·ªØ li·ªáu t·ª´ request
        data = request.get_json()
        
        if not data or 'menu' not in data:
            print(f"[ERROR] No menu data found. Data keys: {data.keys() if data else 'None'}")
            return jsonify({'success': False, 'error': 'Kh√¥ng c√≥ d·ªØ li·ªáu th·ª±c ƒë∆°n'}), 400
        
        menu_data = data['menu']
        overwrite = data.get('overwrite', False)
        
        # S·ª≠ d·ª•ng tu·∫ßn ƒë∆∞·ª£c ch·ªçn ho·∫∑c tu·∫ßn hi·ªán t·∫°i
        if 'target_week' in data and 'target_year' in data:
            week_number = data['target_week']
            year = data['target_year']
        else:
            # Fallback: t√≠nh tu·∫ßn hi·ªán t·∫°i
            today = datetime.now().date()
            week_start = today - timedelta(days=today.weekday())
            week_number = week_start.isocalendar()[1]
            year = week_start.year
        
        # Ki·ªÉm tra xem th·ª±c ƒë∆°n tu·∫ßn n√†y ƒë√£ t·ªìn t·∫°i ch∆∞a (s·ª≠ d·ª•ng model Menu)
        existing_menu = Menu.query.filter_by(week_number=week_number, year=year).first()
        if existing_menu and not overwrite:
            return jsonify({
                'success': False,
                'error': f'Th·ª±c ƒë∆°n tu·∫ßn {week_number}/{year} ƒë√£ t·ªìn t·∫°i',
                'week_number': week_number,
                'existing_id': existing_menu.id
            }), 409
        
        # T·∫°o ho·∫∑c c·∫≠p nh·∫≠t th·ª±c ƒë∆°n
        if existing_menu and overwrite:
            menu_obj = existing_menu

        else:
            menu_obj = Menu(week_number=week_number, year=year)
            db.session.add(menu_obj)

        
        # C·∫≠p nh·∫≠t d·ªØ li·ªáu th·ª±c ƒë∆°n
        for day in ['mon', 'tue', 'wed', 'thu', 'fri', 'sat']:
            day_data = menu_data.get(day, {})
  
            
            if day == 'mon':
                menu_obj.monday_morning = day_data.get('morning', '')
                menu_obj.monday_snack = day_data.get('snack', '')
                menu_obj.monday_dessert = day_data.get('dessert', '')
                menu_obj.monday_lunch = day_data.get('lunch', '')
                menu_obj.monday_afternoon = day_data.get('afternoon', '')
                menu_obj.monday_lateafternoon = day_data.get('lateafternoon', '')
            elif day == 'tue':
                menu_obj.tuesday_morning = day_data.get('morning', '')
                menu_obj.tuesday_snack = day_data.get('snack', '')
                menu_obj.tuesday_dessert = day_data.get('dessert', '')
                menu_obj.tuesday_lunch = day_data.get('lunch', '')
                menu_obj.tuesday_afternoon = day_data.get('afternoon', '')
                menu_obj.tuesday_lateafternoon = day_data.get('lateafternoon', '')
            elif day == 'wed':
                menu_obj.wednesday_morning = day_data.get('morning', '')
                menu_obj.wednesday_snack = day_data.get('snack', '')
                menu_obj.wednesday_dessert = day_data.get('dessert', '')
                menu_obj.wednesday_lunch = day_data.get('lunch', '')
                menu_obj.wednesday_afternoon = day_data.get('afternoon', '')
                menu_obj.wednesday_lateafternoon = day_data.get('lateafternoon', '')
            elif day == 'thu':
                menu_obj.thursday_morning = day_data.get('morning', '')
                menu_obj.thursday_snack = day_data.get('snack', '')
                menu_obj.thursday_dessert = day_data.get('dessert', '')
                menu_obj.thursday_lunch = day_data.get('lunch', '')
                menu_obj.thursday_afternoon = day_data.get('afternoon', '')
                menu_obj.thursday_lateafternoon = day_data.get('lateafternoon', '')
            elif day == 'fri':
                menu_obj.friday_morning = day_data.get('morning', '')
                menu_obj.friday_snack = day_data.get('snack', '')
                menu_obj.friday_dessert = day_data.get('dessert', '')
                menu_obj.friday_lunch = day_data.get('lunch', '')
                menu_obj.friday_afternoon = day_data.get('afternoon', '')
                menu_obj.friday_lateafternoon = day_data.get('lateafternoon', '')
            elif day == 'sat':
                menu_obj.saturday_morning = day_data.get('morning', '')
                menu_obj.saturday_snack = day_data.get('snack', '')
                menu_obj.saturday_dessert = day_data.get('dessert', '')
                menu_obj.saturday_lunch = day_data.get('lunch', '')
                menu_obj.saturday_afternoon = day_data.get('afternoon', '')
                menu_obj.saturday_lateafternoon = day_data.get('lateafternoon', '')
        
        db.session.commit()
 
        return jsonify({
            'success': True,
            'message': f'ƒê√£ {"c·∫≠p nh·∫≠t" if overwrite and existing_menu else "t·∫°o"} th·ª±c ƒë∆°n tu·∫ßn {week_number}/{year} th√†nh c√¥ng!',
            'week_number': week_number,
            'year': year,
            'menu_id': menu_obj.id,
            'overwritten': overwrite and existing_menu
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"[ERROR] create_menu_from_suggestions: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'L·ªói server: {str(e)}'}), 500


@main.route('/ai-dashboard', methods=['GET', 'POST'])
def ai_dashboard():
    """Trang dashboard AI v·ªõi c√°c t√≠nh nƒÉng LLM Farm"""
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()

    # Kh√¥ng t·ª± t·∫°o prompt ·ªü ƒë√¢y n·ªØa, ch·ªâ render dashboard, prompt s·∫Ω l·∫•y t·ª´ API /ai/menu-suggestions
    return render_template('ai_dashboard.html')

# ===== STUDENT ALBUM MANAGEMENT ROUTES =====

@main.route('/student-albums')
def student_albums():
    """Danh s√°ch album c·ªßa t·∫•t c·∫£ h·ªçc sinh"""
    user_role = session.get('role')
    user_id = session.get('user_id')
    students = []
    albums = []
    if user_role == 'parent':
        # Ch·ªâ xem album c·ªßa con m√¨nh
        child = Child.query.filter_by(id=user_id).first()
        if child:
            students = [child]
            albums = StudentAlbum.query.filter_by(student_id=child.id).order_by(StudentAlbum.date_created.desc()).all()
        else:
            students = []
            albums = []
    else:
        # Gi√°o vi√™n, admin xem t·∫•t c·∫£
        students = Child.query.all()
        import os
        updated = False
        for s in students:
            old_avatar = s.avatar
            # S·ª≠a l·∫°i ƒë∆∞·ªùng d·∫´n avatar n·∫øu ph√°t hi·ªán sai ƒë·ªãnh d·∫°ng
            if s.avatar:
                new_avatar = s.avatar
                # Remove app/static/ prefix n·∫øu c√≥
                if new_avatar.startswith('app/static/'):
                    new_avatar = new_avatar.replace('app/static/', '')
                elif new_avatar.startswith('/app/static/'):
                    new_avatar = new_avatar.replace('/app/static/', '')
                # Convert backslashes to forward slashes  
                new_avatar = new_avatar.replace('\\', '/')
                # Ensure proper path structure
                if not new_avatar.startswith('images/students/') and new_avatar:
                    if new_avatar.startswith('students/'):
                        new_avatar = 'images/' + new_avatar
                    elif not new_avatar.startswith('images/'):
                        new_avatar = 'images/students/' + new_avatar
                
                if new_avatar != old_avatar:
                    s.avatar = new_avatar
                    updated = True
                    print(f"[AUTO-FIX] {s.name} ({s.student_code}): {old_avatar} ‚Üí {new_avatar}")
            
            # N·∫øu avatar l√† None/r·ªóng th√¨ t√¨m file theo student_code ho·∫∑c student_id
            if (not s.avatar or s.avatar.strip() == '') and s.student_code:
                import glob
                # T√¨m theo student_code tr∆∞·ªõc (format m·ªõi)
                pattern = os.path.join('app', 'static', 'images', 'students', f'student_{s.student_code}_*')
                matches = glob.glob(pattern)
                
                # N·∫øu kh√¥ng t√¨m th·∫•y, th·ª≠ t√¨m theo student ID (format c≈©)
                if not matches:
                    pattern = os.path.join('app', 'static', 'images', 'students', f'student_{s.id}_*')
                    matches = glob.glob(pattern)
                
                if matches:
                    # L·∫•y t√™n file ƒë·∫ßu ti√™n t√¨m ƒë∆∞·ª£c
                    rel_path = os.path.relpath(matches[0], os.path.join('app', 'static'))
                    s.avatar = rel_path.replace('\\', '/')
                    updated = True
                    print(f"[AUTO-DETECT] {s.name} ({s.student_code}): Found avatar {s.avatar}")
        if updated:
            db.session.commit()
        albums = StudentAlbum.query.join(Child).order_by(StudentAlbum.date_created.desc()).all()
    # ƒê·∫£m b·∫£o students v√† albums lu√¥n l√† list
    students = students or []
    albums = albums or []
    # T√≠nh th·ªëng k√™
    today = date.today()
    albums_today = [album for album in albums if album.date_created == today]
    academic_albums = [album for album in albums if album.milestone_type == 'academic']
    # Mobile detection
    user_agent = request.headers.get('User-Agent', '').lower()
    mobile = any(device in user_agent for device in ['mobile', 'android', 'iphone'])
    return render_template('student_albums.html', 
                         students=students, 
                         albums=albums, 
                         albums_today=albums_today,
                         academic_albums=academic_albums,
                         mobile=mobile)

@main.route('/student/<int:student_id>/albums')
def student_albums_detail(student_id):
    """Album chi ti·∫øt c·ªßa m·ªôt h·ªçc sinh"""
    if session.get('role') not in ['admin', 'teacher', 'parent']:
        return redirect_no_permission()
    
    student = Child.query.get_or_404(student_id)
    
    # N·∫øu l√† parent, ch·ªâ xem ƒë∆∞·ª£c album c·ªßa con m√¨nh
    if session.get('role') == 'parent' and session.get('user_id') != student_id:
        flash('B·∫°n ch·ªâ c√≥ th·ªÉ xem album c·ªßa con m√¨nh!', 'error')
        return redirect(url_for('main.index'))
    
    albums = StudentAlbum.query.filter_by(student_id=student_id).order_by(StudentAlbum.date_created.desc()).all()
    progress_records = StudentProgress.query.filter_by(student_id=student_id).order_by(StudentProgress.evaluation_date.desc()).all()
    
    # T√≠nh th·ªëng k√™ cho student n√†y
    total_photos = sum(len(album.photos) for album in albums)
    today = date.today()
    first_day_of_month = today.replace(day=1)
    albums_this_month = [album for album in albums if album.date_created >= first_day_of_month]
    
    # Mobile detection
    user_agent = request.headers.get('User-Agent', '').lower()
    mobile = any(device in user_agent for device in ['mobile', 'android', 'iphone'])
    
    return render_template('student_album_detail.html', 
                         student=student, 
                         albums=albums, 
                         progress_records=progress_records,
                         total_photos=total_photos,
                         albums_this_month=albums_this_month,
                         mobile=mobile,
                         current_date=date.today().strftime('%Y-%m-%d'))

@main.route('/student/<int:student_id>/album/new', methods=['GET', 'POST'])
def create_student_album(student_id):
    """T·∫°o album m·ªõi cho h·ªçc sinh"""
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    
    student = Child.query.get_or_404(student_id)
    
    if request.method == 'POST':
        # L·∫•y th√¥ng tin album
        title = request.form.get('title')
        description = request.form.get('description', '')
        milestone_type = request.form.get('milestone_type', 'other')
        school_year = request.form.get('school_year', '')
        semester = request.form.get('semester', '')
        age_at_time = request.form.get('age_at_time', '')
        
        # T·∫°o album m·ªõi
        album = StudentAlbum(
            student_id=student_id,
            title=title,
            description=description,
            date_created=date.today(),
            milestone_type=milestone_type,
            school_year=school_year,
            semester=semester,
            age_at_time=age_at_time,
            created_by=session.get('username', 'teacher'),
            is_shared_with_parents=True
        )
        
        db.session.add(album)
        db.session.flush()  # ƒê·ªÉ l·∫•y album.id
        
        # Kh·ªüi t·∫°o R2 storage
        r2 = get_r2_storage()
        
        # X·ª≠ l√Ω upload ·∫£nh
        uploaded_files = request.files.getlist('photos')
        if uploaded_files:
            upload_dir = os.path.join(current_app.static_folder, 'student_albums', str(student_id), str(album.id))
            os.makedirs(upload_dir, exist_ok=True)
            
            for i, file in enumerate(uploaded_files):
                if file and file.filename:
                    filename = secrets.token_hex(16) + '.' + file.filename.rsplit('.', 1)[1].lower()
                    r2_key = f"albums/{student_id}/{album.id}/{filename}"
                    
                    # Upload to R2
                    r2_url = r2.upload_file(file, r2_key)
                    
                    # Fallback to local if R2 fails
                    if not r2_url:
                        filepath = os.path.join(upload_dir, filename)
                        file.save(filepath)
                        file_path_or_url = f"student_albums/{student_id}/{album.id}/{filename}"
                        file_size = os.path.getsize(filepath)
                    else:
                        file_path_or_url = r2_url
                        file_size = len(file.read())
                        file.seek(0)  # Reset file pointer
                    
                    # T·∫°o record ·∫£nh
                    photo = StudentPhoto(
                        album_id=album.id,
                        filename=filename,
                        filepath=file_path_or_url,
                        original_filename=file.filename,
                        caption=request.form.get(f'caption_{i}', ''),
                        upload_date=datetime.now(),
                        file_size=file_size,
                        image_order=i,
                        is_cover_photo=(i == 0)  # ·∫¢nh ƒë·∫ßu ti√™n l√†m ·∫£nh ƒë·∫°i di·ªán
                    )
                    db.session.add(photo)
        
        db.session.commit()
        log_activity('create', 'album', album.id, f'T·∫°o album "{title}" cho h·ªçc sinh {student.name}')
        flash(f'‚úÖ ƒê√£ t·∫°o album "{title}" cho {student.name}!', 'success')
        return redirect(url_for('main.student_albums_detail', student_id=student_id))
    
    # Mobile detection
    user_agent = request.headers.get('User-Agent', '').lower()
    mobile = any(device in user_agent for device in ['mobile', 'android', 'iphone'])
    
    return render_template('create_student_album.html', student=student, mobile=mobile)

@main.route('/album/<int:album_id>')
def view_album(album_id):
    """Xem chi ti·∫øt m·ªôt album"""
    if session.get('role') not in ['admin', 'teacher', 'parent']:
        return redirect_no_permission()
    
    album = StudentAlbum.query.get_or_404(album_id)
    
    # N·∫øu l√† parent, ch·ªâ xem ƒë∆∞·ª£c album c·ªßa con m√¨nh
    if session.get('role') == 'parent' and session.get('user_id') != album.student_id:
        flash('B·∫°n ch·ªâ c√≥ th·ªÉ xem album c·ªßa con m√¨nh!', 'error')
        return redirect(url_for('main.index'))
    
    photos = StudentPhoto.query.filter_by(album_id=album_id).order_by(StudentPhoto.image_order).all()
    
    # Mobile detection
    user_agent = request.headers.get('User-Agent', '').lower()
    mobile = any(device in user_agent for device in ['mobile', 'android', 'iphone'])
    
    return render_template('view_album.html', album=album, photos=photos, mobile=mobile)

@main.route('/student/<int:student_id>/progress/new', methods=['GET', 'POST'])
def add_student_progress(student_id):
    """Th√™m ƒë√°nh gi√° ti·∫øn b·ªô cho h·ªçc sinh"""
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    
    student = Child.query.get_or_404(student_id)
    
    if request.method == 'POST':
        progress = StudentProgress(
            student_id=student_id,
            evaluation_date=datetime.strptime(request.form.get('evaluation_date'), '%Y-%m-%d').date(),
            skill_category=request.form.get('skill_category'),
            skill_name=request.form.get('skill_name'),
            level_achieved=request.form.get('level_achieved'),
            notes=request.form.get('notes', ''),
            teacher_name=session.get('username', 'teacher')
        )
        
        db.session.add(progress)
        db.session.commit()
        
        flash(f'‚úÖ ƒê√£ th√™m ƒë√°nh gi√° ti·∫øn b·ªô cho {student.name}!', 'success')
        return redirect(url_for('main.student_albums_detail', student_id=student_id))
    
    # Mobile detection
    user_agent = request.headers.get('User-Agent', '').lower()
    mobile = any(device in user_agent for device in ['mobile', 'android', 'iphone'])
    
    return render_template('add_student_progress.html', student=student, mobile=mobile)

@main.route('/album/<int:album_id>/delete', methods=['POST'])
def delete_album(album_id):
    """X√≥a album"""
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    
    album = StudentAlbum.query.get_or_404(album_id)
    student_id = album.student_id
    album_title = album.title
    student_name = album.student.name
    
    # Kh·ªüi t·∫°o R2 storage
    r2 = get_r2_storage()
    
    # Ph√¢n lo·∫°i v√† x√≥a ·∫£nh t·ª´ R2 ho·∫∑c local
    r2_images = []
    local_images = []
    
    for photo in album.photos:
        if photo.filepath.startswith('http'):
            r2_images.append(photo.filepath)
        else:
            local_images.append(photo.filepath)
    
    # X√≥a ·∫£nh R2 theo batch
    if r2_images:
        r2.delete_files_batch(r2_images)
    
    # X√≥a th∆∞ m·ª•c local n·∫øu c√≤n
    if local_images:
        album_dir = os.path.join(current_app.static_folder, 'student_albums', str(student_id), str(album_id))
        if os.path.exists(album_dir):
            import shutil
            shutil.rmtree(album_dir)
    
    db.session.delete(album)
    db.session.commit()
    log_activity('delete', 'album', album_id, f'X√≥a album "{album_title}" c·ªßa h·ªçc sinh {student_name}')
    
    flash('‚úÖ ƒê√£ x√≥a album!', 'success')
    return redirect(url_for('main.student_albums_detail', student_id=student_id))

@main.route('/fix-avatars', methods=['GET'])
def fix_avatars():
    """Route ƒë·ªÉ fix t·∫•t c·∫£ avatar paths trong database"""
    if session.get('role') != 'admin':
        return redirect_no_permission()
    
    students = Child.query.all()
    fixed_count = 0
    
    for student in students:
        if student.avatar:
            old_avatar = student.avatar
            # Fix c√°c format sai
            new_avatar = old_avatar
            
            # Remove app/static/ prefix n·∫øu c√≥
            if new_avatar.startswith('app/static/'):
                new_avatar = new_avatar.replace('app/static/', '')
            elif new_avatar.startswith('/app/static/'):
                new_avatar = new_avatar.replace('/app/static/', '')
            
            # Convert backslashes to forward slashes
            new_avatar = new_avatar.replace('\\', '/')
            
            # Ensure starts with images/students/
            if not new_avatar.startswith('images/students/'):
                if new_avatar.startswith('students/'):
                    new_avatar = 'images/' + new_avatar
                elif not new_avatar.startswith('images/'):
                    new_avatar = 'images/students/' + new_avatar
            
            if new_avatar != old_avatar:
                student.avatar = new_avatar
                fixed_count += 1
                print(f"[FIX] {student.name} ({student.student_code}): {old_avatar} ‚Üí {new_avatar}")
    
    if fixed_count > 0:
        db.session.commit()
        flash(f'‚úÖ ƒê√£ fix {fixed_count} avatar paths!', 'success')
    else:
        flash('‚úÖ T·∫•t c·∫£ avatar paths ƒë√£ ƒë√∫ng format!', 'info')
    
    return redirect(url_for('main.student_albums'))

# Route debug upload limits cho 60-70 ·∫£nh
@main.route('/debug-upload-test')
def debug_upload_test():
    if session.get('role') not in ['admin', 'teacher']:
        return redirect_no_permission()
    
    return '''
    <h2>üîç Test Upload Logic - 60-70 ·∫£nh</h2>
    <p><strong>MAX_CONTENT_LENGTH:</strong> ''' + str(current_app.config.get('MAX_CONTENT_LENGTH', 0) // (1024*1024)) + '''MB</p>
    
    <h3>Test Upload Traditional (< 30 ·∫£nh)</h3>
    <form id="traditionalForm" enctype="multipart/form-data" action="/debug-process-upload" method="post">
        <input type="file" name="test_files" multiple accept="image/*" id="traditionalInput">
        <br><br>
        <button type="submit">Test Traditional Upload</button>
        <div id="traditionalResults"></div>
    </form>
    
    <hr>
    
    <h3>Test Client-Side Compression (>= 30 ·∫£nh)</h3>
    <form id="compressionForm" enctype="multipart/form-data">
        <input type="file" name="test_files" multiple accept="image/*" id="compressionInput">
        <br><br>
        <button type="button" onclick="testCompression()">Test Compression Upload</button>
        <div id="compressionResults"></div>
    </form>
    
    <script>
    function testCompression() {
        const files = document.getElementById('compressionInput').files;
        document.getElementById('compressionResults').innerHTML = 
            `<p>üîÑ Testing compression for ${files.length} files...</p>
             <p><strong>Logic:</strong> ${files.length >= 30 ? 'Client-side compression' : 'Traditional upload'}</p>`;
        
        if (files.length >= 30) {
            alert('Would use client-side compression for ' + files.length + ' files!');
        } else {
            alert('Would use traditional upload for ' + files.length + ' files!');
        }
    }
    </script>
    '''

@main.route('/debug-process-upload', methods=['POST'])
def debug_process_upload():
    if session.get('role') not in ['admin', 'teacher']:
        return jsonify({'error': 'Unauthorized'}), 403
    
    files = request.files.getlist('test_files')
    file_count = len(files)
    
    results = {
        'file_count': file_count,
        'method': 'traditional' if file_count < 30 else 'client_compression',
        'max_content_mb': current_app.config.get('MAX_CONTENT_LENGTH', 0) // (1024*1024),
        'files_with_content': len([f for f in files if f.filename]),
        'status': 'success'
    }
    
    return f'''
    <h3>‚úÖ Upload Test Results</h3>
    <pre>{json.dumps(results, indent=2)}</pre>
    <p><a href="/debug-upload-test">‚Üê Back to test</a></p>
    '''